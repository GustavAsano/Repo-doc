"""
app/src/doc_gen_export.py

Standalone PDF and Word (docx) export utilities extracted from the original
Streamlit app.py.  No Streamlit dependency — pure Python + fpdf2 + python-docx.
"""

from __future__ import annotations

import os
from datetime import datetime
from io import BytesIO
from pathlib import Path

# ── Module-level render-mode state (mirrors original app.py globals) ─────────
PDF_RENDER_BACKEND: str = "fpdf"
PDF_RENDER_FALLBACK_REASON: str = ""
PDF_RENDER_FONT: str = ""


# ── Text helpers ──────────────────────────────────────────────────────────────

def _sanitize_pdf_text(text: str) -> str:
    replacements = {
        "\u2013": "-", "\u2014": "-",
        "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"',
        "\u2022": "-", "\u2026": "...", "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", "ignore").decode("latin-1")


def _strip_inline_markdown(text: str) -> str:
    import re
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = re.sub(r"[*_]{1,3}(.+?)[*_]{1,3}", r"\1", text)
    text = text.replace("`", "")
    return text


def _normalize_doc_language(language: str | None) -> str:
    value = (language or "").strip().upper()
    if value.startswith("PT"): return "PT-BR"
    if value.startswith("EN"): return "EN-US"
    if value.startswith("ES"): return "ES-ES"
    if value.startswith("FR"): return "FR-FR"
    if value.startswith("DE"): return "DE-DE"
    return "EN-US"


def _markdown_to_pdf_blocks(md_text: str) -> list[tuple[str, str]]:
    import re
    lines = md_text.splitlines()
    blocks: list[tuple[str, str]] = []
    in_code = False
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue
        if in_code:
            blocks.append(("code", _sanitize_pdf_text(line)))
            continue
        if not stripped:
            blocks.append(("blank", ""))
            continue
        if re.match(r"^\s*([-*_]\s*){3,}$", line):
            blocks.append(("hr", ""))
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading:
            level = len(heading.group(1))
            text = _sanitize_pdf_text(_strip_inline_markdown(heading.group(2).strip()))
            blocks.append((f"h{level}", text))
            continue
        bullet = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if bullet:
            indent = len(bullet.group(1))
            marker = bullet.group(2).strip()
            text = _sanitize_pdf_text(_strip_inline_markdown(bullet.group(3).strip()))
            blocks.append(("li", f"{indent}\x1f{marker}\x1f{text}"))
            continue
        if stripped.startswith(">"):
            text = _sanitize_pdf_text(_strip_inline_markdown(stripped[1:].strip()))
            blocks.append(("quote", text))
            continue
        blocks.append(("p", _sanitize_pdf_text(_strip_inline_markdown(stripped))))
    return blocks


def _normalize_pdf_blocks(blocks: list[tuple[str, str]]) -> list[tuple[str, str]]:
    merged: list[tuple[str, str]] = []
    buffer: list[str] = []
    code_buffer: list[str] = []
    for kind, text in blocks:
        if kind == "code":
            code_buffer.append(text)
            continue
        if code_buffer:
            merged.append(("code", "\n".join(code_buffer)))
            code_buffer = []
        if kind == "p":
            buffer.append(text)
            continue
        if buffer:
            merged.append(("p", " ".join(buffer)))
            buffer = []
        merged.append((kind, text))
    if buffer:
        merged.append(("p", " ".join(buffer)))
    if code_buffer:
        merged.append(("code", "\n".join(code_buffer)))
    return merged


def _markdown_to_docx_blocks(md_text: str) -> list[tuple[str, str]]:
    import re
    lines = md_text.splitlines()
    blocks: list[tuple[str, str]] = []
    in_code = False
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue
        if in_code:
            blocks.append(("code", line))
            continue
        if not stripped:
            blocks.append(("blank", ""))
            continue
        if re.match(r"^\s*([-*_]\s*){3,}$", line):
            blocks.append(("hr", ""))
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading:
            level = len(heading.group(1))
            text = _strip_inline_markdown(heading.group(2).strip())
            blocks.append((f"h{level}", text))
            continue
        bullet = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if bullet:
            indent = len(bullet.group(1))
            marker = bullet.group(2).strip()
            text = _strip_inline_markdown(bullet.group(3).strip())
            blocks.append(("li", f"{indent}\x1f{marker}\x1f{text}"))
            continue
        if stripped.startswith(">"):
            text = _strip_inline_markdown(stripped[1:].strip())
            blocks.append(("quote", text))
            continue
        blocks.append(("p", _strip_inline_markdown(stripped)))
    return blocks


def _split_markdown_sections(md_text: str) -> list[tuple[str, str]]:
    import re
    sections: list[tuple[str, str]] = []
    current_title: str | None = None
    buffer: list[str] = []
    for raw in md_text.splitlines():
        line = raw.rstrip("\n")
        stripped = line.strip()
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue
        heading = re.match(r"^#\s+(.*)", stripped)
        if heading:
            if current_title is not None:
                sections.append((current_title, "\n".join(buffer).strip()))
            current_title = _sanitize_pdf_text(_strip_inline_markdown(heading.group(1).strip()))
            buffer = []
            continue
        if current_title is not None:
            buffer.append(line)
    if current_title is not None:
        sections.append((current_title, "\n".join(buffer).strip()))
    if not sections:
        cleaned = md_text.strip()
        if cleaned:
            sections.append(("Overview", cleaned))
    return sections


def _pdf_i18n(language: str, repo_name: str, generated_date: str, doc_kind: str = "technical") -> dict:
    lang = _normalize_doc_language(language)
    texts = {
        "PT-BR": {
            "doc_subtitle": "Documentação Técnica",
            "generated_on": "Gerado em",
            "abstract_title": "Resumo",
            "abstract_body": (
                f"O objetivo deste documento é detalhar tecnicamente o repositório {repo_name}. "
                f"Esta documentação foi gerada automaticamente em {generated_date} com uso do Repository Documentation "
                "e necessita de revisão humana antes de publicação final."
            ),
            "summary_title": "Sumário",
            "version": "Versão",
            "changed_by": "Alterado por",
            "change_date": "Data da alteração",
            "page_label": "Página",
        },
        "EN-US": {
            "doc_subtitle": "Technical Documentation",
            "generated_on": "Generated on",
            "abstract_title": "Abstract",
            "abstract_body": (
                f"The objective of this document is to provide a technical description of the {repo_name} repository. "
                f"This documentation was automatically generated on {generated_date} using Repository Documentation "
                "and requires human review before final publication."
            ),
            "summary_title": "Summary",
            "version": "Version",
            "changed_by": "Changed by",
            "change_date": "Change date",
            "page_label": "Page",
        },
        "ES-ES": {
            "doc_subtitle": "Documentación Técnica",
            "generated_on": "Generado el",
            "abstract_title": "Resumen",
            "abstract_body": (
                f"El objetivo de este documento es detallar técnicamente el repositorio {repo_name}. "
                f"Esta documentación fue generada automáticamente el {generated_date} usando Repository Documentation "
                "y requiere revisión humana antes de su publicación final."
            ),
            "summary_title": "Resumen",
            "version": "Versión",
            "changed_by": "Modificado por",
            "change_date": "Fecha de cambio",
            "page_label": "Página",
        },
        "FR-FR": {
            "doc_subtitle": "Documentation Technique",
            "generated_on": "Généré le",
            "abstract_title": "Résumé",
            "abstract_body": (
                f"L'objectif de ce document est de décrire techniquement le dépôt {repo_name}. "
                f"Cette documentation a été générée automatiquement le {generated_date} avec Repository Documentation "
                "et nécessite une révision humaine avant publication finale."
            ),
            "summary_title": "Sommaire",
            "version": "Version",
            "changed_by": "Modifié par",
            "change_date": "Date de modification",
            "page_label": "Page",
        },
        "DE-DE": {
            "doc_subtitle": "Technische Dokumentation",
            "generated_on": "Erstellt am",
            "abstract_title": "Zusammenfassung",
            "abstract_body": (
                f"Ziel dieses Dokuments ist die technische Beschreibung des Repositorys {repo_name}. "
                f"Diese Dokumentation wurde am {generated_date} automatisch mit Repository Documentation erstellt "
                "und muss vor der finalen Veröffentlichung manuell geprüft werden."
            ),
            "summary_title": "Inhaltsverzeichnis",
            "version": "Version",
            "changed_by": "Geändert von",
            "change_date": "Änderungsdatum",
            "page_label": "Seite",
        },
    }
    payload = dict(texts.get(lang, texts["EN-US"]))
    payload["changed_by_value"] = "Repository Documentation"

    if str(doc_kind or "technical").strip().lower().startswith("func"):
        func_overrides = {
            "PT-BR": {"doc_subtitle": "Documentação Funcional"},
            "EN-US": {"doc_subtitle": "Functional Documentation"},
            "ES-ES": {"doc_subtitle": "Documentación Funcional"},
            "FR-FR": {"doc_subtitle": "Documentation Fonctionnelle"},
            "DE-DE": {"doc_subtitle": "Funktionale Dokumentation"},
        }
        payload.update(func_overrides.get(lang, func_overrides["EN-US"]))
    return payload


# ── PDF builder (fpdf2 only — no Playwright dep for basic test) ──────────────

def _build_pdf_bytes_fpdf(md_text: str, title: str, language: str = "EN-US", doc_kind: str = "technical") -> bytes:
    global PDF_RENDER_FONT
    try:
        from fpdf import FPDF, XPos, YPos
    except Exception as e:
        raise RuntimeError("fpdf2 is not installed. pip install fpdf2") from e

    lang = _normalize_doc_language(language)
    date_format = "%Y-%m-%d" if lang == "EN-US" else "%d/%m/%Y"
    generated_date = datetime.now().strftime(date_format)
    i18n = _pdf_i18n(lang, title, generated_date, doc_kind=doc_kind)
    blocks = _normalize_pdf_blocks(_markdown_to_pdf_blocks(md_text))

    main_font = "Helvetica"
    PDF_RENDER_FONT = "Helvetica (core)"
    page_margin = 18
    bottom_margin = 18
    paragraph_after_mm = 0.8
    heading_after_mm = 3.6

    class DocPDF(FPDF):
        def footer(self):
            if self.page_no() <= 3:
                return
            self.set_y(-10)
            self.set_font(main_font, size=9)
            self.set_text_color(120, 120, 120)
            self.cell(0, 5, f"{i18n['page_label']} {self.page_no()}", align="R")
            self.set_text_color(0, 0, 0)

    helper = DocPDF()
    helper.set_margins(page_margin, page_margin, page_margin)
    helper.set_auto_page_break(auto=True, margin=bottom_margin)
    helper.set_font(main_font, size=10)
    page_width = helper.w - helper.l_margin - helper.r_margin
    if page_width <= 0:
        page_width = 100

    def safe_width(pdf: DocPDF, width: float) -> float:
        min_w = max(pdf.get_string_width("W"), pdf.get_string_width("m")) + 0.5
        return max(width, min_w)

    def wrap_line(pdf: DocPDF, text: str, width: float) -> list[str]:
        width = safe_width(pdf, width)
        words = text.split(" ")
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip() if current else word
            if pdf.get_string_width(candidate) <= width:
                current = candidate
                continue
            if current:
                lines.append(current)
                current = ""
            if pdf.get_string_width(word) <= width:
                current = word
                continue
            chunk = ""
            for ch in word:
                if pdf.get_string_width(chunk + ch) <= width:
                    chunk += ch
                else:
                    if chunk:
                        lines.append(chunk)
                    chunk = ch
            if chunk:
                current = chunk
        if current:
            lines.append(current)
        return lines

    def render_lines(pdf, lines, line_height=6, width=None, x=None, align="L"):
        use_width = safe_width(pdf, width or page_width)
        for entry in lines:
            if not entry:
                pdf.ln(line_height)
                continue
            if x is not None:
                pdf.set_x(x)
            else:
                pdf.set_x(pdf.l_margin)
            pdf.cell(use_width, line_height, entry, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align=align)

    def render_paragraph_justified(pdf, text, line_height=6):
        use_width = safe_width(pdf, page_width)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(use_width, line_height, text, align="J", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(paragraph_after_mm)

    def render_content(pdf, collect_toc=False, page_offset=0):
        toc_entries = []
        pdf.add_page()
        pdf.set_font(main_font, size=10)
        max_width = page_width
        has_content = False
        for kind, text in blocks:
            if kind == "blank":
                pdf.ln(4)
                continue
            if kind == "hr":
                y = pdf.get_y()
                pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
                pdf.ln(4)
                continue
            if kind.startswith("h"):
                level = int(kind[1:])
                if level == 1 and has_content:
                    pdf.add_page()
                if collect_toc and level in (1, 2):
                    toc_entries.append((level, text, pdf.page_no() + page_offset))
                heading_size = {1: 18, 2: 13, 3: 11, 4: 10.5, 5: 10, 6: 9.5}.get(level, 10.5)
                heading_line_height = 8 if level == 1 else (6 if level == 2 else 5.5)
                pdf.set_font(main_font, "B", heading_size)
                render_lines(pdf, wrap_line(pdf, text, max_width), line_height=heading_line_height, width=max_width)
                pdf.ln(heading_after_mm)
                pdf.set_font(main_font, size=10)
                has_content = True
                continue
            if kind == "code":
                pdf.set_font(main_font, size=9)
                code_lines = []
                for raw in text.splitlines():
                    if not raw.strip():
                        code_lines.append("")
                        continue
                    code_lines.extend(wrap_line(pdf, raw, max_width - 6))
                if not code_lines:
                    code_lines = [""]
                line_height = 5
                block_height = line_height * len(code_lines) + 6
                x = pdf.l_margin
                y = pdf.get_y()
                if y + block_height > pdf.h - pdf.b_margin:
                    pdf.add_page()
                    y = pdf.get_y()
                pdf.set_fill_color(245, 247, 250)
                pdf.set_draw_color(220, 224, 230)
                pdf.rect(x, y, max_width, block_height, "DF")
                inner_x = x + 3
                inner_width = max_width - 6
                pdf.set_xy(inner_x, y + 3)
                for line in code_lines:
                    render_lines(pdf, [line], line_height=line_height, width=inner_width, x=inner_x)
                pdf.set_y(y + block_height + 2)
                pdf.set_x(pdf.l_margin)
                pdf.set_font(main_font, size=10)
                has_content = True
                continue
            if kind == "quote":
                pdf.set_font(main_font, "I", 10)
                render_paragraph_justified(pdf, f"> {text}", line_height=6)
                pdf.set_font(main_font, size=10)
                has_content = True
                continue
            if kind == "li":
                indent, marker, item_text = 0, "-", text
                if "\x1f" in text:
                    parts = text.split("\x1f", 2)
                    if len(parts) == 3:
                        try:
                            indent = max(0, int(parts[0]))
                        except Exception:
                            indent = 0
                        marker = parts[1].strip() or "-"
                        item_text = parts[2]
                list_level = min(indent // 2, 6)
                indent_mm = list_level * 4
                ordered = marker.endswith(".") and marker[:-1].isdigit()
                prefix = f"{marker} " if ordered else "· "
                prefix_width = pdf.get_string_width(prefix)
                line_width = max(10, max_width - indent_mm)
                first_width = max(10, line_width - prefix_width)
                wrapped = wrap_line(pdf, item_text, first_width) or [""]
                x_start = pdf.l_margin + indent_mm
                render_lines(pdf, [prefix + wrapped[0]], width=line_width, x=x_start)
                if len(wrapped) > 1:
                    cont_x = x_start + prefix_width
                    cont_w = max(10, line_width - prefix_width)
                    for cont in wrapped[1:]:
                        render_lines(pdf, [cont], width=cont_w, x=cont_x)
                pdf.ln(paragraph_after_mm)
                has_content = True
                continue
            render_paragraph_justified(pdf, text, line_height=6)
            has_content = True
        return toc_entries

    # Pass 1: TOC page numbers
    probe = DocPDF()
    probe.set_margins(page_margin, page_margin, page_margin)
    probe.set_auto_page_break(auto=True, margin=bottom_margin)
    toc_entries = render_content(probe, collect_toc=True, page_offset=3)

    # Pass 2: final
    pdf = DocPDF()
    pdf.set_margins(page_margin, page_margin, page_margin)
    pdf.set_auto_page_break(auto=True, margin=bottom_margin)

    # Cover
    pdf.add_page()
    pdf.set_font(main_font, "B", 26)
    center_y = (pdf.h / 2) - 10
    pdf.set_y(center_y)
    render_lines(pdf, wrap_line(pdf, title, page_width), line_height=12, width=page_width, align="C")
    pdf.set_font(main_font, size=18)
    render_lines(pdf, [i18n["doc_subtitle"]], line_height=10, width=page_width, align="C")
    pdf.set_y(pdf.h - pdf.b_margin - 6)
    pdf.set_font(main_font, "I", 10)
    pdf.cell(page_width, 4, f"{i18n['generated_on']} {generated_date}", align="R")

    # Back cover
    pdf.add_page()
    pdf.set_font(main_font, "B", 16)
    render_lines(pdf, [i18n["abstract_title"]], line_height=8, width=page_width)
    pdf.ln(heading_after_mm)
    pdf.set_font(main_font, "", 10)
    render_paragraph_justified(pdf, i18n["abstract_body"], line_height=6)

    # Summary
    pdf.add_page()
    pdf.set_font(main_font, "B", 10)
    render_lines(pdf, [i18n["summary_title"]], line_height=7, width=page_width)
    pdf.ln(2)
    pdf.set_font(main_font, size=10)
    toc_page_col = 14.0
    for level, item, page_no in toc_entries:
        label = item.strip()
        indent_mm = max(0, level - 1) * 4.0
        x_start = pdf.l_margin + indent_mm
        usable_width = max(20.0, page_width - indent_mm)
        page_col = min(toc_page_col, usable_width * 0.35)
        left_col = max(10.0, usable_width - page_col - 2.0)
        max_label_width = max(6.0, left_col * 0.62)
        while label and pdf.get_string_width(label) > max_label_width:
            label = label[:-1]
        dot_unit = max(pdf.get_string_width("."), 0.3)
        dot_width = max(0.0, left_col - pdf.get_string_width(label) - 1.0)
        dot_count = max(3, int(dot_width / dot_unit))
        left_text = f"{label} {'.' * dot_count}"
        y = pdf.get_y()
        pdf.set_xy(x_start, y)
        pdf.cell(left_col, 6, left_text, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.cell(page_col, 6, str(page_no), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Content
    render_content(pdf, collect_toc=False, page_offset=0)

    data = pdf.output()
    if isinstance(data, str):
        data = data.encode("latin-1", "ignore")
    if isinstance(data, bytearray):
        data = bytes(data)
    return data


def build_pdf_bytes(md_text: str, title: str, language: str = "EN-US", doc_kind: str = "technical") -> bytes:
    global PDF_RENDER_BACKEND, PDF_RENDER_FALLBACK_REASON, PDF_RENDER_FONT
    PDF_RENDER_BACKEND = "fpdf"
    PDF_RENDER_FALLBACK_REASON = ""
    return _build_pdf_bytes_fpdf(md_text, title, language=language, doc_kind=doc_kind)


def build_docx_bytes(md_text: str, title: str, language: str = "EN-US", doc_kind: str = "technical") -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except Exception as e:
        raise RuntimeError("python-docx is required. pip install python-docx") from e

    lang = _normalize_doc_language(language)
    date_format = "%Y-%m-%d" if lang == "EN-US" else "%d/%m/%Y"
    generated_date = datetime.now().strftime(date_format)
    i18n = _pdf_i18n(lang, title, generated_date, doc_kind=doc_kind)
    blocks = _normalize_pdf_blocks(_markdown_to_docx_blocks(md_text))

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.line_spacing = 1.15

    heading_sizes = {1: 16, 2: 12, 3: 11, 4: 10.5, 5: 10, 6: 9.5}

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(i18n["doc_subtitle"])
    r.font.name = "Calibri"
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(f"{i18n['generated_on']} {generated_date}")
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)

    doc.add_page_break()

    # Abstract
    p = doc.add_paragraph()
    r = p.add_run(i18n["abstract_title"])
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph(i18n["abstract_body"]).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate([i18n["version"], i18n["changed_by"], i18n["change_date"]]):
        table.rows[0].cells[idx].paragraphs[0].add_run(header).bold = True
    for idx, value in enumerate(["1.0", i18n["changed_by_value"], generated_date]):
        table.rows[1].cells[idx].text = value

    doc.add_page_break()

    # Summary
    p = doc.add_paragraph()
    r = p.add_run(i18n["summary_title"])
    r.bold = True
    r.font.size = Pt(16)
    for sec_title, _ in _split_markdown_sections(md_text):
        doc.add_paragraph(sec_title, style="List Bullet")

    doc.add_page_break()

    # Content
    has_content = False
    for kind, text in blocks:
        if kind == "blank":
            doc.add_paragraph("")
            continue
        if kind == "hr":
            doc.add_paragraph("—" * 48)
            continue
        if kind.startswith("h"):
            level = int(kind[1:])
            if level == 1 and has_content:
                doc.add_page_break()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(heading_sizes.get(level, 12))
            p.paragraph_format.space_after = Pt(10)
            has_content = True
            continue
        if kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            has_content = True
            continue
        if kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(text).italic = True
            has_content = True
            continue
        if kind == "li":
            indent, marker, item_text = 0, "-", text
            if "\x1f" in text:
                parts = text.split("\x1f", 2)
                if len(parts) == 3:
                    try:
                        indent = max(0, int(parts[0]))
                    except Exception:
                        indent = 0
                    marker = parts[1].strip() or "-"
                    item_text = parts[2]
            level = min(indent // 2, 6)
            is_ordered = marker.endswith(".") and marker[:-1].isdigit()
            style_name = "List Number" if is_ordered else "List Bullet"
            try:
                p = doc.add_paragraph(style=style_name)
            except Exception:
                p = doc.add_paragraph()
                p.add_run(f"{'•' if not is_ordered else marker} ")
            p.paragraph_format.left_indent = Inches(0.2 * level)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(item_text)
            has_content = True
            continue
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        has_content = True

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
