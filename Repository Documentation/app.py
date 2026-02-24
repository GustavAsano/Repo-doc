import streamlit as st
from pathlib import Path
import json
import os
import yaml
import socket
import uuid
import base64
import shutil
from io import BytesIO
from datetime import datetime, timedelta, timezone
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
import streamlit.components.v1 as components

# import your backend logic
from src.loader import load_repository
from src.repo_mapping import repo_map
from src.doc_gen import (
    generate_doc,
    generate_functional_doc,
    generate_functional_doc_from_technical,
    separate_output,
    apply_standard_ai_disclaimer,
    DOCUMENTATION_SECTIONS as DEFAULT_DOCUMENTATION_SECTIONS,
    FUNCTIONAL_DOCUMENTATION_SECTIONS as DEFAULT_FUNCTIONAL_DOCUMENTATION_SECTIONS,
)
from src.mkdocs_ui import generate_mkdocs_config, serve_mkdocs
from src.llm_utils import init_llm, resolve_api_key, PROVIDER_LABELS

LAST_REPO_FILE = Path.cwd() / ".last_repo"
LLM_CONFIG_FILE = Path.cwd() / ".llm_config.json"
LEGACY_REPO_LIBRARY_FILE = Path.cwd() / ".repo_library.json"
LEGACY_LIBRARY_DIR = Path.cwd() / "biblioteca"
LEGACY_LIBRARY_FILE = LEGACY_LIBRARY_DIR / "library.json"
LIBRARY_DIR = Path.cwd() / "library"
REPO_LIBRARY_FILE = LIBRARY_DIR / "library.json"
REPO_ASSETS_DIR = LIBRARY_DIR / "repos"
FUNCTIONAL_LIBRARY_DIR = Path.cwd() / "library_functional"
FUNCTIONAL_REPO_LIBRARY_FILE = FUNCTIONAL_LIBRARY_DIR / "library.json"
FUNCTIONAL_REPO_ASSETS_DIR = FUNCTIONAL_LIBRARY_DIR / "repos"
LLM_RUNTIME_KEYS = {}
BEDROCK_RUNTIME_ENV_BACKUP: dict[str, str | None] = {}
BEDROCK_RUNTIME_ENV_ACTIVE = False
CHAT_HISTORY_DIR = Path.cwd() / ".chat_history"
MAX_CHAT_HISTORY = 200
TOP_LOGO_PATH = Path(r"C:\Users\rafael.maia.nunes\Pictures\genMS.svg")
MODEL_OPTIONS = {
    "gemini": ["gemini-2.5-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"],
    "openai": ["gpt-5-nano", "gpt-5-mini", "gpt-4o-mini", "gpt-4o"],
    "bedrock": [
        "moonshotai.kimi-k2.5",
        "moonshot.kimi-k2-thinking",
        "deepseek.v3.2",
        "minimax.minimax-m2.1",
        "openai.gpt-oss-120b-1:0",
        "openai.gpt-oss-20b-1:0",
        "zai.glm-4.7",
    ],
}
BEDROCK_DEFAULT_MODEL = "moonshotai.kimi-k2.5"
SOURCE_TECH_LIBRARY = "Technical Library"
SOURCE_TECH_LIBRARY_LEGACY = "Repository Library"
SOURCE_FUNCTIONAL_LIBRARY = "Functional Library"
GEN_MODE_TECHNICAL_ONLY = "Technical only"
GEN_MODE_TECHNICAL_AND_FUNCTIONAL = "Technical + Functional"
GEN_MODE_FUNCTIONAL_ONLY = "Functional only"
GEN_MODE_FUNCTIONAL_ONLY_LEGACY = "Functional only (from existing technical doc)"
PDF_RENDER_BACKEND = "fpdf"
PDF_RENDER_FALLBACK_REASON = ""
PDF_RENDER_FONT = ""
EXPORT_RENDER_CACHE_VERSION = "2026-02-19-7"


def _safe_rerun() -> None:
    rerun_fn = getattr(st, "rerun", None)
    if callable(rerun_fn):
        rerun_fn()
        return
    legacy_rerun = getattr(st, "experimental_rerun", None)
    if callable(legacy_rerun):
        legacy_rerun()


def _normalize_generation_mode_value(value: str | None) -> str:
    mode = str(value or "").strip()
    if mode == GEN_MODE_FUNCTIONAL_ONLY_LEGACY:
        return GEN_MODE_FUNCTIONAL_ONLY
    return mode


def _is_functional_only_mode(value: str | None) -> bool:
    return _normalize_generation_mode_value(value) == GEN_MODE_FUNCTIONAL_ONLY


def _load_llm_config() -> dict:
    try:
        if LLM_CONFIG_FILE.exists():
            data = json.loads(LLM_CONFIG_FILE.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}
    return {}


def _save_llm_config(
    provider: str,
    model: str,
    use_system_key: bool,
    api_key: str | None = None,
    bedrock_access_key: str | None = None,
    bedrock_secret_key: str | None = None,
) -> None:
    try:
        existing = _load_llm_config()
        payload = {
            "provider": provider,
            "model": model,
            "use_system_key": bool(use_system_key),
            "api_key": "" if use_system_key else (api_key or ""),
            "bedrock_access_key": str(existing.get("bedrock_access_key") or ""),
            "bedrock_secret_key": str(existing.get("bedrock_secret_key") or ""),
        }
        if provider == "bedrock":
            payload["bedrock_access_key"] = (
                "" if use_system_key else str(bedrock_access_key or "")
            )
            payload["bedrock_secret_key"] = (
                "" if use_system_key else str(bedrock_secret_key or "")
            )
        LLM_CONFIG_FILE.write_text(json.dumps(payload), encoding="utf-8")
    except Exception:
        pass


def _set_runtime_key(provider: str, api_key: str | None) -> None:
    if not provider:
        return
    if api_key:
        LLM_RUNTIME_KEYS[provider] = api_key
    else:
        LLM_RUNTIME_KEYS.pop(provider, None)


def _get_runtime_key(provider: str) -> str | None:
    if not provider:
        return None
    value = LLM_RUNTIME_KEYS.get(provider)
    return value or None


def _restore_runtime_bedrock_credentials() -> None:
    global BEDROCK_RUNTIME_ENV_BACKUP, BEDROCK_RUNTIME_ENV_ACTIVE
    if not BEDROCK_RUNTIME_ENV_ACTIVE:
        return
    for env_key in ("AWS_ACCESS_KEY_ID", "AWS_SECRET_ACCESS_KEY"):
        previous = BEDROCK_RUNTIME_ENV_BACKUP.get(env_key)
        if previous is None:
            os.environ.pop(env_key, None)
        else:
            os.environ[env_key] = previous
    BEDROCK_RUNTIME_ENV_BACKUP = {}
    BEDROCK_RUNTIME_ENV_ACTIVE = False


def _set_runtime_bedrock_credentials(
    provider: str,
    use_system_key: bool,
    access_key: str | None = None,
    secret_key: str | None = None,
) -> None:
    global BEDROCK_RUNTIME_ENV_BACKUP, BEDROCK_RUNTIME_ENV_ACTIVE
    if provider != "bedrock":
        _restore_runtime_bedrock_credentials()
        return

    access = str(access_key or "").strip()
    secret = str(secret_key or "").strip()
    if use_system_key or not access or not secret:
        _restore_runtime_bedrock_credentials()
        return

    if not BEDROCK_RUNTIME_ENV_ACTIVE:
        BEDROCK_RUNTIME_ENV_BACKUP = {
            "AWS_ACCESS_KEY_ID": os.environ.get("AWS_ACCESS_KEY_ID"),
            "AWS_SECRET_ACCESS_KEY": os.environ.get("AWS_SECRET_ACCESS_KEY"),
        }
    os.environ["AWS_ACCESS_KEY_ID"] = access
    os.environ["AWS_SECRET_ACCESS_KEY"] = secret
    BEDROCK_RUNTIME_ENV_ACTIVE = True


def _load_chat_history(repo_name: str) -> list[dict]:
    try:
        path = CHAT_HISTORY_DIR / f"{repo_name}.json"
        if path.exists():
            data = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(data, list):
                return data
    except Exception:
        return []
    return []


def _save_chat_history(repo_name: str, history: list[dict]) -> None:
    try:
        CHAT_HISTORY_DIR.mkdir(parents=True, exist_ok=True)
        trimmed = history[-MAX_CHAT_HISTORY:]
        path = CHAT_HISTORY_DIR / f"{repo_name}.json"
        path.write_text(json.dumps(trimmed, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _ensure_llm_state():
    config = _load_llm_config()
    if "llm_provider" not in st.session_state:
        st.session_state["llm_provider"] = config.get("provider") or "gemini"
    provider = st.session_state["llm_provider"]
    options = MODEL_OPTIONS.get(provider, MODEL_OPTIONS["gemini"])
    default_model = options[0]
    if provider == "bedrock" and BEDROCK_DEFAULT_MODEL in options:
        default_model = BEDROCK_DEFAULT_MODEL
    if "llm_model" not in st.session_state:
        configured_model = str(config.get("model") or "").strip()
        st.session_state["llm_model"] = configured_model if configured_model in options else default_model
    if st.session_state["llm_model"] not in options:
        st.session_state["llm_model"] = default_model
    if "llm_use_system_key" not in st.session_state:
        st.session_state["llm_use_system_key"] = config.get("use_system_key", True)
    if "llm_api_key" not in st.session_state:
        st.session_state["llm_api_key"] = (
            "" if st.session_state["llm_use_system_key"] else (config.get("api_key") or "")
        )
    if "llm_bedrock_access_key" not in st.session_state:
        st.session_state["llm_bedrock_access_key"] = (
            "" if st.session_state["llm_use_system_key"] else (config.get("bedrock_access_key") or "")
        )
    if "llm_bedrock_secret_key" not in st.session_state:
        st.session_state["llm_bedrock_secret_key"] = (
            "" if st.session_state["llm_use_system_key"] else (config.get("bedrock_secret_key") or "")
        )
    if provider == "bedrock":
        _set_runtime_key(provider, None)
    elif st.session_state["llm_use_system_key"]:
        _set_runtime_key(provider, None)
    else:
        _set_runtime_key(provider, st.session_state.get("llm_api_key") or None)
    _set_runtime_bedrock_credentials(
        provider,
        bool(st.session_state.get("llm_use_system_key", True)),
        st.session_state.get("llm_bedrock_access_key"),
        st.session_state.get("llm_bedrock_secret_key"),
    )


def _new_doc_section_row(
    *,
    key: str = "",
    title: str = "",
    description: str = "",
    row_id: str | None = None,
) -> dict:
    return {
        "id": row_id or uuid.uuid4().hex,
        "key": str(key or "").strip(),
        "title": str(title or "").strip(),
        "description": str(description or "").strip(),
    }


def _localized_standard_section_titles(language: str | None) -> dict[str, str]:
    lang = _normalize_doc_language(language)
    base = {
        key: str(meta.get("title") or key)
        for key, meta in DEFAULT_DOCUMENTATION_SECTIONS.items()
    }
    localized_by_lang = {
        "PT-BR": {
            "INDEX": "Visão Geral",
            "GETTING_STARTED": "Primeiros Passos",
            "INSTALLATION": "Instalação",
            "USAGE": "Uso",
            "ARCHITECTURE": "Arquitetura",
            "TECHNOLOGIES": "Tecnologias",
            "API_REFERENCE": "Referência de API",
            "CONFIGURATION": "Configuração",
            "TESTING": "Testes",
            "FILE_ANALYSIS": "Análise Arquivo a Arquivo",
        },
        "EN-US": {},
        "ES-ES": {
            "INDEX": "Descripción General",
            "GETTING_STARTED": "Primeros Pasos",
            "INSTALLATION": "Instalación",
            "USAGE": "Uso",
            "ARCHITECTURE": "Arquitectura",
            "TECHNOLOGIES": "Tecnologías",
            "API_REFERENCE": "Referencia de API",
            "CONFIGURATION": "Configuración",
            "TESTING": "Pruebas",
            "FILE_ANALYSIS": "Análisis Archivo por Archivo",
        },
        "FR-FR": {
            "INDEX": "Vue d'ensemble",
            "GETTING_STARTED": "Prise en Main",
            "INSTALLATION": "Installation",
            "USAGE": "Utilisation",
            "ARCHITECTURE": "Architecture",
            "TECHNOLOGIES": "Technologies",
            "API_REFERENCE": "Référence API",
            "CONFIGURATION": "Configuration",
            "TESTING": "Tests",
            "FILE_ANALYSIS": "Analyse Fichier par Fichier",
        },
        "DE-DE": {
            "INDEX": "Überblick",
            "GETTING_STARTED": "Erste Schritte",
            "INSTALLATION": "Installation",
            "USAGE": "Nutzung",
            "ARCHITECTURE": "Architektur",
            "TECHNOLOGIES": "Technologien",
            "API_REFERENCE": "API-Referenz",
            "CONFIGURATION": "Konfiguration",
            "TESTING": "Tests",
            "FILE_ANALYSIS": "Datei-für-Datei-Analyse",
        },
    }
    out = dict(base)
    out.update(localized_by_lang.get(lang, {}))
    return out


def _localized_default_doc_sections(language: str | None) -> dict[str, dict]:
    title_map = _localized_standard_section_titles(language)
    sections: dict[str, dict] = {}
    for key, meta in DEFAULT_DOCUMENTATION_SECTIONS.items():
        sections[key] = {
            "title": title_map.get(key, str(meta.get("title") or key)),
            "description": str(meta.get("description") or ""),
        }
    return sections


def _localized_standard_functional_section_titles(language: str | None) -> dict[str, str]:
    lang = _normalize_doc_language(language)
    base = {
        key: str(meta.get("title") or key)
        for key, meta in DEFAULT_FUNCTIONAL_DOCUMENTATION_SECTIONS.items()
    }
    localized_by_lang = {
        "PT-BR": {
            "OVERVIEW": "Visão Geral",
            "BUSINESS_SCOPE": "Escopo de Negócio",
            "FUNCTIONAL_FLOWS": "Fluxos Funcionais",
            "BUSINESS_RULES_EXCEPTIONS": "Regras de Negócio e Exceções",
            "OPERATIONAL_PROCEDURES": "Procedimentos Operacionais",
        },
        "EN-US": {},
        "ES-ES": {
            "OVERVIEW": "Descripción General",
            "BUSINESS_SCOPE": "Alcance de Negocio",
            "FUNCTIONAL_FLOWS": "Flujos Funcionales",
            "BUSINESS_RULES_EXCEPTIONS": "Reglas de Negocio y Excepciones",
            "OPERATIONAL_PROCEDURES": "Procedimientos Operativos",
        },
        "FR-FR": {
            "OVERVIEW": "Vue d'ensemble",
            "BUSINESS_SCOPE": "Périmètre Métier",
            "FUNCTIONAL_FLOWS": "Flux Fonctionnels",
            "BUSINESS_RULES_EXCEPTIONS": "Règles Métier et Exceptions",
            "OPERATIONAL_PROCEDURES": "Procédures Opérationnelles",
        },
        "DE-DE": {
            "OVERVIEW": "Überblick",
            "BUSINESS_SCOPE": "Geschäftsumfang",
            "FUNCTIONAL_FLOWS": "Funktionale Abläufe",
            "BUSINESS_RULES_EXCEPTIONS": "Geschäftsregeln und Ausnahmen",
            "OPERATIONAL_PROCEDURES": "Betriebliche Verfahren",
        },
    }
    out = dict(base)
    out.update(localized_by_lang.get(lang, {}))
    return out


def _localized_default_functional_sections(language: str | None) -> dict[str, dict]:
    title_map = _localized_standard_functional_section_titles(language)
    sections: dict[str, dict] = {}
    for key, meta in DEFAULT_FUNCTIONAL_DOCUMENTATION_SECTIONS.items():
        sections[key] = {
            "title": title_map.get(key, str(meta.get("title") or key)),
            "description": str(meta.get("description") or ""),
        }
    return sections


def _default_doc_section_rows(language: str | None = None) -> list[dict]:
    _ = language
    rows: list[dict] = []
    for key, meta in DEFAULT_DOCUMENTATION_SECTIONS.items():
        rows.append(
            _new_doc_section_row(
                key=key,
                title=str(meta.get("title") or key),
                description=str(meta.get("description") or ""),
            )
        )
    return rows


def _default_functional_section_rows(language: str | None = None) -> list[dict]:
    _ = language
    rows: list[dict] = []
    for key, meta in DEFAULT_FUNCTIONAL_DOCUMENTATION_SECTIONS.items():
        rows.append(
            _new_doc_section_row(
                key=key,
                title=str(meta.get("title") or key),
                description=str(meta.get("description") or ""),
            )
        )
    return rows


def _default_doc_sections_for_generation() -> dict[str, dict]:
    return {
        key: {
            "title": str(meta.get("title") or key),
            "description": str(meta.get("description") or ""),
        }
        for key, meta in DEFAULT_DOCUMENTATION_SECTIONS.items()
    }


def _default_functional_sections_for_generation() -> dict[str, dict]:
    return {
        key: {
            "title": str(meta.get("title") or key),
            "description": str(meta.get("description") or ""),
        }
        for key, meta in DEFAULT_FUNCTIONAL_DOCUMENTATION_SECTIONS.items()
    }


def _sanitize_doc_section_rows(rows: list[dict]) -> list[dict]:
    defaults = _default_doc_sections_for_generation()
    sanitized: list[dict] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        raw_key = str(row.get("key") or "").strip()
        key_norm = _normalize_doc_section_key(raw_key)
        key_value = raw_key or key_norm
        title = str(row.get("title") or "").strip()
        description = str(row.get("description") or "").strip()
        if key_norm in defaults:
            if not title:
                title = defaults[key_norm]["title"]
            if not description:
                description = defaults[key_norm]["description"]
        sanitized.append(
            _new_doc_section_row(
                key=key_value,
                title=title,
                description=description,
                row_id=str(row.get("id") or uuid.uuid4().hex),
            )
        )
    return sanitized


def _sanitize_functional_section_rows(rows: list[dict]) -> list[dict]:
    defaults = _default_functional_sections_for_generation()
    sanitized: list[dict] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        raw_key = str(row.get("key") or "").strip()
        key_norm = _normalize_doc_section_key(raw_key)
        key_value = raw_key or key_norm
        title = str(row.get("title") or "").strip()
        description = str(row.get("description") or "").strip()
        if key_norm in defaults:
            if not title:
                title = defaults[key_norm]["title"]
            if not description:
                description = defaults[key_norm]["description"]
        sanitized.append(
            _new_doc_section_row(
                key=key_value,
                title=title,
                description=description,
                row_id=str(row.get("id") or uuid.uuid4().hex),
            )
        )
    return sanitized


def _section_rows_signature(rows: list[dict]) -> str:
    payload: list[dict[str, str]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        payload.append(
            {
                "key": str(row.get("key") or "").strip(),
                "title": str(row.get("title") or "").strip(),
                "description": str(row.get("description") or "").strip(),
            }
        )
    return json.dumps(payload, ensure_ascii=False, sort_keys=True)


def _standard_section_title_aliases(*, functional: bool) -> dict[str, set[str]]:
    title_getter = (
        _localized_standard_functional_section_titles
        if functional
        else _localized_standard_section_titles
    )
    aliases: dict[str, set[str]] = {}
    for lang in ("EN-US", "PT-BR", "ES-ES", "FR-FR", "DE-DE"):
        for key, title in title_getter(lang).items():
            aliases.setdefault(str(key), set()).add(str(title).strip().casefold())
    return aliases


def _build_confirmed_section_rows(rows: list[dict], *, functional: bool) -> list[dict]:
    sanitizer = _sanitize_functional_section_rows if functional else _sanitize_doc_section_rows
    aliases_by_key = _standard_section_title_aliases(functional=functional)
    sanitized_rows = sanitizer(list(rows or []))
    confirmed_rows: list[dict] = []

    for row in sanitized_rows:
        row_id = str(row.get("id") or uuid.uuid4().hex)
        raw_key = str(row.get("key") or "").strip()
        key_norm = _normalize_doc_section_key(raw_key)
        title = str(row.get("title") or "").strip()
        description = str(row.get("description") or "").strip()

        key_for_runtime = raw_key
        # If a standard section was renamed, force runtime key to derive from the custom title
        # so downstream logic won't map it back to the standard title.
        if title:
            if not key_for_runtime:
                key_for_runtime = title
            elif key_norm in aliases_by_key:
                if title.casefold() not in aliases_by_key.get(key_norm, set()):
                    key_for_runtime = title

        confirmed_rows.append(
            _new_doc_section_row(
                key=key_for_runtime,
                title=title,
                description=description,
                row_id=row_id,
            )
        )
    return confirmed_rows


def _get_effective_section_rows(
    *,
    rows_state_key: str,
    confirmed_rows_key: str,
    confirmed_signature_key: str,
    functional: bool,
) -> tuple[list[dict], bool]:
    sanitizer = _sanitize_functional_section_rows if functional else _sanitize_doc_section_rows
    current_rows = sanitizer(list(st.session_state.get(rows_state_key) or []))
    current_signature = _section_rows_signature(current_rows)
    confirmed_signature = str(st.session_state.get(confirmed_signature_key) or "").strip()
    confirmed_rows_raw = st.session_state.get(confirmed_rows_key)
    if isinstance(confirmed_rows_raw, list) and confirmed_signature == current_signature and confirmed_signature:
        confirmed_rows = sanitizer(list(confirmed_rows_raw))
        if confirmed_rows:
            return confirmed_rows, True
    return current_rows, False


def _sections_pending_confirmation(*, functional: bool) -> bool:
    rows_state_key = "functional_sections_rows" if functional else "doc_sections_rows"
    confirmed_signature_key = (
        "functional_sections_confirmed_signature"
        if functional
        else "doc_sections_confirmed_signature"
    )
    sanitizer = _sanitize_functional_section_rows if functional else _sanitize_doc_section_rows
    default_rows_factory = (
        _default_functional_section_rows if functional else _default_doc_section_rows
    )

    current_rows = sanitizer(list(st.session_state.get(rows_state_key) or []))
    if not current_rows:
        return False

    current_signature = _section_rows_signature(current_rows)
    default_signature = _section_rows_signature(sanitizer(default_rows_factory()))
    if current_signature == default_signature:
        return False

    confirmed_signature = str(st.session_state.get(confirmed_signature_key) or "").strip()
    return confirmed_signature != current_signature


def _clear_doc_section_widget_state(row_ids: set[str]) -> None:
    for row_id in row_ids:
        st.session_state.pop(f"doc_sec_name_{row_id}", None)
        st.session_state.pop(f"doc_sec_desc_{row_id}", None)


def _ensure_doc_sections_state(language: str | None = None) -> None:
    if "doc_sections_rows" not in st.session_state:
        st.session_state["doc_sections_rows"] = _default_doc_section_rows(language)
        return
    rows_raw = st.session_state.get("doc_sections_rows")
    if not isinstance(rows_raw, list) or not rows_raw:
        st.session_state["doc_sections_rows"] = _default_doc_section_rows(language)
        return
    rows = _sanitize_doc_section_rows(list(rows_raw))
    if not rows:
        st.session_state["doc_sections_rows"] = _default_doc_section_rows(language)
        return
    coerced_rows = _coerce_doc_section_rows_to_english(rows)
    if coerced_rows != rows_raw:
        row_ids = {str(row.get("id")) for row in rows_raw if isinstance(row, dict)}
        _clear_doc_section_widget_state(row_ids)
    st.session_state["doc_sections_rows"] = coerced_rows


def _clear_functional_section_widget_state(row_ids: set[str]) -> None:
    for row_id in row_ids:
        st.session_state.pop(f"func_sec_name_{row_id}", None)
        st.session_state.pop(f"func_sec_desc_{row_id}", None)


def _normalize_doc_section_key(value: str) -> str:
    import re
    import unicodedata

    raw = str(value or "").strip()
    raw = unicodedata.normalize("NFKD", raw)
    raw = "".join(ch for ch in raw if not unicodedata.combining(ch))
    key = raw.upper().replace("-", "_").replace(" ", "_")
    key = re.sub(r"[^A-Z0-9_]", "", key)
    return key


def _coerce_doc_section_rows_to_english(rows: list[dict]) -> list[dict]:
    english_titles = _localized_standard_section_titles("EN-US")
    localized_titles_by_key: dict[str, set[str]] = {
        key: {str(title).strip().casefold()}
        for key, title in english_titles.items()
    }
    for lang in ("PT-BR", "ES-ES", "FR-FR", "DE-DE"):
        localized_map = _localized_standard_section_titles(lang)
        for key, title in localized_map.items():
            localized_titles_by_key.setdefault(key, set()).add(str(title).strip().casefold())

    updated_rows: list[dict] = []
    for row in rows:
        normalized_row = _new_doc_section_row(
            key=str(row.get("key") or ""),
            title=str(row.get("title") or ""),
            description=str(row.get("description") or ""),
            row_id=str(row.get("id") or uuid.uuid4().hex),
        )
        key = _normalize_doc_section_key(str(normalized_row.get("key") or ""))
        title = str(normalized_row.get("title") or "").strip()
        if key in english_titles and title.casefold() in localized_titles_by_key.get(key, set()):
            normalized_row["title"] = english_titles[key]
        updated_rows.append(normalized_row)
    return updated_rows


def _coerce_functional_section_rows_to_english(rows: list[dict]) -> list[dict]:
    english_titles = _localized_standard_functional_section_titles("EN-US")
    localized_titles_by_key: dict[str, set[str]] = {
        key: {str(title).strip().casefold()}
        for key, title in english_titles.items()
    }
    for lang in ("PT-BR", "ES-ES", "FR-FR", "DE-DE"):
        localized_map = _localized_standard_functional_section_titles(lang)
        for key, title in localized_map.items():
            localized_titles_by_key.setdefault(key, set()).add(str(title).strip().casefold())

    updated_rows: list[dict] = []
    for row in rows:
        normalized_row = _new_doc_section_row(
            key=str(row.get("key") or ""),
            title=str(row.get("title") or ""),
            description=str(row.get("description") or ""),
            row_id=str(row.get("id") or uuid.uuid4().hex),
        )
        key = _normalize_doc_section_key(str(normalized_row.get("key") or ""))
        title = str(normalized_row.get("title") or "").strip()
        if key in english_titles and title.casefold() in localized_titles_by_key.get(key, set()):
            normalized_row["title"] = english_titles[key]
        updated_rows.append(normalized_row)
    return updated_rows


def _ensure_functional_sections_state(language: str | None = None) -> None:
    if "functional_sections_rows" not in st.session_state:
        st.session_state["functional_sections_rows"] = _default_functional_section_rows(language)
        return
    rows_raw = st.session_state.get("functional_sections_rows")
    if not isinstance(rows_raw, list) or not rows_raw:
        st.session_state["functional_sections_rows"] = _default_functional_section_rows(language)
        return
    rows = _sanitize_functional_section_rows(list(rows_raw))
    if not rows:
        st.session_state["functional_sections_rows"] = _default_functional_section_rows(language)
        return
    coerced_rows = _coerce_functional_section_rows_to_english(rows)
    if coerced_rows != rows_raw:
        row_ids = {str(row.get("id")) for row in rows_raw if isinstance(row, dict)}
        _clear_functional_section_widget_state(row_ids)
    st.session_state["functional_sections_rows"] = coerced_rows


def _resolve_runtime_doc_sections() -> tuple[dict, list[str]]:
    rows, _ = _get_effective_section_rows(
        rows_state_key="doc_sections_rows",
        confirmed_rows_key="doc_sections_confirmed_rows",
        confirmed_signature_key="doc_sections_confirmed_signature",
        functional=False,
    )
    errors: list[str] = []
    sections: dict[str, dict] = {}
    used_keys: set[str] = set()

    for idx, row in enumerate(rows, start=1):
        title = str(row.get("title") or "").strip()
        description = str(row.get("description") or "").strip()
        if not title and not description:
            continue
        if not title:
            errors.append(f"Section {idx}: name is required.")
        if not description:
            errors.append(f"Section {idx}: description is required.")
        if not title or not description:
            continue

        raw_key = str(row.get("key") or "").strip() or title
        key = _normalize_doc_section_key(raw_key) or f"SECTION_{idx}"
        base_key = key
        suffix = 2
        while key in used_keys:
            key = f"{base_key}_{suffix}"
            suffix += 1
        used_keys.add(key)
        sections[key] = {"title": title, "description": description}

    if sections:
        return sections, errors
    if errors:
        errors.append("Add at least one section to generate documentation.")
        return sections, errors
    return _default_doc_sections_for_generation(), []


def _resolve_runtime_functional_sections() -> tuple[dict, list[str]]:
    rows, _ = _get_effective_section_rows(
        rows_state_key="functional_sections_rows",
        confirmed_rows_key="functional_sections_confirmed_rows",
        confirmed_signature_key="functional_sections_confirmed_signature",
        functional=True,
    )
    errors: list[str] = []
    sections: dict[str, dict] = {}
    used_keys: set[str] = set()

    for idx, row in enumerate(rows, start=1):
        title = str(row.get("title") or "").strip()
        description = str(row.get("description") or "").strip()
        if not title and not description:
            continue
        if not title:
            errors.append(f"Functional section {idx}: name is required.")
        if not description:
            errors.append(f"Functional section {idx}: description is required.")
        if not title or not description:
            continue

        raw_key = str(row.get("key") or "").strip() or title
        key = _normalize_doc_section_key(raw_key) or f"FUNCTIONAL_SECTION_{idx}"
        base_key = key
        suffix = 2
        while key in used_keys:
            key = f"{base_key}_{suffix}"
            suffix += 1
        used_keys.add(key)
        sections[key] = {"title": title, "description": description}

    if sections:
        return sections, errors
    if errors:
        errors.append("Add at least one functional section to generate documentation.")
        return sections, errors
    return _default_functional_sections_for_generation(), []


def _render_doc_sections_editor(language: str | None = None) -> None:
    _ensure_doc_sections_state(language)
    with st.expander("[OPTIONAL] Customize documentation sections", expanded=False):
        st.caption(
            "Applies only to this documentation generation. Name and description are required for each section."
        )

        rows = list(st.session_state.get("doc_sections_rows") or [])
        updated_rows: list[dict] = []
        remove_ids: set[str] = set()

        for idx, row in enumerate(rows, start=1):
            row_id = str(row.get("id") or uuid.uuid4().hex)
            name_key = f"doc_sec_name_{row_id}"
            desc_key = f"doc_sec_desc_{row_id}"

            if name_key not in st.session_state:
                st.session_state[name_key] = str(row.get("title") or "")
            if desc_key not in st.session_state:
                st.session_state[desc_key] = str(row.get("description") or "")

            cols = st.columns([3, 5, 1])
            name = cols[0].text_input(f"Name {idx}", key=name_key)
            description = cols[1].text_area(f"Description {idx}", key=desc_key, height=78)
            remove_clicked = cols[2].button(
                "❌",
                key=f"doc_sec_remove_{row_id}",
                disabled=(len(rows) <= 1),
            )
            if remove_clicked:
                remove_ids.add(row_id)

            updated_rows.append(
                _new_doc_section_row(
                    key=str(row.get("key") or ""),
                    title=name,
                    description=description,
                    row_id=row_id,
                )
            )

        if remove_ids:
            updated_rows = [row for row in updated_rows if str(row.get("id")) not in remove_ids]
            _clear_doc_section_widget_state(remove_ids)
            st.session_state["doc_sections_rows"] = updated_rows
            _safe_rerun()
            return

        current_rows_sanitized = _sanitize_doc_section_rows(updated_rows)
        current_signature = _section_rows_signature(current_rows_sanitized)
        confirmed_signature = str(st.session_state.get("doc_sections_confirmed_signature") or "").strip()
        confirmed_rows_raw = st.session_state.get("doc_sections_confirmed_rows")
        is_confirmed_current = (
            isinstance(confirmed_rows_raw, list)
            and bool(confirmed_signature)
            and confirmed_signature == current_signature
        )

        col_add, col_reset, col_confirm = st.columns(3)
        if col_add.button("Add section", key="doc_sec_add"):
            updated_rows.append(_new_doc_section_row())
            st.session_state["doc_sections_rows"] = updated_rows
            _safe_rerun()
            return
        if col_reset.button("Reset to default", key="doc_sec_reset"):
            old_ids = {str(row.get("id")) for row in rows}
            _clear_doc_section_widget_state(old_ids)
            updated_rows = _default_doc_section_rows()
            st.session_state["doc_sections_rows"] = updated_rows
            st.session_state.pop("doc_sections_confirmed_rows", None)
            st.session_state.pop("doc_sections_confirmed_signature", None)
            _safe_rerun()
            return
        if col_confirm.button("Confirm custom sections", key="doc_sec_confirm"):
            st.session_state["doc_sections_rows"] = updated_rows
            st.session_state["doc_sections_confirmed_rows"] = _build_confirmed_section_rows(
                updated_rows,
                functional=False,
            )
            st.session_state["doc_sections_confirmed_signature"] = current_signature
            _safe_rerun()
            return

        st.session_state["doc_sections_rows"] = updated_rows
        st.caption(f"Sections configured: {len(updated_rows)}")
        if _sections_pending_confirmation(functional=False):
            st.caption("Changes detected. Click 'Confirm custom sections' before generating.")
        elif is_confirmed_current:
            st.caption("Custom section definitions confirmed for this generation.")


def _render_functional_sections_editor(language: str | None = None) -> None:
    _ensure_functional_sections_state(language)
    with st.expander("[OPTIONAL] Customize functional documentation sections", expanded=False):
        st.caption(
            "Applies only to this functional documentation generation. "
            "Name and description are required for each section."
        )

        rows = list(st.session_state.get("functional_sections_rows") or [])
        updated_rows: list[dict] = []
        remove_ids: set[str] = set()

        for idx, row in enumerate(rows, start=1):
            row_id = str(row.get("id") or uuid.uuid4().hex)
            name_key = f"func_sec_name_{row_id}"
            desc_key = f"func_sec_desc_{row_id}"

            if name_key not in st.session_state:
                st.session_state[name_key] = str(row.get("title") or "")
            if desc_key not in st.session_state:
                st.session_state[desc_key] = str(row.get("description") or "")

            cols = st.columns([3, 5, 1])
            name = cols[0].text_input(f"Functional Name {idx}", key=name_key)
            description = cols[1].text_area(
                f"Functional Description {idx}",
                key=desc_key,
                height=78,
            )
            remove_clicked = cols[2].button(
                "❌",
                key=f"func_sec_remove_{row_id}",
                disabled=(len(rows) <= 1),
            )
            if remove_clicked:
                remove_ids.add(row_id)

            updated_rows.append(
                _new_doc_section_row(
                    key=str(row.get("key") or ""),
                    title=name,
                    description=description,
                    row_id=row_id,
                )
            )

        if remove_ids:
            updated_rows = [row for row in updated_rows if str(row.get("id")) not in remove_ids]
            _clear_functional_section_widget_state(remove_ids)
            st.session_state["functional_sections_rows"] = updated_rows
            _safe_rerun()
            return

        current_rows_sanitized = _sanitize_functional_section_rows(updated_rows)
        current_signature = _section_rows_signature(current_rows_sanitized)
        confirmed_signature = str(
            st.session_state.get("functional_sections_confirmed_signature") or ""
        ).strip()
        confirmed_rows_raw = st.session_state.get("functional_sections_confirmed_rows")
        is_confirmed_current = (
            isinstance(confirmed_rows_raw, list)
            and bool(confirmed_signature)
            and confirmed_signature == current_signature
        )

        col_add, col_reset, col_confirm = st.columns(3)
        if col_add.button("Add section", key="func_sec_add"):
            updated_rows.append(_new_doc_section_row())
            st.session_state["functional_sections_rows"] = updated_rows
            _safe_rerun()
            return
        if col_reset.button("Reset to default", key="func_sec_reset"):
            old_ids = {str(row.get("id")) for row in rows}
            _clear_functional_section_widget_state(old_ids)
            updated_rows = _default_functional_section_rows()
            st.session_state["functional_sections_rows"] = updated_rows
            st.session_state.pop("functional_sections_confirmed_rows", None)
            st.session_state.pop("functional_sections_confirmed_signature", None)
            _safe_rerun()
            return
        if col_confirm.button("Confirm custom sections", key="func_sec_confirm"):
            st.session_state["functional_sections_rows"] = updated_rows
            st.session_state["functional_sections_confirmed_rows"] = _build_confirmed_section_rows(
                updated_rows,
                functional=True,
            )
            st.session_state["functional_sections_confirmed_signature"] = current_signature
            _safe_rerun()
            return

        st.session_state["functional_sections_rows"] = updated_rows
        st.caption(f"Functional sections configured: {len(updated_rows)}")
        if _sections_pending_confirmation(functional=True):
            st.caption("Changes detected. Click 'Confirm custom sections' before generating.")
        elif is_confirmed_current:
            st.caption("Functional section definitions confirmed for this generation.")

def _get_query_param(name: str):
    try:
        return st.query_params.get(name)
    except Exception:
        params = st.experimental_get_query_params()
        value = params.get(name)
        if isinstance(value, list):
            return value[0] if value else None
        return value


def _save_last_repo(repo_name: str) -> None:
    try:
        LAST_REPO_FILE.write_text(repo_name, encoding="utf-8")
    except Exception:
        pass


def _load_last_repo() -> str | None:
    try:
        if LAST_REPO_FILE.exists():
            value = LAST_REPO_FILE.read_text(encoding="utf-8").strip()
            return value or None
    except Exception:
        return None
    return None


def _render_top_logo() -> None:
    if not TOP_LOGO_PATH.exists():
        return
    try:
        svg_b64 = base64.b64encode(TOP_LOGO_PATH.read_bytes()).decode("ascii")
        st.markdown(
            f"""
            <div style="display:flex; justify-content:center; margin: 0.0rem 0 0.45rem 0;">
              <img src="data:image/svg+xml;base64,{svg_b64}" alt="genMS logo" style="max-height:68px; width:auto;" />
            </div>
            """,
            unsafe_allow_html=True,
        )
    except Exception:
        pass


def _repo_slug(repo_name: str) -> str:
    import re

    value = (repo_name or "").strip()
    if not value:
        return "repo"
    return re.sub(r"[^A-Za-z0-9._-]+", "_", value)


def _load_repo_library() -> dict:
    try:
        if REPO_LIBRARY_FILE.exists():
            data = json.loads(REPO_LIBRARY_FILE.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
        if LEGACY_LIBRARY_FILE.exists():
            data = json.loads(LEGACY_LIBRARY_FILE.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
        if LEGACY_REPO_LIBRARY_FILE.exists():
            data = json.loads(LEGACY_REPO_LIBRARY_FILE.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}
    return {}


def _save_repo_library(data: dict) -> None:
    try:
        REPO_LIBRARY_FILE.parent.mkdir(parents=True, exist_ok=True)
        REPO_LIBRARY_FILE.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        pass


def _load_functional_repo_library() -> dict:
    try:
        if FUNCTIONAL_REPO_LIBRARY_FILE.exists():
            data = json.loads(FUNCTIONAL_REPO_LIBRARY_FILE.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}
    return {}


def _save_functional_repo_library(data: dict) -> None:
    try:
        FUNCTIONAL_REPO_LIBRARY_FILE.parent.mkdir(parents=True, exist_ok=True)
        FUNCTIONAL_REPO_LIBRARY_FILE.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        pass


def _library_repo_dir(repo_name: str) -> Path:
    return REPO_ASSETS_DIR / _repo_slug(repo_name)


def _functional_library_repo_dir(repo_name: str) -> Path:
    return FUNCTIONAL_REPO_ASSETS_DIR / _repo_slug(repo_name)


def _library_entry_key(repo_name: str, language: str | None) -> str:
    lang = _normalize_doc_language(language)
    return f"{repo_name}::{lang}"


def _parse_library_entry_key(entry_key: str) -> tuple[str, str | None]:
    if "::" in (entry_key or ""):
        repo, lang = entry_key.split("::", 1)
        return repo, _normalize_doc_language(lang)
    return entry_key, None


def _library_repo_dir_for_language(repo_name: str, language: str | None) -> Path:
    lang = _normalize_doc_language(language)
    return REPO_ASSETS_DIR / f"{_repo_slug(repo_name)}__{_repo_slug(lang)}"


def _functional_library_repo_dir_for_language(repo_name: str, language: str | None) -> Path:
    lang = _normalize_doc_language(language)
    return FUNCTIONAL_REPO_ASSETS_DIR / f"{_repo_slug(repo_name)}__{_repo_slug(lang)}"


def _reset_docs_workspace() -> None:
    docs_dst = Path.cwd() / "docs"
    md_dst = Path.cwd() / "documentation.md"
    functional_docs_dst = Path.cwd() / "docs_functional"
    functional_md_dst = Path.cwd() / "functional_documentation.md"
    try:
        if docs_dst.exists():
            shutil.rmtree(docs_dst)
        if md_dst.exists():
            md_dst.unlink()
        if functional_docs_dst.exists():
            shutil.rmtree(functional_docs_dst)
        if functional_md_dst.exists():
            functional_md_dst.unlink()
    except Exception:
        pass


def _upsert_repo_library_entry(repo_name: str, updates: dict, language: str | None = None) -> tuple[str, dict]:
    repo_name = (repo_name or "").strip()
    if not repo_name:
        return "", {}
    normalized_lang = _normalize_doc_language(language or updates.get("language"))
    entry_key = updates.get("entry_key") or _library_entry_key(repo_name, normalized_lang)
    library = _load_repo_library()
    current = library.get(entry_key, {})
    if not current and entry_key != repo_name:
        # Backward compatibility: migrate plain repo key to language key.
        current = library.get(repo_name, {})
        if current:
            library.pop(repo_name, None)
    merged = dict(current)
    merged.update(updates)
    merged["repo_name"] = repo_name
    merged["language"] = normalized_lang
    merged["entry_key"] = entry_key
    merged["updated_at"] = (
        datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    )
    library[entry_key] = merged
    _save_repo_library(library)
    return entry_key, merged


def _upsert_functional_repo_library_entry(
    repo_name: str,
    updates: dict,
    language: str | None = None,
) -> tuple[str, dict]:
    repo_name = (repo_name or "").strip()
    if not repo_name:
        return "", {}
    normalized_lang = _normalize_doc_language(language or updates.get("language"))
    entry_key = updates.get("entry_key") or _library_entry_key(repo_name, normalized_lang)
    library = _load_functional_repo_library()
    current = library.get(entry_key, {})
    merged = dict(current)
    merged.update(updates)
    merged["repo_name"] = repo_name
    merged["language"] = normalized_lang
    merged["entry_key"] = entry_key
    merged["updated_at"] = (
        datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")
    )
    library[entry_key] = merged
    _save_functional_repo_library(library)
    return entry_key, merged


def _format_library_updated_at(updated_at: str | None) -> str:
    if not updated_at:
        return "-"
    value = str(updated_at).strip()
    try:
        if value.endswith("Z"):
            dt_utc = datetime.fromisoformat(value.replace("Z", "+00:00"))
        else:
            dt_utc = datetime.fromisoformat(value)
            if dt_utc.tzinfo is None:
                dt_utc = dt_utc.replace(tzinfo=timezone.utc)
        utc_minus_3 = timezone(timedelta(hours=-3))
        dt_local = dt_utc.astimezone(utc_minus_3)
        return dt_local.strftime("%d/%m/%y %H:%M:%S")
    except Exception:
        return value


def _snapshot_repo_assets(
    repo_name: str,
    code_json_path: Path,
    language: str,
    graph_path: str | None = None,
    include_functional: bool = False,
) -> dict:
    repo_dir = _library_repo_dir_for_language(repo_name, language)
    docs_src = Path.cwd() / "docs"
    md_src = Path.cwd() / "documentation.md"
    functional_docs_src = Path.cwd() / "docs_functional"
    functional_md_src = Path.cwd() / "functional_documentation.md"
    docs_dst = repo_dir / "docs"
    code_dst = repo_dir / "code.json"
    md_dst = repo_dir / "documentation.md"
    functional_docs_dst = repo_dir / "docs_functional"
    functional_md_dst = repo_dir / "functional_documentation.md"
    graph_dst = repo_dir / "graph.json"

    payload = {
        "library_repo_dir": str(repo_dir),
        "library_docs_dir": str(docs_dst),
        "library_code_json": str(code_dst),
        "library_documentation_md": str(md_dst),
        "library_functional_docs_dir": str(functional_docs_dst),
        "library_functional_documentation_md": str(functional_md_dst),
        "library_graph_json": "",
        "docs_available": False,
        "functional_docs_available": False,
    }

    try:
        repo_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(code_json_path, code_dst)

        if docs_dst.exists():
            shutil.rmtree(docs_dst)
        if docs_src.exists():
            shutil.copytree(docs_src, docs_dst)
            payload["docs_available"] = any(docs_dst.glob("*.md"))

        if md_src.exists():
            shutil.copy2(md_src, md_dst)

        if include_functional:
            if functional_docs_dst.exists():
                shutil.rmtree(functional_docs_dst)
            if functional_docs_src.exists():
                shutil.copytree(functional_docs_src, functional_docs_dst)
                payload["functional_docs_available"] = any(functional_docs_dst.glob("*.md"))
            if functional_md_src.exists():
                shutil.copy2(functional_md_src, functional_md_dst)
                payload["functional_docs_available"] = bool(
                    payload["functional_docs_available"] or functional_md_dst.exists()
                )

        graph_src = Path(graph_path) if graph_path else None
        if graph_src and graph_src.exists():
            shutil.copy2(graph_src, graph_dst)
            payload["library_graph_json"] = str(graph_dst)
        elif graph_dst.exists():
            graph_dst.unlink()

        return payload
    except Exception:
        return payload


def _snapshot_functional_assets(
    repo_name: str,
    language: str,
    graph_path: str | None = None,
    code_json_path: Path | None = None,
) -> dict:
    repo_dir = _functional_library_repo_dir_for_language(repo_name, language)
    docs_src = Path.cwd() / "docs_functional"
    md_src = Path.cwd() / "functional_documentation.md"
    docs_dst = repo_dir / "docs_functional"
    md_dst = repo_dir / "functional_documentation.md"
    code_dst = repo_dir / "code.json"
    graph_dst = repo_dir / "graph.json"

    payload = {
        "library_repo_dir": str(repo_dir),
        "library_docs_dir": str(docs_dst),
        "library_documentation_md": str(md_dst),
        "library_functional_docs_dir": str(docs_dst),
        "library_functional_documentation_md": str(md_dst),
        "library_code_json": str(code_dst),
        "library_graph_json": "",
        "docs_available": False,
        "functional_docs_available": False,
    }

    try:
        repo_dir.mkdir(parents=True, exist_ok=True)

        if docs_dst.exists():
            shutil.rmtree(docs_dst)
        if docs_src.exists():
            shutil.copytree(docs_src, docs_dst)
            payload["functional_docs_available"] = any(docs_dst.glob("*.md"))

        if md_dst.exists():
            md_dst.unlink()
        if md_src.exists():
            shutil.copy2(md_src, md_dst)
            payload["functional_docs_available"] = bool(
                payload["functional_docs_available"] or md_dst.exists()
            )

        payload["docs_available"] = bool(payload["functional_docs_available"])

        if code_json_path is not None and Path(code_json_path).exists():
            shutil.copy2(code_json_path, code_dst)

        graph_src = Path(graph_path) if graph_path else None
        if graph_src and graph_src.exists():
            shutil.copy2(graph_src, graph_dst)
            payload["library_graph_json"] = str(graph_dst)
        elif graph_dst.exists():
            graph_dst.unlink()

        return payload
    except Exception:
        return payload


def _resolve_library_entry_assets(repo_name: str, entry: dict) -> dict:
    resolved = dict(entry or {})
    lang = resolved.get("language")
    repo_dir = _library_repo_dir_for_language(repo_name, lang)
    legacy_repo_dir = _library_repo_dir(repo_name)
    defaults = {
        "library_repo_dir": str(repo_dir),
        "library_docs_dir": str(repo_dir / "docs"),
        "library_code_json": str(repo_dir / "code.json"),
        "library_documentation_md": str(repo_dir / "documentation.md"),
        "library_functional_docs_dir": str(repo_dir / "docs_functional"),
        "library_functional_documentation_md": str(repo_dir / "functional_documentation.md"),
        "library_graph_json": str(repo_dir / "graph.json"),
    }
    for key, fallback in defaults.items():
        current = str(resolved.get(key) or "").strip()
        if current and Path(current).exists():
            resolved[key] = current
            continue
        if Path(fallback).exists():
            resolved[key] = fallback
            continue
        legacy_fallback = fallback.replace(str(repo_dir), str(legacy_repo_dir))
        if Path(legacy_fallback).exists():
            resolved[key] = legacy_fallback
            continue
        # Keep deterministic non-empty path to avoid treating empty path as "." later.
        resolved[key] = fallback

    docs_dir_raw = str(resolved.get("library_docs_dir") or "").strip()
    docs_dir = Path(docs_dir_raw) if docs_dir_raw else None
    resolved["docs_available"] = bool(
        docs_dir is not None and docs_dir.exists() and any(docs_dir.glob("*.md"))
    )

    functional_docs_raw = str(resolved.get("library_functional_docs_dir") or "").strip()
    functional_docs_dir = Path(functional_docs_raw) if functional_docs_raw else None
    resolved["functional_docs_available"] = bool(
        functional_docs_dir is not None
        and functional_docs_dir.exists()
        and any(functional_docs_dir.glob("*.md"))
    )
    return resolved


def _resolve_functional_library_entry_assets(repo_name: str, entry: dict) -> dict:
    resolved = dict(entry or {})
    lang = resolved.get("language")
    repo_dir = _functional_library_repo_dir_for_language(repo_name, lang)
    defaults = {
        "library_repo_dir": str(repo_dir),
        "library_docs_dir": str(repo_dir / "docs_functional"),
        "library_code_json": str(repo_dir / "code.json"),
        "library_documentation_md": str(repo_dir / "functional_documentation.md"),
        "library_functional_docs_dir": str(repo_dir / "docs_functional"),
        "library_functional_documentation_md": str(repo_dir / "functional_documentation.md"),
        "library_graph_json": str(repo_dir / "graph.json"),
    }
    for key, fallback in defaults.items():
        current = str(resolved.get(key) or "").strip()
        if current and Path(current).exists():
            resolved[key] = current
            continue
        resolved[key] = fallback

    functional_docs_raw = str(resolved.get("library_functional_docs_dir") or "").strip()
    functional_docs_dir = Path(functional_docs_raw) if functional_docs_raw else None
    resolved["functional_docs_available"] = bool(
        functional_docs_dir is not None
        and functional_docs_dir.exists()
        and any(functional_docs_dir.glob("*.md"))
    )
    resolved["docs_available"] = bool(resolved["functional_docs_available"])
    return resolved


def _activate_repo_assets(entry: dict, doc_variant: str = "technical") -> bool:
    def _coerce_path(value: str | None) -> Path | None:
        raw = str(value or "").strip()
        if not raw or raw in {".", "./", ".\\"}:
            return None
        try:
            return Path(raw)
        except Exception:
            return None

    def _same_path(left: Path | None, right: Path | None) -> bool:
        if left is None or right is None:
            return False
        try:
            return left.resolve() == right.resolve()
        except Exception:
            return str(left) == str(right)

    variant = str(doc_variant or "technical").strip().lower()
    if variant == "functional":
        docs_src = _coerce_path(entry.get("library_functional_docs_dir"))
        md_src = _coerce_path(entry.get("library_functional_documentation_md"))
    else:
        docs_src = _coerce_path(entry.get("library_docs_dir"))
        md_src = _coerce_path(entry.get("library_documentation_md"))
    docs_dst = Path.cwd() / "docs"
    md_dst = Path.cwd() / "documentation.md"
    functional_docs_src = _coerce_path(entry.get("library_functional_docs_dir"))
    functional_md_src = _coerce_path(entry.get("library_functional_documentation_md"))
    functional_docs_dst = Path.cwd() / "docs_functional"
    functional_md_dst = Path.cwd() / "functional_documentation.md"

    # Mandatory activation for the selected variant.
    restored_docs = False
    try:
        if docs_dst.exists():
            shutil.rmtree(docs_dst)
        if docs_src is not None and docs_src.exists() and not _same_path(docs_src, docs_dst):
            shutil.copytree(docs_src, docs_dst)
            restored_docs = True

        if md_dst.exists():
            md_dst.unlink()
        if md_src is not None and md_src.exists() and not _same_path(md_src, md_dst):
            shutil.copy2(md_src, md_dst)
    except Exception:
        return False

    # Optional sidecar sync for functional artifacts.
    # Never delete local functional outputs unless we actually have a valid source to copy from.
    try:
        if (
            functional_docs_src is not None
            and functional_docs_src.exists()
            and not _same_path(functional_docs_src, functional_docs_dst)
        ):
            if functional_docs_dst.exists():
                shutil.rmtree(functional_docs_dst)
            shutil.copytree(functional_docs_src, functional_docs_dst)
        if (
            functional_md_src is not None
            and functional_md_src.exists()
            and not _same_path(functional_md_src, functional_md_dst)
        ):
            if functional_md_dst.exists():
                functional_md_dst.unlink()
            shutil.copy2(functional_md_src, functional_md_dst)
    except Exception:
        pass
    return restored_docs


def _activate_local_functional_workspace() -> bool:
    docs_src = Path.cwd() / "docs_functional"
    md_src = Path.cwd() / "functional_documentation.md"
    docs_dst = Path.cwd() / "docs"
    md_dst = Path.cwd() / "documentation.md"
    restored = False
    try:
        if docs_dst.exists():
            shutil.rmtree(docs_dst)
        if docs_src.exists():
            shutil.copytree(docs_src, docs_dst)
            restored = True
        if md_dst.exists():
            md_dst.unlink()
        if md_src.exists():
            shutil.copy2(md_src, md_dst)
            restored = True
    except Exception:
        return False
    return restored


def _resolve_repo_code_json(repo_name: str, preferred_language: str | None = None) -> Path | None:
    library = _load_repo_library()
    preferred_lang = _normalize_doc_language(
        preferred_language or st.session_state.get("repo_language") or st.session_state.get("language")
    )
    candidates: list[tuple[int, str, Path]] = []
    for entry_key, entry in library.items():
        if not isinstance(entry, dict):
            continue
        entry_repo, _ = _parse_library_entry_key(entry_key)
        entry_repo = entry.get("repo_name") or entry_repo
        if entry_repo != repo_name:
            continue
        resolved = _resolve_library_entry_assets(repo_name, entry)
        path = resolved.get("library_code_json")
        if not path or not Path(path).exists():
            continue
        entry_lang = _normalize_doc_language(resolved.get("language"))
        score_lang = 1 if entry_lang == preferred_lang else 0
        updated = str(resolved.get("updated_at") or "")
        candidates.append((score_lang, updated, Path(path)))
    if candidates:
        candidates.sort(key=lambda row: (row[0], row[1]), reverse=True)
        return candidates[0][2]
    out_path = Path.cwd() / f"out/{repo_name}/code.json"
    if out_path.exists():
        return out_path
    return None


def _sync_repo_library_from_disk() -> None:
    library = _load_repo_library()
    if not isinstance(library, dict):
        _save_repo_library({})
        return
    migrated: dict = {}
    changed = False
    for old_key, entry in library.items():
        if not isinstance(entry, dict):
            continue
        repo_name, lang_from_key = _parse_library_entry_key(old_key)
        repo_name = (entry.get("repo_name") or repo_name or "").strip()
        if not repo_name:
            continue
        lang = _normalize_doc_language(entry.get("language") or lang_from_key)
        resolved = _resolve_library_entry_assets(repo_name, entry)
        resolved["repo_name"] = repo_name
        resolved["language"] = lang
        new_key = _library_entry_key(repo_name, lang)
        resolved["entry_key"] = new_key
        if new_key != old_key:
            changed = True
        if resolved != entry:
            changed = True
        migrated[new_key] = resolved
    if changed or migrated.keys() != library.keys():
        _save_repo_library(migrated)


def _sync_functional_repo_library_from_disk() -> None:
    library = _load_functional_repo_library()
    if not isinstance(library, dict):
        _save_functional_repo_library({})
        return
    migrated: dict = {}
    changed = False
    for old_key, entry in library.items():
        if not isinstance(entry, dict):
            continue
        repo_name, lang_from_key = _parse_library_entry_key(old_key)
        repo_name = (entry.get("repo_name") or repo_name or "").strip()
        if not repo_name:
            continue
        lang = _normalize_doc_language(entry.get("language") or lang_from_key)
        resolved = _resolve_functional_library_entry_assets(repo_name, entry)
        resolved["repo_name"] = repo_name
        resolved["language"] = lang
        new_key = _library_entry_key(repo_name, lang)
        resolved["entry_key"] = new_key
        if new_key != old_key:
            changed = True
        if resolved != entry:
            changed = True
        migrated[new_key] = resolved
    if changed or migrated.keys() != library.keys():
        _save_functional_repo_library(migrated)


def _migrate_functional_assets_from_technical_library() -> None:
    technical_library = _load_repo_library()
    functional_library = _load_functional_repo_library()
    if not isinstance(technical_library, dict) or not isinstance(functional_library, dict):
        return

    def _path_or_none(value: str | None) -> Path | None:
        raw = str(value or "").strip()
        if not raw or raw in {".", "./", ".\\"}:
            return None
        try:
            return Path(raw)
        except Exception:
            return None

    for entry_key, entry in technical_library.items():
        if not isinstance(entry, dict):
            continue
        parsed_repo, lang_from_key = _parse_library_entry_key(entry_key)
        repo_name = (entry.get("repo_name") or parsed_repo or "").strip()
        if not repo_name:
            continue
        language = _normalize_doc_language(entry.get("language") or lang_from_key)
        resolved = _resolve_library_entry_assets(repo_name, entry)
        if not bool(resolved.get("functional_docs_available")):
            continue

        target_key = _library_entry_key(repo_name, language)
        existing = functional_library.get(target_key)
        if isinstance(existing, dict):
            existing_resolved = _resolve_functional_library_entry_assets(repo_name, existing)
            if bool(existing_resolved.get("functional_docs_available")):
                continue

        functional_docs_src = _path_or_none(resolved.get("library_functional_docs_dir"))
        functional_md_src = _path_or_none(resolved.get("library_functional_documentation_md"))
        if (
            (functional_docs_src is None or not functional_docs_src.exists())
            and (functional_md_src is None or not functional_md_src.exists())
        ):
            continue

        repo_dir = _functional_library_repo_dir_for_language(repo_name, language)
        functional_docs_dst = repo_dir / "docs_functional"
        functional_md_dst = repo_dir / "functional_documentation.md"
        code_json_src = _path_or_none(resolved.get("library_code_json"))
        code_json_dst = repo_dir / "code.json"
        graph_src = _path_or_none(resolved.get("library_graph_json"))
        graph_dst = repo_dir / "graph.json"

        try:
            repo_dir.mkdir(parents=True, exist_ok=True)

            if functional_docs_dst.exists():
                shutil.rmtree(functional_docs_dst)
            if functional_docs_src is not None and functional_docs_src.exists():
                shutil.copytree(functional_docs_src, functional_docs_dst)

            if functional_md_dst.exists():
                functional_md_dst.unlink()
            if functional_md_src is not None and functional_md_src.exists():
                shutil.copy2(functional_md_src, functional_md_dst)

            if code_json_src is not None and code_json_src.exists():
                shutil.copy2(code_json_src, code_json_dst)

            if graph_src is not None and graph_src.exists():
                shutil.copy2(graph_src, graph_dst)
        except Exception:
            continue

        docs_available = bool(
            functional_docs_dst.exists() and any(functional_docs_dst.glob("*.md"))
        )
        functional_docs_available = bool(docs_available or functional_md_dst.exists())
        snapshot_payload = {
            "library_repo_dir": str(repo_dir),
            "library_docs_dir": str(functional_docs_dst),
            "library_documentation_md": str(functional_md_dst),
            "library_functional_docs_dir": str(functional_docs_dst),
            "library_functional_documentation_md": str(functional_md_dst),
            "library_code_json": str(code_json_dst),
            "library_graph_json": str(graph_dst) if graph_dst.exists() else "",
            "docs_available": docs_available,
            "functional_docs_available": functional_docs_available,
        }
        _upsert_functional_repo_library_entry(
            repo_name,
            {
                "repo_path": resolved.get("repo_path"),
                "owner": resolved.get("owner"),
                "output_dir": resolved.get("output_dir"),
                "graph_path": resolved.get("graph_path"),
                "repo_url": resolved.get("repo_url"),
                "source": resolved.get("source"),
                "source_type": resolved.get("source_type"),
                "language": language,
                "docs_available": docs_available,
                "functional_docs_available": functional_docs_available,
                **snapshot_payload,
            },
            language=language,
        )


def _iter_library_repo_variants(repo_name: str, library: dict | None = None) -> list[tuple[str, dict]]:
    repo_name = (repo_name or "").strip()
    if not repo_name:
        return []
    data = library if isinstance(library, dict) else _load_repo_library()
    variants: list[tuple[str, dict]] = []
    for entry_key, entry in data.items():
        if not isinstance(entry, dict):
            continue
        parsed_repo, _ = _parse_library_entry_key(entry_key)
        entry_repo = (entry.get("repo_name") or parsed_repo or "").strip()
        if entry_repo != repo_name:
            continue
        resolved = _resolve_library_entry_assets(entry_repo, entry)
        resolved["repo_name"] = entry_repo
        resolved["language"] = _normalize_doc_language(resolved.get("language"))
        resolved["entry_key"] = entry_key
        variants.append((entry_key, resolved))
    return variants


def _iter_functional_library_repo_variants(
    repo_name: str,
    library: dict | None = None,
) -> list[tuple[str, dict]]:
    repo_name = (repo_name or "").strip()
    if not repo_name:
        return []
    data = library if isinstance(library, dict) else _load_functional_repo_library()
    variants: list[tuple[str, dict]] = []
    for entry_key, entry in data.items():
        if not isinstance(entry, dict):
            continue
        parsed_repo, _ = _parse_library_entry_key(entry_key)
        entry_repo = (entry.get("repo_name") or parsed_repo or "").strip()
        if entry_repo != repo_name:
            continue
        resolved = _resolve_functional_library_entry_assets(entry_repo, entry)
        resolved["repo_name"] = entry_repo
        resolved["language"] = _normalize_doc_language(resolved.get("language"))
        resolved["entry_key"] = entry_key
        variants.append((entry_key, resolved))
    return variants


def _group_library_entries_by_repo(library: dict | None = None) -> dict[str, list[tuple[str, dict]]]:
    data = library if isinstance(library, dict) else _load_repo_library()
    grouped: dict[str, list[tuple[str, dict]]] = {}
    for entry_key, entry in data.items():
        if not isinstance(entry, dict):
            continue
        parsed_repo, _ = _parse_library_entry_key(entry_key)
        repo_name = (entry.get("repo_name") or parsed_repo or "").strip()
        if not repo_name:
            continue
        resolved = _resolve_library_entry_assets(repo_name, entry)
        resolved["repo_name"] = repo_name
        resolved["language"] = _normalize_doc_language(resolved.get("language"))
        resolved["entry_key"] = entry_key
        grouped.setdefault(repo_name, []).append((entry_key, resolved))
    return grouped


def _group_functional_library_entries_by_repo(
    library: dict | None = None,
) -> dict[str, list[tuple[str, dict]]]:
    data = library if isinstance(library, dict) else _load_functional_repo_library()
    grouped: dict[str, list[tuple[str, dict]]] = {}
    for entry_key, entry in data.items():
        if not isinstance(entry, dict):
            continue
        parsed_repo, _ = _parse_library_entry_key(entry_key)
        repo_name = (entry.get("repo_name") or parsed_repo or "").strip()
        if not repo_name:
            continue
        resolved = _resolve_functional_library_entry_assets(repo_name, entry)
        resolved["repo_name"] = repo_name
        resolved["language"] = _normalize_doc_language(resolved.get("language"))
        resolved["entry_key"] = entry_key
        grouped.setdefault(repo_name, []).append((entry_key, resolved))
    return grouped


def _pick_library_repo_variant(
    repo_name: str,
    preferred_language: str | None = None,
    require_docs: bool = False,
    docs_kind: str = "technical",
    library: dict | None = None,
) -> tuple[str | None, dict | None]:
    variants = _iter_library_repo_variants(repo_name, library=library)
    if not variants:
        return None, None

    docs_kind_norm = str(docs_kind or "technical").strip().lower()

    def _entry_has_docs(entry: dict) -> bool:
        if docs_kind_norm == "functional":
            return bool(entry.get("functional_docs_available"))
        return bool(entry.get("docs_available"))

    preferred_lang = _normalize_doc_language(preferred_language)
    scored: list[tuple[int, int, str, str, dict]] = []
    for entry_key, entry in variants:
        docs_available = _entry_has_docs(entry)
        if require_docs and not docs_available:
            continue
        entry_lang = _normalize_doc_language(entry.get("language"))
        score_docs = 1 if docs_available else 0
        score_lang = 1 if preferred_language and entry_lang == preferred_lang else 0
        updated = str(entry.get("updated_at") or "")
        scored.append((score_docs, score_lang, updated, entry_key, entry))

    if not scored and require_docs:
        return None, None
    if not scored:
        for entry_key, entry in variants:
            entry_lang = _normalize_doc_language(entry.get("language"))
            score_docs = 1 if _entry_has_docs(entry) else 0
            score_lang = 1 if preferred_language and entry_lang == preferred_lang else 0
            updated = str(entry.get("updated_at") or "")
            scored.append((score_docs, score_lang, updated, entry_key, entry))

    scored.sort(key=lambda row: (row[0], row[1], row[2]), reverse=True)
    best = scored[0]
    return best[3], best[4]


def _pick_functional_library_repo_variant(
    repo_name: str,
    preferred_language: str | None = None,
    require_docs: bool = False,
    library: dict | None = None,
) -> tuple[str | None, dict | None]:
    variants = _iter_functional_library_repo_variants(repo_name, library=library)
    if not variants:
        return None, None

    preferred_lang = _normalize_doc_language(preferred_language)
    scored: list[tuple[int, int, str, str, dict]] = []
    for entry_key, entry in variants:
        docs_available = bool(entry.get("functional_docs_available"))
        if require_docs and not docs_available:
            continue
        entry_lang = _normalize_doc_language(entry.get("language"))
        score_docs = 1 if docs_available else 0
        score_lang = 1 if preferred_language and entry_lang == preferred_lang else 0
        updated = str(entry.get("updated_at") or "")
        scored.append((score_docs, score_lang, updated, entry_key, entry))

    if not scored and require_docs:
        return None, None
    if not scored:
        for entry_key, entry in variants:
            entry_lang = _normalize_doc_language(entry.get("language"))
            score_docs = 1 if bool(entry.get("functional_docs_available")) else 0
            score_lang = 1 if preferred_language and entry_lang == preferred_lang else 0
            updated = str(entry.get("updated_at") or "")
            scored.append((score_docs, score_lang, updated, entry_key, entry))

    scored.sort(key=lambda row: (row[0], row[1], row[2]), reverse=True)
    best = scored[0]
    return best[3], best[4]


def _load_code_index(code_json_path: Path):
    with open(code_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    items = []
    for item in data:
        if item.get("ignore"):
            continue
        text = " ".join(
            [
                item.get("path", ""),
                item.get("tipo", ""),
                item.get("nome", "") or "",
                item.get("resumo", "") or "",
            ]
        ).lower()
        items.append({"item": item, "text": text})
    return items


def _simple_search(index, query, k=5):
    import re
    terms = re.findall(r"[a-zA-Z_][a-zA-Z0-9_]*", query.lower())
    terms = [t for t in terms if len(t) >= 2]
    if not terms:
        return []
    scored = []
    for it in index:
        item = it["item"]
        text = it["text"]
        code = (item.get("code") or "").lower()
        name = (item.get("nome") or "").lower()
        score = 0
        for t in terms:
            score += text.count(t)
            if code:
                score += code.count(t)
            if name and t == name:
                score += 50
        if score:
            scored.append((score, item))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [i for _, i in scored[:k]]


def _extract_documentation_fallback_sections(doc_text: str, max_chars: int = 9000) -> str:
    import re

    text = str(doc_text or "")
    if not text.strip():
        return ""
    budget = max(1, int(max_chars))

    marker_re = re.compile(r"^\s*<!--\s*SECTION:([A-Z_]+)\s*-->\s*$")
    wanted_keys = ["INDEX", "USAGE", "ARCHITECTURE", "TECHNOLOGIES"]
    sections: dict[str, list[str]] = {}
    section_order: list[str] = []
    current_key = None
    current_lines: list[str] = []

    for raw_line in text.splitlines():
        match = marker_re.match(raw_line.strip())
        if match:
            if current_key is not None:
                sections[current_key] = list(current_lines)
                section_order.append(current_key)
            current_key = match.group(1).strip().upper()
            current_lines = [raw_line]
            continue
        if current_key is not None:
            current_lines.append(raw_line)

    if current_key is not None:
        sections[current_key] = list(current_lines)
        section_order.append(current_key)

    if not sections:
        return text.strip()[:budget].strip()

    selected_blocks: list[str] = []
    selected_keys: set[str] = set()
    used_chars = 0

    def _append_block(block: str) -> None:
        nonlocal used_chars
        part = str(block or "").strip()
        if not part or used_chars >= budget:
            return
        extra = len(part) + (2 if selected_blocks else 0)
        if not selected_blocks and len(part) > budget:
            selected_blocks.append(part[:budget].strip())
            used_chars = len(selected_blocks[-1])
            return
        if selected_blocks and used_chars + extra > budget:
            remaining = max(0, budget - used_chars - 2)
            if remaining > 0:
                selected_blocks.append(part[:remaining].strip())
                used_chars = budget
            return
        selected_blocks.append(part)
        used_chars += extra

    for key in wanted_keys:
        block_lines = sections.get(key)
        if not block_lines:
            continue
        _append_block("\n".join(block_lines))
        selected_keys.add(key)

    for key in section_order:
        if key in selected_keys:
            continue
        block_lines = sections.get(key)
        if not block_lines:
            continue
        _append_block("\n".join(block_lines))
        if used_chars >= budget:
            break

    if selected_blocks:
        return "\n\n".join(selected_blocks).strip()

    return text.strip()[:budget].strip()


def _extract_documentation_context(doc_md_path: Path, query: str, k: int = 5, max_chars: int = 9000) -> str:
    import re

    if not doc_md_path.exists():
        return ""

    try:
        doc_text = doc_md_path.read_text(encoding="utf-8")
    except Exception:
        return ""

    terms = re.findall(r"[a-zA-Z_][a-zA-Z0-9_]*", (query or "").lower())
    terms = [t for t in terms if len(t) >= 2]
    if not terms:
        return _extract_documentation_fallback_sections(doc_text, max_chars=max_chars)

    blocks = [b.strip() for b in re.split(r"\n\s*\n", doc_text) if b.strip()]
    if not blocks:
        return ""

    scored = []
    for idx, block in enumerate(blocks):
        lower = block.lower()
        score = 0
        for t in terms:
            score += lower.count(t)
        if score > 0:
            # Preserve original order as tie-breaker after relevance.
            scored.append((score, -idx, block))

    if not scored:
        return _extract_documentation_fallback_sections(doc_text, max_chars=max_chars)

    scored.sort(reverse=True)
    selected = []
    used = 0
    for _, _, block in scored[: max(1, int(k))]:
        part = block.strip()
        if not part:
            continue
        extra = len(part) + (2 if selected else 0)
        if selected and used + extra > max_chars:
            break
        if not selected and len(part) > max_chars:
            selected.append(part[:max_chars].strip())
            used = len(selected[-1])
            break
        selected.append(part)
        used += extra

    return "\n\n".join(selected).strip()


def _sanitize_pdf_text(text: str) -> str:
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2022": "-",
        "\u2026": "...",
        "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", "ignore").decode("latin-1")


def _strip_inline_markdown(text: str) -> str:
    import re

    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = re.sub(r"[*_]{1,3}(.+?)[*_]{1,3}", r"\1", text)
    text = text.replace("`", "")
    return text


def _markdown_to_pdf_blocks(md_text: str) -> list[tuple[str, str]]:
    import re

    lines = md_text.splitlines()
    blocks: list[tuple[str, str]] = []
    in_code = False
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue
        if in_code:
            blocks.append(("code", _sanitize_pdf_text(line)))
            continue
        if not stripped:
            blocks.append(("blank", ""))
            continue
        if re.match(r"^\s*([-*_]\s*){3,}$", line):
            blocks.append(("hr", ""))
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading:
            level = len(heading.group(1))
            text = _sanitize_pdf_text(_strip_inline_markdown(heading.group(2).strip()))
            blocks.append((f"h{level}", text))
            continue
        bullet = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if bullet:
            indent = len(bullet.group(1))
            marker = bullet.group(2).strip()
            text = _sanitize_pdf_text(_strip_inline_markdown(bullet.group(3).strip()))
            blocks.append(("li", f"{indent}\x1f{marker}\x1f{text}"))
            continue
        if stripped.startswith(">"):
            text = _sanitize_pdf_text(_strip_inline_markdown(stripped[1:].strip()))
            blocks.append(("quote", text))
            continue
        blocks.append(("p", _sanitize_pdf_text(_strip_inline_markdown(stripped))))
    return blocks


def _normalize_pdf_blocks(blocks: list[tuple[str, str]]) -> list[tuple[str, str]]:
    merged: list[tuple[str, str]] = []
    buffer: list[str] = []
    code_buffer: list[str] = []
    for kind, text in blocks:
        if kind == "code":
            code_buffer.append(text)
            continue
        if code_buffer:
            merged.append(("code", "\n".join(code_buffer)))
            code_buffer = []
        if kind == "p":
            buffer.append(text)
            continue
        if buffer:
            merged.append(("p", " ".join(buffer)))
            buffer = []
        merged.append((kind, text))
    if buffer:
        merged.append(("p", " ".join(buffer)))
    if code_buffer:
        merged.append(("code", "\n".join(code_buffer)))
    return merged


def _markdown_to_docx_blocks(md_text: str) -> list[tuple[str, str]]:
    import re

    lines = md_text.splitlines()
    blocks: list[tuple[str, str]] = []
    in_code = False
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue
        if in_code:
            blocks.append(("code", line))
            continue
        if not stripped:
            blocks.append(("blank", ""))
            continue
        if re.match(r"^\s*([-*_]\s*){3,}$", line):
            blocks.append(("hr", ""))
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading:
            level = len(heading.group(1))
            text = _strip_inline_markdown(heading.group(2).strip())
            blocks.append((f"h{level}", text))
            continue
        bullet = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if bullet:
            indent = len(bullet.group(1))
            marker = bullet.group(2).strip()
            text = _strip_inline_markdown(bullet.group(3).strip())
            blocks.append(("li", f"{indent}\x1f{marker}\x1f{text}"))
            continue
        if stripped.startswith(">"):
            text = _strip_inline_markdown(stripped[1:].strip())
            blocks.append(("quote", text))
            continue
        blocks.append(("p", _strip_inline_markdown(stripped)))
    return blocks


def _normalize_doc_language(language: str | None) -> str:
    value = (language or "").strip().upper()
    if value.startswith("PT"):
        return "PT-BR"
    if value.startswith("EN"):
        return "EN-US"
    if value.startswith("ES"):
        return "ES-ES"
    if value.startswith("FR"):
        return "FR-FR"
    if value.startswith("DE"):
        return "DE-DE"
    return "EN-US"


def _pdf_i18n(
    language: str,
    repo_name: str,
    generated_date: str,
    doc_kind: str = "technical",
) -> dict:
    lang = _normalize_doc_language(language)
    texts = {
        "PT-BR": {
            "doc_subtitle": "Documentação Técnica",
            "generated_on": "Gerado em",
            "abstract_title": "Resumo",
            "abstract_body": (
                f"O objetivo deste documento é detalhar tecnicamente o repositório {repo_name}. "
                f"Esta documentação foi gerada automaticamente em {generated_date} com uso do Repository Documentation "
                "e necessita de revisão humana antes de publicação final. "
                "O conteúdo inclui visão geral, instalação, uso, arquitetura, tecnologias, referência de API, "
                "configuração, testes e análise arquivo a arquivo, além de artefatos para consulta técnica."
            ),
            "summary_title": "Sumário",
            "version": "Versão",
            "changed_by": "Alterado por",
            "change_date": "Data da alteração",
            "page_label": "Página",
        },
        "EN-US": {
            "doc_subtitle": "Technical Documentation",
            "generated_on": "Generated on",
            "abstract_title": "Abstract",
            "abstract_body": (
                f"The objective of this document is to provide a technical description of the {repo_name} repository. "
                f"This documentation was automatically generated on {generated_date} using Repository Documentation "
                "and requires human review before final publication. "
                "It includes overview, installation, usage, architecture, technology stack, API reference, "
                "configuration, testing, and file-by-file analysis, plus artifacts for technical consultation."
            ),
            "summary_title": "Summary",
            "version": "Version",
            "changed_by": "Changed by",
            "change_date": "Change date",
            "page_label": "Page",
        },
        "ES-ES": {
            "doc_subtitle": "Documentación Técnica",
            "generated_on": "Generado el",
            "abstract_title": "Resumen",
            "abstract_body": (
                f"El objetivo de este documento es detallar técnicamente el repositorio {repo_name}. "
                f"Esta documentación fue generada automáticamente el {generated_date} usando Repository Documentation "
                "y requiere revisión humana antes de su publicación final. "
                "Incluye visión general, instalación, uso, arquitectura, tecnologías, referencia de API, "
                "configuración, pruebas y análisis archivo por archivo, además de artefactos para consulta técnica."
            ),
            "summary_title": "Resumen",
            "version": "Versión",
            "changed_by": "Modificado por",
            "change_date": "Fecha de cambio",
            "page_label": "Página",
        },
        "FR-FR": {
            "doc_subtitle": "Documentation Technique",
            "generated_on": "Généré le",
            "abstract_title": "Résumé",
            "abstract_body": (
                f"L'objectif de ce document est de décrire techniquement le dépôt {repo_name}. "
                f"Cette documentation a été générée automatiquement le {generated_date} avec Repository Documentation "
                "et nécessite une révision humaine avant publication finale. "
                "Elle couvre la vue d'ensemble, l'installation, l'utilisation, l'architecture, les technologies, "
                "la référence API, la configuration, les tests et l'analyse fichier par fichier, avec des artefacts de consultation technique."
            ),
            "summary_title": "Sommaire",
            "version": "Version",
            "changed_by": "Modifié par",
            "change_date": "Date de modification",
            "page_label": "Page",
        },
        "DE-DE": {
            "doc_subtitle": "Technische Dokumentation",
            "generated_on": "Erstellt am",
            "abstract_title": "Zusammenfassung",
            "abstract_body": (
                f"Ziel dieses Dokuments ist die technische Beschreibung des Repositorys {repo_name}. "
                f"Diese Dokumentation wurde am {generated_date} automatisch mit Repository Documentation erstellt "
                "und muss vor der finalen Veröffentlichung manuell geprüft werden. "
                "Sie umfasst Überblick, Installation, Nutzung, Architektur, Technologien, API-Referenz, "
                "Konfiguration, Tests und eine dateiweise Analyse sowie Artefakte für die technische Konsultation."
            ),
            "summary_title": "Inhaltsverzeichnis",
            "version": "Version",
            "changed_by": "Geändert von",
            "change_date": "Änderungsdatum",
            "page_label": "Seite",
        },
    }
    payload = dict(texts.get(lang, texts["EN-US"]))
    payload["changed_by_value"] = "Repository Documentation"

    doc_kind_norm = str(doc_kind or "technical").strip().lower()
    if doc_kind_norm.startswith("func"):
        functional_overrides = {
            "PT-BR": {
                "doc_subtitle": "Documentação Funcional",
                "abstract_body": (
                    f"O objetivo deste documento é descrever funcionalmente o repositório {repo_name}. "
                    f"Esta documentação foi gerada automaticamente em {generated_date} com uso do Repository Documentation "
                    "e necessita de revisão humana antes de publicação final. "
                    "O foco está em escopo de negócio, fluxos funcionais, regras e exceções, e procedimentos operacionais."
                ),
            },
            "EN-US": {
                "doc_subtitle": "Functional Documentation",
                "abstract_body": (
                    f"The objective of this document is to describe the {repo_name} repository from a functional perspective. "
                    f"This documentation was automatically generated on {generated_date} using Repository Documentation "
                    "and requires human review before final publication. "
                    "It focuses on business scope, functional flows, rules and exceptions, and operational procedures."
                ),
            },
            "ES-ES": {
                "doc_subtitle": "Documentación Funcional",
                "abstract_body": (
                    f"El objetivo de este documento es describir funcionalmente el repositorio {repo_name}. "
                    f"Esta documentación fue generada automáticamente el {generated_date} usando Repository Documentation "
                    "y requiere revisión humana antes de su publicación final. "
                    "Se centra en alcance de negocio, flujos funcionales, reglas y excepciones, y procedimientos operativos."
                ),
            },
            "FR-FR": {
                "doc_subtitle": "Documentation Fonctionnelle",
                "abstract_body": (
                    f"L'objectif de ce document est de décrire fonctionnellement le dépôt {repo_name}. "
                    f"Cette documentation a été générée automatiquement le {generated_date} avec Repository Documentation "
                    "et nécessite une révision humaine avant publication finale. "
                    "Elle se concentre sur le périmètre métier, les flux fonctionnels, les règles/exceptions et les procédures opérationnelles."
                ),
            },
            "DE-DE": {
                "doc_subtitle": "Funktionale Dokumentation",
                "abstract_body": (
                    f"Ziel dieses Dokuments ist die funktionale Beschreibung des Repositorys {repo_name}. "
                    f"Diese Dokumentation wurde am {generated_date} automatisch mit Repository Documentation erstellt "
                    "und muss vor der finalen Veröffentlichung manuell geprüft werden. "
                    "Der Fokus liegt auf fachlichem Umfang, funktionalen Abläufen, Regeln/Ausnahmen und operativen Verfahren."
                ),
            },
        }
        payload.update(functional_overrides.get(lang, functional_overrides["EN-US"]))
    return payload


def _build_pdf_bytes_fpdf(
    md_text: str,
    title: str,
    language: str = "EN-US",
    doc_kind: str = "technical",
) -> bytes:
    global PDF_RENDER_FONT
    try:
        from fpdf import FPDF, XPos, YPos
    except Exception as e:
        raise RuntimeError(
            "fpdf2 is not installed. Install it with: pip install fpdf2"
        ) from e

    lang = _normalize_doc_language(language)
    date_format = "%Y-%m-%d" if lang == "EN-US" else "%d/%m/%Y"
    generated_date = datetime.now().strftime(date_format)
    i18n = _pdf_i18n(lang, title, generated_date, doc_kind=doc_kind)

    blocks = _normalize_pdf_blocks(_markdown_to_pdf_blocks(md_text))
    user_fonts_dir = Path.home() / "AppData" / "Local" / "Microsoft" / "Windows" / "Fonts"

    def _first_existing_font(candidates: list[Path]) -> Path | None:
        for candidate in candidates:
            try:
                if candidate.exists():
                    return candidate
            except Exception:
                continue
        return None

    def _resolve_font_family_paths() -> list[
        tuple[str, str, Path | None, Path | None, Path | None, Path | None]
    ]:
        return [
            (
                "HubiaCalibri",
                "Calibri",
                _first_existing_font([Path(r"C:\Windows\Fonts\calibri.ttf"), user_fonts_dir / "calibri.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\calibrib.ttf"), user_fonts_dir / "calibrib.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\calibrii.ttf"), user_fonts_dir / "calibrii.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\calibriz.ttf"), user_fonts_dir / "calibriz.ttf"]),
            ),
            (
                "HubiaArial",
                "Arial",
                _first_existing_font([Path(r"C:\Windows\Fonts\arial.ttf"), user_fonts_dir / "arial.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\arialbd.ttf"), user_fonts_dir / "arialbd.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\ariali.ttf"), user_fonts_dir / "ariali.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\arialbi.ttf"), user_fonts_dir / "arialbi.ttf"]),
            ),
            (
                "HubiaSegoeUI",
                "Segoe UI",
                _first_existing_font([Path(r"C:\Windows\Fonts\segoeui.ttf"), user_fonts_dir / "segoeui.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\segoeuib.ttf"), user_fonts_dir / "segoeuib.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\segoeuii.ttf"), user_fonts_dir / "segoeuii.ttf"]),
                _first_existing_font([Path(r"C:\Windows\Fonts\seguisbi.ttf"), user_fonts_dir / "seguisbi.ttf"]),
            ),
        ]

    def _register_font_family(
        pdf: FPDF,
        family: str,
        regular_src: Path | None,
        bold_src: Path | None,
        italic_src: Path | None,
        bold_italic_src: Path | None,
    ) -> bool:
        if regular_src is None:
            return False
        try:
            use_bold = bold_src or regular_src
            use_italic = italic_src or regular_src
            use_bold_italic = bold_italic_src or use_bold or use_italic or regular_src
            pdf.add_font(family, "", str(regular_src))
            pdf.add_font(family, "B", str(use_bold))
            pdf.add_font(family, "I", str(use_italic))
            pdf.add_font(family, "BI", str(use_bold_italic))
            return True
        except Exception:
            return False

    def _font_family_is_crisp(
        family: str,
        regular_src: Path | None,
        bold_src: Path | None,
        italic_src: Path | None,
        bold_italic_src: Path | None,
    ) -> bool:
        if regular_src is None:
            return False
        try:
            probe = FPDF()
            probe.set_compression(False)
            if not _register_font_family(
                probe,
                family,
                regular_src=regular_src,
                bold_src=bold_src,
                italic_src=italic_src,
                bold_italic_src=bold_italic_src,
            ):
                return False
            probe.add_page()
            probe.set_font(family, size=10)
            probe.cell(0, 6, "Font probe")
            out = probe.output()
            if isinstance(out, bytearray):
                out = bytes(out)
            elif isinstance(out, str):
                out = out.encode("latin-1", "ignore")
            return b"/Subtype /Type3" not in out
        except Exception:
            return False

    main_font = "Helvetica"
    chosen_family = ""
    chosen_family_label = ""
    chosen_regular = None
    chosen_bold = None
    chosen_italic = None
    chosen_bold_italic = None

    for fam, fam_label, reg, bld, ita, bi in _resolve_font_family_paths():
        if reg is None:
            continue
        if _font_family_is_crisp(fam, reg, bld, ita, bi):
            chosen_family = fam
            chosen_family_label = fam_label
            chosen_regular, chosen_bold, chosen_italic, chosen_bold_italic = reg, bld, ita, bi
            break

    paragraph_after_mm = 0.8
    heading_after_mm = 3.6
    page_margin = 18
    bottom_margin = 18

    def register_fonts(pdf: FPDF):
        if not chosen_family or chosen_regular is None:
            return False
        return _register_font_family(
            pdf,
            chosen_family,
            regular_src=chosen_regular,
            bold_src=chosen_bold,
            italic_src=chosen_italic,
            bold_italic_src=chosen_bold_italic,
        )

    class DocPDF(FPDF):
        def footer(self):
            if self.page_no() <= 3:
                return
            self.set_y(-10)
            self.set_font(main_font, size=9)
            self.set_text_color(120, 120, 120)
            self.cell(0, 5, f"{i18n['page_label']} {self.page_no()}", align="R")
            self.set_text_color(0, 0, 0)

    # Base helper PDF for width calculations
    helper = DocPDF()
    if register_fonts(helper):
        main_font = chosen_family
        PDF_RENDER_FONT = f"{chosen_family_label or chosen_family} (embedded)"
    else:
        PDF_RENDER_FONT = "Helvetica (core)"
    helper.set_margins(page_margin, page_margin, page_margin)
    helper.set_auto_page_break(auto=True, margin=bottom_margin)
    helper.set_font(main_font, size=10)
    page_width = helper.w - helper.l_margin - helper.r_margin
    if page_width <= 0:
        page_width = 100

    def safe_width(pdf: DocPDF, width: float) -> float:
        min_w = max(
            pdf.get_string_width("W"),
            pdf.get_string_width("M"),
            pdf.get_string_width("m"),
        ) + 0.5
        return max(width, min_w)

    def wrap_line(pdf: DocPDF, text: str, width: float) -> list[str]:
        width = safe_width(pdf, width)
        words = text.split(" ")
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip() if current else word
            if pdf.get_string_width(candidate) <= width:
                current = candidate
                continue
            if current:
                lines.append(current)
                current = ""
            if pdf.get_string_width(word) <= width:
                current = word
                continue
            chunk = ""
            for ch in word:
                if pdf.get_string_width(chunk + ch) <= width:
                    chunk += ch
                else:
                    if chunk:
                        lines.append(chunk)
                    chunk = ch
            if chunk:
                current = chunk
        if current:
            lines.append(current)
        return lines

    def normalize_text_for_justify(pdf: DocPDF, text: str, width: float) -> str:
        # Ensures no token exceeds the render width to avoid FPDF width errors.
        words = text.split(" ")
        normalized: list[str] = []
        for word in words:
            if not word:
                continue
            if pdf.get_string_width(word) <= width:
                normalized.append(word)
                continue
            normalized.extend(wrap_line(pdf, word, width))
        return " ".join(normalized)

    def render_lines(
        pdf: DocPDF,
        lines: list[str],
        line_height: float = 6,
        width: float | None = None,
        x: float | None = None,
        align: str = "L",
    ):
        use_width = safe_width(pdf, width or page_width)
        for entry in lines:
            if not entry:
                pdf.ln(line_height)
                continue
            if x is not None:
                pdf.set_x(x)
            else:
                pdf.set_x(pdf.l_margin)
            pdf.cell(use_width, line_height, entry, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align=align)

    def render_paragraph_justified(pdf: DocPDF, text: str, line_height: float = 6):
        use_width = safe_width(pdf, page_width)
        normalized = normalize_text_for_justify(pdf, text, use_width)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(use_width, line_height, normalized, align="J", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(paragraph_after_mm)

    def render_change_table(pdf: DocPDF):
        col1 = 28.0
        col2 = 92.0
        col3 = page_width - col1 - col2
        headers = [i18n["version"], i18n["changed_by"], i18n["change_date"]]
        row = ["1.0", i18n["changed_by_value"], generated_date]

        pdf.set_fill_color(240, 242, 245)
        pdf.set_draw_color(190, 196, 205)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font(main_font, "B", 10)
        pdf.cell(col1, 8, headers[0], border=1, fill=True)
        pdf.cell(col2, 8, headers[1], border=1, fill=True)
        pdf.cell(col3, 8, headers[2], border=1, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font(main_font, "", 10)
        pdf.cell(col1, 8, row[0], border=1)
        pdf.cell(col2, 8, row[1], border=1)
        pdf.cell(col3, 8, row[2], border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def render_content(pdf: DocPDF, collect_toc: bool = False, page_offset: int = 0) -> list[tuple[int, str, int]]:
        toc_entries: list[tuple[int, str, int]] = []
        pdf.add_page()
        pdf.set_font(main_font, size=10)
        max_width = page_width
        has_content = False

        for kind, text in blocks:
            if kind == "blank":
                pdf.ln(4)
                continue
            if kind == "hr":
                y = pdf.get_y()
                pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
                pdf.ln(4)
                continue
            if kind.startswith("h"):
                level = int(kind[1:])
                if level == 1 and has_content:
                    pdf.add_page()
                if collect_toc and level in (1, 2):
                    toc_entries.append((level, text, pdf.page_no() + page_offset))
                heading_size = {1: 18, 2: 13, 3: 11, 4: 10.5, 5: 10, 6: 9.5}.get(level, 10.5)
                heading_line_height = 8 if level == 1 else (6 if level == 2 else 5.5)
                pdf.set_font(main_font, "B", heading_size)
                render_lines(
                    pdf,
                    wrap_line(pdf, text, max_width),
                    line_height=heading_line_height,
                    width=max_width,
                )
                pdf.ln(heading_after_mm)
                pdf.set_font(main_font, size=10)
                has_content = True
                continue
            if kind == "code":
                # Keep code blocks in the same visual family used by HTML/PDF body text.
                pdf.set_font(main_font, size=9)
                code_lines: list[str] = []
                for raw in text.splitlines():
                    if not raw.strip():
                        code_lines.append("")
                        continue
                    code_lines.extend(wrap_line(pdf, raw, max_width - 6))
                if not code_lines:
                    code_lines = [""]
                line_height = 5
                block_height = line_height * len(code_lines) + 6
                x = pdf.l_margin
                y = pdf.get_y()
                if y + block_height > pdf.h - pdf.b_margin:
                    pdf.add_page()
                    y = pdf.get_y()
                pdf.set_fill_color(245, 247, 250)
                pdf.set_draw_color(220, 224, 230)
                pdf.rect(x, y, max_width, block_height, "DF")
                inner_x = x + 3
                inner_width = max_width - 6
                pdf.set_xy(inner_x, y + 3)
                for line in code_lines:
                    render_lines(pdf, [line], line_height=line_height, width=inner_width, x=inner_x)
                pdf.set_y(y + block_height + 2)
                pdf.set_x(pdf.l_margin)
                pdf.set_font(main_font, size=10)
                has_content = True
                continue
            if kind == "quote":
                pdf.set_font(main_font, "I", 10)
                render_paragraph_justified(pdf, f"> {text}", line_height=6)
                pdf.set_font(main_font, size=10)
                has_content = True
                continue
            if kind == "li":
                indent = 0
                marker = "-"
                item_text = text
                if "\x1f" in text:
                    parts = text.split("\x1f", 2)
                    if len(parts) == 3:
                        indent_raw, marker_raw, item_text = parts
                        try:
                            indent = max(0, int(indent_raw))
                        except Exception:
                            indent = 0
                        marker = marker_raw.strip() or "-"

                list_level = min(indent // 2, 6)
                indent_mm = list_level * 4
                ordered = marker.endswith(".") and marker[:-1].isdigit()
                prefix = f"{marker} " if ordered else "· "
                prefix_width = pdf.get_string_width(prefix)
                line_width = max(10, max_width - indent_mm)
                first_width = max(10, line_width - prefix_width)
                wrapped = wrap_line(pdf, item_text, first_width)
                if not wrapped:
                    wrapped = [""]

                x_start = pdf.l_margin + indent_mm
                render_lines(pdf, [prefix + wrapped[0]], width=line_width, x=x_start)
                if len(wrapped) > 1:
                    continuation_x = x_start + prefix_width
                    continuation_w = max(10, line_width - prefix_width)
                    for continuation in wrapped[1:]:
                        render_lines(pdf, [continuation], width=continuation_w, x=continuation_x)
                pdf.ln(paragraph_after_mm)
                has_content = True
                continue
            render_paragraph_justified(pdf, text, line_height=6)
            has_content = True

        return toc_entries

    # Pass 1: compute exact section page numbers (content starts after cover+back cover+summary)
    probe_pdf = DocPDF()
    if main_font != "Helvetica":
        if not register_fonts(probe_pdf):
            main_font = "Helvetica"
            PDF_RENDER_FONT = "Helvetica (core)"
    probe_pdf.set_margins(page_margin, page_margin, page_margin)
    probe_pdf.set_auto_page_break(auto=True, margin=bottom_margin)
    toc_entries = render_content(probe_pdf, collect_toc=True, page_offset=3)

    # Pass 2: final document
    pdf = DocPDF()
    if main_font != "Helvetica":
        if not register_fonts(pdf):
            main_font = "Helvetica"
            PDF_RENDER_FONT = "Helvetica (core)"
    pdf.set_margins(page_margin, page_margin, page_margin)
    pdf.set_auto_page_break(auto=True, margin=bottom_margin)

    # Cover: centered title/subtitle
    pdf.add_page()
    pdf.set_font(main_font, "B", 26)
    center_y = (pdf.h / 2) - 10
    pdf.set_y(center_y)
    render_lines(pdf, wrap_line(pdf, title, page_width), line_height=12, width=page_width, align="C")
    pdf.set_font(main_font, size=18)
    render_lines(pdf, [i18n["doc_subtitle"]], line_height=10, width=page_width, align="C")
    pdf.set_y(pdf.h - pdf.b_margin - 6)
    pdf.set_font(main_font, "I", 10)
    pdf.cell(page_width, 4, f"{i18n['generated_on']} {generated_date}", align="R")

    # Back cover: abstract + change table
    pdf.add_page()
    pdf.set_font(main_font, "B", 16)
    render_lines(pdf, [i18n["abstract_title"]], line_height=8, width=page_width)
    pdf.ln(heading_after_mm)
    pdf.set_font(main_font, "", 10)
    render_paragraph_justified(pdf, i18n["abstract_body"], line_height=6)
    pdf.ln(2)
    render_change_table(pdf)

    # Summary with dotted leaders
    pdf.add_page()
    pdf.set_font(main_font, "B", 10)
    render_lines(pdf, [i18n["summary_title"]], line_height=7, width=page_width)
    pdf.ln(2)
    pdf.set_font(main_font, size=10)
    toc_page_col = 14.0
    toc_gap = 2.0
    for level, item, page_no in toc_entries:
        label = item.strip()
        indent_mm = max(0, level - 1) * 4.0
        x_start = pdf.l_margin + indent_mm
        usable_width = max(20.0, page_width - indent_mm)
        page_col = min(toc_page_col, usable_width * 0.35)
        left_col = max(10.0, usable_width - page_col - toc_gap)
        page_text = str(page_no)
        max_label_width = max(6.0, left_col * 0.62)
        while label and pdf.get_string_width(label) > max_label_width:
            label = label[:-1]
        dot_width = max(0.0, left_col - pdf.get_string_width(label) - 1.0)
        dot_unit = max(pdf.get_string_width("."), 0.3)
        dot_count = max(3, int(dot_width / dot_unit))
        dots = "." * dot_count
        left_text = f"{label} {dots}"
        y = pdf.get_y()
        pdf.set_xy(x_start, y)
        pdf.cell(left_col, 6, left_text, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.cell(page_col, 6, page_text, align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Content
    render_content(pdf, collect_toc=False, page_offset=0)

    data = pdf.output()
    if isinstance(data, str):
        data = data.encode("latin-1", "ignore")
    if isinstance(data, bytearray):
        data = bytes(data)
    return data


def _split_markdown_sections(md_text: str) -> list[tuple[str, str]]:
    import re

    sections: list[tuple[str, str]] = []
    current_title: str | None = None
    buffer: list[str] = []

    for raw in md_text.splitlines():
        line = raw.rstrip("\n")
        stripped = line.strip()
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue
        heading = re.match(r"^#\s+(.*)", stripped)
        if heading:
            if current_title is not None:
                sections.append((current_title, "\n".join(buffer).strip()))
            current_title = _sanitize_pdf_text(_strip_inline_markdown(heading.group(1).strip()))
            buffer = []
            continue
        if current_title is not None:
            buffer.append(line)

    if current_title is not None:
        sections.append((current_title, "\n".join(buffer).strip()))

    if not sections:
        cleaned = md_text.strip()
        if cleaned:
            sections.append(("Overview", cleaned))
    return sections


def _markdown_to_html_fragment(md_text: str) -> str:
    try:
        import markdown as md
    except Exception as e:
        raise RuntimeError(
            "Markdown renderer unavailable. Install with: pip install markdown"
        ) from e

    extensions = [
        "extra",
        "sane_lists",
        "tables",
        "fenced_code",
        "codehilite",
        "admonition",
        "attr_list",
        "nl2br",
    ]
    return md.markdown(md_text or "", extensions=extensions, output_format="html5")


def _playwright_runtime_precheck() -> None:
    # Streamlit on Windows can run with Selector policy, which does not support
    # subprocess creation used by Playwright. Try switching to Proactor policy
    # before giving up.
    if os.name != "nt":
        return
    try:
        import asyncio

        selector_policy_cls = getattr(asyncio, "WindowsSelectorEventLoopPolicy", None)
        proactor_policy_cls = getattr(asyncio, "WindowsProactorEventLoopPolicy", None)
        policy = asyncio.get_event_loop_policy()
        if selector_policy_cls and isinstance(policy, selector_policy_cls):
            if proactor_policy_cls is None:
                raise RuntimeError(
                    "Playwright HTML backend is not available with Windows Selector event loop policy in this runtime."
                )
            try:
                asyncio.set_event_loop_policy(proactor_policy_cls())
            except Exception as switch_err:
                raise RuntimeError(
                    "Playwright HTML backend is not available with Windows Selector event loop policy in this runtime."
                ) from switch_err
    except RuntimeError:
        raise
    except Exception:
        # If the check itself fails, keep default behavior and let Playwright decide.
        return


def _build_pdf_bytes_html(
    md_text: str,
    title: str,
    language: str = "EN-US",
    doc_kind: str = "technical",
) -> bytes:
    global PDF_RENDER_FONT
    import base64
    import html as html_lib
    import json as json_lib

    try:
        from playwright.sync_api import sync_playwright
    except Exception as e:
        raise RuntimeError(
            "Playwright is required for HTML/CSS PDF. Install with: pip install playwright "
            "and then run: python -m playwright install chromium"
        ) from e

    _playwright_runtime_precheck()

    lang = _normalize_doc_language(language)
    date_format = "%Y-%m-%d" if lang == "EN-US" else "%d/%m/%Y"
    generated_date = datetime.now().strftime(date_format)
    i18n = _pdf_i18n(lang, title, generated_date, doc_kind=doc_kind)
    sections = _split_markdown_sections(md_text)

    if not sections:
        sections = [("Overview", "")]

    toc_rows: list[str] = []
    content_sections: list[str] = []
    for idx, (section_title, section_md) in enumerate(sections, start=1):
        sec_id = f"sec-{idx}"
        safe_title = html_lib.escape(section_title or f"Section {idx}")
        body_html = _markdown_to_html_fragment(section_md or "")
        toc_rows.append(
            (
                '<div class="toc-row">'
                f'<span class="toc-label">{safe_title}</span>'
                '<span class="toc-dots"></span>'
                f'<span class="toc-page-no" data-target="{sec_id}">-</span>'
                "</div>"
            )
        )
        content_sections.append(
            (
                f'<section class="doc-section" id="{sec_id}">'
                f"<h1>{safe_title}</h1>"
                f"{body_html}"
                "</section>"
            )
        )

    safe_repo = html_lib.escape(title)
    safe_subtitle = html_lib.escape(i18n["doc_subtitle"])
    safe_generated_on = html_lib.escape(i18n["generated_on"])
    safe_abstract_title = html_lib.escape(i18n["abstract_title"])
    safe_abstract_body = html_lib.escape(i18n["abstract_body"])
    safe_summary_title = html_lib.escape(i18n["summary_title"])
    safe_version = html_lib.escape(i18n["version"])
    safe_changed_by = html_lib.escape(i18n["changed_by"])
    safe_changed_by_value = html_lib.escape(i18n["changed_by_value"])
    safe_change_date = html_lib.escape(i18n["change_date"])
    safe_page_label = html_lib.escape(i18n["page_label"])

    user_fonts_dir = Path.home() / "AppData" / "Local" / "Microsoft" / "Windows" / "Fonts"

    def _first_existing_font(candidates: list[Path]) -> Path | None:
        for candidate in candidates:
            try:
                if candidate.exists():
                    return candidate
            except Exception:
                continue
        return None

    def _font_data_uri(path: Path | None) -> str:
        if not path:
            return ""
        try:
            payload = base64.b64encode(path.read_bytes()).decode("ascii")
        except Exception:
            return ""
        return f"data:font/ttf;base64,{payload}"

    calibri_regular = _first_existing_font(
        [Path(r"C:\Windows\Fonts\calibri.ttf"), user_fonts_dir / "calibri.ttf"]
    )
    calibri_bold = _first_existing_font(
        [Path(r"C:\Windows\Fonts\calibrib.ttf"), user_fonts_dir / "calibrib.ttf"]
    )
    calibri_italic = _first_existing_font(
        [Path(r"C:\Windows\Fonts\calibrii.ttf"), user_fonts_dir / "calibrii.ttf"]
    )
    calibri_bold_italic = _first_existing_font(
        [Path(r"C:\Windows\Fonts\calibriz.ttf"), user_fonts_dir / "calibriz.ttf"]
    )

    embedded_font_css = ""
    body_font_stack = 'Calibri, "Segoe UI", Arial, sans-serif'
    regular_data = _font_data_uri(calibri_regular)
    if regular_data:
        embedded_font_css = (
            "@font-face {"
            "font-family: 'HubiaCalibri'; font-style: normal; font-weight: 400; "
            f"src: url('{regular_data}') format('truetype');"
            "}\n"
        )
        bold_data = _font_data_uri(calibri_bold)
        if bold_data:
            embedded_font_css += (
                "@font-face {"
                "font-family: 'HubiaCalibri'; font-style: normal; font-weight: 700; "
                f"src: url('{bold_data}') format('truetype');"
                "}\n"
            )
        italic_data = _font_data_uri(calibri_italic)
        if italic_data:
            embedded_font_css += (
                "@font-face {"
                "font-family: 'HubiaCalibri'; font-style: italic; font-weight: 400; "
                f"src: url('{italic_data}') format('truetype');"
                "}\n"
            )
        bold_italic_data = _font_data_uri(calibri_bold_italic)
        if bold_italic_data:
            embedded_font_css += (
                "@font-face {"
                "font-family: 'HubiaCalibri'; font-style: italic; font-weight: 700; "
                f"src: url('{bold_italic_data}') format('truetype');"
                "}\n"
            )
        body_font_stack = '"HubiaCalibri", Calibri, "Segoe UI", Arial, sans-serif'
        PDF_RENDER_FONT = "HubiaCalibri"
    else:
        PDF_RENDER_FONT = "Calibri/Segoe UI/Arial"

    html_doc = f"""<!doctype html>
<html lang="{lang}">
<head>
  <meta charset="utf-8" />
  <style>
    {embedded_font_css}
    @page {{
      size: A4;
      margin: 18mm;
    }}
    * {{ box-sizing: border-box; }}
    html, body {{
      margin: 0;
      padding: 0;
      color: #111827;
      font-family: {body_font_stack};
      font-size: 10pt;
      line-height: 1.45;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
      background: #fff;
    }}
    .fixed-page {{
      min-height: calc(297mm - 36mm);
      display: flex;
      flex-direction: column;
      page-break-after: always;
      break-after: page;
    }}
    .cover {{
      justify-content: center;
      align-items: center;
      text-align: center;
    }}
    .cover h1 {{
      margin: 0;
      font-size: 26pt;
      line-height: 1.2;
      font-weight: 700;
    }}
    .cover .subtitle {{
      margin-top: 10pt;
      font-size: 18pt;
      line-height: 1.2;
      font-weight: 400;
    }}
    .cover .generated {{
      margin-top: auto;
      width: 100%;
      text-align: right;
      font-size: 10pt;
      font-style: italic;
      color: #6b7280;
    }}
    .back-cover h2,
    .toc-page h2 {{
      margin: 0 0 12pt 0;
      font-size: 16pt;
      line-height: 1.25;
      font-weight: 700;
    }}
    .back-cover p {{
      margin: 0 0 6pt 0;
      text-align: justify;
    }}
    .meta-table {{
      margin-top: 8pt;
      width: 100%;
      border-collapse: collapse;
      font-size: 10pt;
    }}
    .meta-table th,
    .meta-table td {{
      border: 1px solid #cbd5e1;
      padding: 6px 8px;
      text-align: left;
      vertical-align: top;
    }}
    .meta-table th {{
      background: #f3f4f6;
      font-weight: 700;
    }}
    .toc-list {{
      margin: 0;
      padding: 0;
    }}
    .toc-row {{
      display: grid;
      grid-template-columns: auto 1fr 18mm;
      align-items: end;
      column-gap: 2mm;
      margin: 0 0 2pt 0;
      font-size: 10pt;
      line-height: 1.35;
    }}
    .toc-label {{
      white-space: nowrap;
      overflow: hidden;
      text-overflow: ellipsis;
    }}
    .toc-dots {{
      border-bottom: 1px dotted #6b7280;
      transform: translateY(-2px);
    }}
    .toc-page-no {{
      text-align: right;
      font-variant-numeric: tabular-nums;
    }}
    .doc-section {{
      page-break-before: always;
      break-before: page;
    }}
    .doc-section h1 {{
      margin: 0 0 12pt 0;
      font-size: 16pt;
      line-height: 1.25;
      font-weight: 700;
    }}
    .doc-section h2,
    .doc-section h3,
    .doc-section h4,
    .doc-section h5,
    .doc-section h6 {{
      margin: 0 0 10pt 0;
      font-weight: 700;
    }}
    .doc-section h2 {{ font-size: 12.5pt; }}
    .doc-section h3 {{ font-size: 11pt; }}
    .doc-section h4 {{ font-size: 10.4pt; }}
    .doc-section h5,
    .doc-section h6 {{ font-size: 10pt; }}
    .doc-section p {{
      margin: 0 0 1.8pt 0;
      text-align: justify;
    }}
    .doc-section ul,
    .doc-section ol {{
      margin: 0 0 4pt 18pt;
      padding: 0;
    }}
    .doc-section li {{
      margin: 0 0 1pt 0;
      text-align: justify;
    }}
    .doc-section code {{
      font-family: Consolas, "Courier New", monospace;
      background: #f3f4f6;
      padding: 1px 3px;
      border-radius: 3px;
      word-break: break-word;
    }}
    .doc-section pre {{
      margin: 0 0 6pt 0;
      padding: 8px 10px;
      border: 1px solid #d1d5db;
      border-radius: 6px;
      background: #f9fafb;
      overflow-wrap: anywhere;
      white-space: pre-wrap;
      word-break: break-word;
    }}
    .doc-section pre code {{
      background: transparent;
      padding: 0;
      border-radius: 0;
    }}
    .doc-section blockquote {{
      margin: 0 0 6pt 0;
      padding: 6px 10px;
      border-left: 3px solid #9ca3af;
      color: #374151;
      background: #f9fafb;
    }}
    .doc-section table {{
      margin: 0 0 6pt 0;
      width: 100%;
      border-collapse: collapse;
      font-size: 10pt;
    }}
    .doc-section th,
    .doc-section td {{
      border: 1px solid #d1d5db;
      padding: 5px 7px;
      vertical-align: top;
    }}
    .doc-section th {{
      background: #f3f4f6;
      font-weight: 700;
    }}
    #page-footers {{
      position: absolute;
      top: 0;
      left: 0;
      width: 100%;
      pointer-events: none;
    }}
    .page-footer-number {{
      position: absolute;
      right: 18mm;
      width: 40mm;
      text-align: right;
      font-size: 9pt;
      color: #6b7280;
      font-family: Calibri, "Segoe UI", Arial, sans-serif;
      font-variant-numeric: tabular-nums;
    }}
  </style>
</head>
<body>
  <section class="fixed-page cover">
    <h1>{safe_repo}</h1>
    <div class="subtitle">{safe_subtitle}</div>
    <div class="generated">{safe_generated_on} {generated_date}</div>
  </section>

  <section class="fixed-page back-cover">
    <h2>{safe_abstract_title}</h2>
    <p>{safe_abstract_body}</p>
    <table class="meta-table">
      <thead>
        <tr>
          <th>{safe_version}</th>
          <th>{safe_changed_by}</th>
          <th>{safe_change_date}</th>
        </tr>
      </thead>
      <tbody>
        <tr>
          <td>1.0</td>
          <td>{safe_changed_by_value}</td>
          <td>{generated_date}</td>
        </tr>
      </tbody>
    </table>
  </section>

  <section class="fixed-page toc-page">
    <h2>{safe_summary_title}</h2>
    <div class="toc-list">
      {''.join(toc_rows)}
    </div>
  </section>

  {''.join(content_sections)}

  <div id="page-footers"></div>

  <script>
    window.__preparePdfLayout = () => {{
      const mmToPx = 96 / 25.4;
      const printablePageHeight = (297 - 36) * mmToPx;
      const firstSection = document.querySelector(".doc-section");

      if (firstSection) {{
        const contentStartTop = firstSection.offsetTop;
        document.querySelectorAll(".toc-page-no").forEach((node) => {{
          const targetId = node.getAttribute("data-target");
          const target = targetId ? document.getElementById(targetId) : null;
          if (!target) {{
            node.textContent = "-";
            return;
          }}
          const relativeTop = Math.max(0, target.offsetTop - contentStartTop);
          const pageIndex = Math.floor(relativeTop / printablePageHeight);
          node.textContent = String(4 + pageIndex);
        }});
      }}

      const footerHost = document.getElementById("page-footers");
      if (footerHost) {{
        footerHost.innerHTML = "";
        const totalPages = Math.max(1, Math.ceil(document.documentElement.scrollHeight / printablePageHeight));
        for (let pageNo = 4; pageNo <= totalPages; pageNo += 1) {{
          const footer = document.createElement("div");
          footer.className = "page-footer-number";
          footer.style.top = `${{(pageNo * printablePageHeight) - 24}}px`;
          footer.textContent = {json_lib.dumps(f"{safe_page_label} ")} + pageNo;
          footerHost.appendChild(footer);
        }}
      }}
    }};
    window.__preparePdfLayout();
  </script>
</body>
</html>
"""

    try:
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch()
            page = browser.new_page(viewport={"width": 1400, "height": 1900})
            page.set_content(html_doc, wait_until="networkidle")
            page.evaluate("window.__preparePdfLayout && window.__preparePdfLayout()")
            pdf_bytes = page.pdf(
                format="A4",
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0mm", "bottom": "0mm", "left": "0mm", "right": "0mm"},
            )
            browser.close()
            return pdf_bytes
    except Exception as e:
        if isinstance(e, NotImplementedError):
            raise RuntimeError(
                "Playwright HTML backend is not supported by the current event-loop runtime."
            ) from e
        message = str(e)
        if "Executable doesn't exist" in message or "playwright install" in message:
            raise RuntimeError(
                "Playwright Chromium is not installed. Run: python -m playwright install chromium"
            ) from e
        raise


def _build_pdf_bytes(
    md_text: str,
    title: str,
    language: str = "EN-US",
    doc_kind: str = "technical",
) -> bytes:
    global PDF_RENDER_BACKEND, PDF_RENDER_FALLBACK_REASON, PDF_RENDER_FONT
    try:
        data = _build_pdf_bytes_html(md_text, title, language=language, doc_kind=doc_kind)
        PDF_RENDER_BACKEND = "html"
        PDF_RENDER_FALLBACK_REASON = ""
        return data
    except Exception as e:
        PDF_RENDER_BACKEND = "fpdf"
        PDF_RENDER_FALLBACK_REASON = str(e)
        if not PDF_RENDER_FONT:
            PDF_RENDER_FONT = ""
        return _build_pdf_bytes_fpdf(md_text, title, language=language, doc_kind=doc_kind)


def _build_docx_bytes(
    md_text: str,
    title: str,
    language: str = "EN-US",
    doc_kind: str = "technical",
) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except Exception as e:
        raise RuntimeError(
            "python-docx is required for Word export. Install with: pip install python-docx"
        ) from e

    lang = _normalize_doc_language(language)
    date_format = "%Y-%m-%d" if lang == "EN-US" else "%d/%m/%Y"
    generated_date = datetime.now().strftime(date_format)
    i18n = _pdf_i18n(lang, title, generated_date, doc_kind=doc_kind)
    blocks = _normalize_pdf_blocks(_markdown_to_docx_blocks(md_text))

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.line_spacing = 1.15

    heading_sizes = {1: 16, 2: 12, 3: 11, 4: 10.5, 5: 10, 6: 9.5}

    # Cover
    usable_height_pt = 650.0
    try:
        page_height_pt = float(section.page_height.pt)
        top_margin_pt = float(section.top_margin.pt)
        bottom_margin_pt = float(section.bottom_margin.pt)
        usable_height_pt = max(0.0, page_height_pt - top_margin_pt - bottom_margin_pt)
        # Keep cover content visually centered on most page sizes.
        cover_block_estimate_pt = 130.0
        cover_top_padding_pt = max(72.0, (usable_height_pt - cover_block_estimate_pt) * 0.42)
    except Exception:
        cover_top_padding_pt = 180.0
    # Push generated date close to bottom-right on cover.
    cover_date_space_before_pt = max(120.0, usable_height_pt - cover_top_padding_pt - 190.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(cover_top_padding_pt)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(i18n["doc_subtitle"])
    r.font.name = "Calibri"
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(cover_date_space_before_pt)
    r = p.add_run(f"{i18n['generated_on']} {generated_date}")
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)

    doc.add_page_break()

    # Back cover
    p = doc.add_paragraph()
    r = p.add_run(i18n["abstract_title"])
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph(i18n["abstract_body"])
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = [i18n["version"], i18n["changed_by"], i18n["change_date"]]
    values = ["1.0", i18n["changed_by_value"], generated_date]
    for idx, header in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(header)
        run.bold = True
    for idx, value in enumerate(values):
        table.rows[1].cells[idx].text = value

    doc.add_page_break()

    # Summary
    p = doc.add_paragraph()
    r = p.add_run(i18n["summary_title"])
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(8)

    for sec_title, _ in _split_markdown_sections(md_text):
        row = doc.add_paragraph(style="List Bullet")
        row.add_run(sec_title)

    doc.add_page_break()

    # Content
    has_content = False
    for kind, text in blocks:
        if kind == "blank":
            doc.add_paragraph("")
            continue
        if kind == "hr":
            doc.add_paragraph("—" * 48)
            continue
        if kind.startswith("h"):
            level = int(kind[1:])
            if level == 1 and has_content:
                doc.add_page_break()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(heading_sizes.get(level, 12))
            p.paragraph_format.space_after = Pt(10)
            has_content = True
            continue
        if kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            has_content = True
            continue
        if kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(text)
            r.italic = True
            has_content = True
            continue
        if kind == "li":
            indent = 0
            marker = "-"
            item_text = text
            if "\x1f" in text:
                parts = text.split("\x1f", 2)
                if len(parts) == 3:
                    indent_raw, marker_raw, item_text = parts
                    try:
                        indent = max(0, int(indent_raw))
                    except Exception:
                        indent = 0
                    marker = marker_raw.strip() or "-"
            level = min(indent // 2, 6)
            is_ordered = marker.endswith(".") and marker[:-1].isdigit()
            style_name = "List Number" if is_ordered else "List Bullet"
            try:
                p = doc.add_paragraph(style=style_name)
            except Exception:
                p = doc.add_paragraph()
                p.add_run(f"{marker if is_ordered else '•'} ")
            p.paragraph_format.left_indent = Inches(0.2 * level)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(item_text)
            has_content = True
            continue
        p = doc.add_paragraph(text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        has_content = True

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@st.cache_data(show_spinner=False)
def _get_pdf_bytes_from_file(
    md_path: Path,
    title: str,
    mtime: float,
    language: str,
    doc_kind: str = "technical",
    render_cache_version: str = "",
) -> bytes:
    _ = render_cache_version
    md_text = md_path.read_text(encoding="utf-8")
    return _build_pdf_bytes(md_text, title, language=language, doc_kind=doc_kind)


@st.cache_data(show_spinner=False)
def _get_docx_bytes_from_file(
    md_path: Path,
    title: str,
    mtime: float,
    language: str,
    doc_kind: str = "technical",
    render_cache_version: str = "",
) -> bytes:
    _ = render_cache_version
    md_text = md_path.read_text(encoding="utf-8")
    return _build_docx_bytes(md_text, title, language=language, doc_kind=doc_kind)


def _answer_with_llm(
    question,
    contexts,
    documentation_context="",
    language="PT-BR",
    provider="gemini",
    model_name="gemini-2.5-flash",
    api_key=None,
    use_system_key=True,
    temperature=0,
):
    model = init_llm(
        provider=provider,
        model_name=model_name,
        api_key=api_key,
        use_system_key=use_system_key,
        temperature=temperature,
    )
    parser = StrOutputParser()
    code_ctx_text = ""
    for c in contexts:
        code_ctx_text += (
            f"PATH: {c.get('path')}\n"
            f"TYPE: {c.get('tipo')}\n"
            f"NAME: {c.get('nome')}\n"
            f"CODE:\n{c.get('code','')}\n\n"
        )
    documentation_context = (documentation_context or "").strip()
    if not code_ctx_text.strip():
        code_ctx_text = "(no code.json matches found)\n"
    if not documentation_context:
        documentation_context = "(documentation context not available)\n"
    ctx_text = (
        "CODE.JSON CONTEXT (HIGH PRIORITY):\n"
        f"{code_ctx_text}\n"
        "DOCUMENTATION.MD CONTEXT (SECONDARY):\n"
        f"{documentation_context}\n"
    )
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                f"Responda em {language}. Seja direto e técnico. "
                "Use primeiro o contexto de CODE.JSON (prioridade alta). "
                "Use DOCUMENTATION.MD apenas para complementar. "
                "Se houver conflito, siga o CODE.JSON."
            ),
            ("user", "PERGUNTA:\n{question}\n\nCONTEXTO:\n{context}"),
        ]
    )
    chain = prompt | model | parser
    return chain.invoke({"question": question, "context": ctx_text})


def _rewrite_standard_section_titles(documentation_md: str, language: str) -> str:
    import re

    text = str(documentation_md or "")
    if not text.strip():
        return text
    title_map = _localized_standard_section_titles(language)
    marker_re = re.compile(r"^\s*<!--\s*SECTION:([A-Z_]+)\s*-->\s*$")
    lines = text.splitlines()

    for idx, line in enumerate(lines):
        match = marker_re.match(line.strip())
        if not match:
            continue
        section_key = match.group(1).strip().upper()
        target_title = title_map.get(section_key)
        if not target_title:
            continue
        j = idx + 1
        while j < len(lines) and not lines[j].strip():
            j += 1
        if j < len(lines) and lines[j].lstrip().startswith("# "):
            indent = lines[j][: len(lines[j]) - len(lines[j].lstrip())]
            current_heading = lines[j].lstrip()[2:].strip()
            prefix_match = re.match(r"^(\d+(?:\.\d+)*\.?)\s+(.+)$", current_heading)
            number_prefix = f"{prefix_match.group(1)} " if prefix_match else ""
            lines[j] = f"{indent}# {number_prefix}{target_title}"
    return "\n".join(lines)


def _derive_sections_from_marked_documentation(
    documentation_md: str,
    fallback_sections: dict | None = None,
) -> dict[str, dict]:
    text = str(documentation_md or "")
    fallback = fallback_sections if isinstance(fallback_sections, dict) else {}
    sections: dict[str, dict] = {}

    parts = text.split("<!-- SECTION:")
    for part in parts[1:]:
        header, *body = part.split("-->", 1)
        key = _normalize_doc_section_key(header.strip())
        if not key:
            continue

        body_text = body[0] if body else ""
        title = ""
        for raw_line in body_text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("# "):
                title = line[2:].strip()
            break

        fallback_meta = fallback.get(key, {}) if isinstance(fallback.get(key), dict) else {}
        if not title:
            title = str(fallback_meta.get("title") or key.replace("_", " ").title())
        description = str(fallback_meta.get("description") or "")
        sections[key] = {"title": title, "description": description}

    return sections if sections else dict(fallback)


def _translate_documentation_with_llm(
    documentation_md: str,
    target_language: str = "PT-BR",
    provider: str = "gemini",
    model_name: str = "gemini-2.5-flash",
    api_key: str | None = None,
    use_system_key: bool = True,
    temperature: float = 0.0,
) -> str:
    def _invoke_with_timeout(fn, timeout_seconds: float, timeout_message: str):
        import threading

        result_holder: dict[str, str] = {}
        error_holder: dict[str, Exception] = {}

        def _runner():
            try:
                result_holder["value"] = fn()
            except Exception as exc:  # noqa: BLE001
                error_holder["error"] = exc

        worker = threading.Thread(target=_runner, daemon=True)
        worker.start()
        worker.join(timeout=max(1.0, float(timeout_seconds)))
        if worker.is_alive():
            raise TimeoutError(timeout_message)
        if "error" in error_holder:
            raise error_holder["error"]
        return result_holder.get("value", "")

    lang = _normalize_doc_language(target_language)
    lang_names = {
        "PT-BR": "Português (Brasil)",
        "EN-US": "English (US)",
        "ES-ES": "Español (España)",
        "FR-FR": "Français (France)",
        "DE-DE": "Deutsch (Deutschland)",
    }
    strict_rule = {
        "PT-BR": "Escreva EXCLUSIVAMENTE em Português (Brasil). Nunca use outro idioma.",
        "EN-US": "Write EXCLUSIVELY in English (US). Never switch to another language.",
        "ES-ES": "Escribe EXCLUSIVAMENTE en Español (España). Nunca cambies a otro idioma.",
        "FR-FR": "Rédige EXCLUSIVEMENT en Français (France). N'utilise jamais une autre langue.",
        "DE-DE": "Schreibe AUSSCHLIESSLICH auf Deutsch (Deutschland). Nutze keine andere Sprache.",
    }[lang]
    standard_titles = _localized_standard_section_titles(lang)
    standard_title_rules = "\n".join(
        f"- {key}: # {title}" for key, title in standard_titles.items() if key in DEFAULT_DOCUMENTATION_SECTIONS
    )

    model = init_llm(
        provider=provider,
        model_name=model_name,
        api_key=api_key,
        use_system_key=use_system_key,
        temperature=temperature,
    )
    parser = StrOutputParser()
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                f"""
                You are a technical documentation translator.
                Target language: {lang_names[lang]}.
                {strict_rule}

                TRANSLATION RULES:
                - Translate the markdown documentation to the target language.
                - Preserve section delimiters exactly (e.g., <!-- SECTION:... -->).
                - For default section keys, enforce the translated section heading text exactly as:
{standard_title_rules}
                - For non-default/custom sections, keep the existing section heading text and meaning.
                - Preserve markdown structure and hierarchy.
                - Preserve code blocks, inline code, file paths, commands, and identifiers.
                - Remove temporary process markers and workflow artifacts (e.g., [PENDING], [pendente],
                  [EVIDENCIA NOVA], [NEW EVIDENCE], "NOVO TRECHO DO REPOSITORIO").
                - Do not add commentary outside the documentation.
                - Output only the final translated markdown.
                """,
            ),
            ("user", "DOCUMENTACAO:\n{documentation}"),
        ]
    )
    chain = prompt | model | parser
    try:
        translation_timeout_seconds = float(
            str(os.getenv("DOC_TRANSLATION_TIMEOUT_SECONDS") or "420").strip() or "420"
        )
    except Exception:
        translation_timeout_seconds = 420.0
    translated = _invoke_with_timeout(
        lambda: chain.invoke({"documentation": documentation_md}),
        timeout_seconds=translation_timeout_seconds,
        timeout_message=(
            f"Translation timed out after {int(translation_timeout_seconds)}s. "
            "Try again, switch model/provider, or reduce document size."
        ),
    )

    # Safety cleanup for known intermediate tags that may survive translation.
    if translated:
        import re

        marker_pattern = re.compile(
            r"\[(?:\s*pending\s*|\s*pendente\s*|\s*evid[eê]ncia\s+nova\s*|\s*new\s+evidence\s*)\]",
            flags=re.IGNORECASE,
        )
        translated = marker_pattern.sub("", translated)
        translated = re.sub(
            r"\bNOVO TRECHO DO REPOSITORIO\b:?\s*",
            "",
            translated,
            flags=re.IGNORECASE,
        )
        translated = re.sub(r"[ \t]+\n", "\n", translated)
        translated = re.sub(r"\n{3,}", "\n\n", translated)
        translated = _rewrite_standard_section_titles(translated, lang)
        translated = apply_standard_ai_disclaimer(translated, lang)
    return translated


def _normalize_path(path: str) -> str:
    return path.replace("\\", "/")


def _graph_groups(graph: dict) -> tuple[dict, dict]:
    nodes = graph.get("nodes", [])
    edges = graph.get("edges", [])
    group_nodes = {}
    for n in nodes:
        path = n.get("path", "") or ""
        path_norm = _normalize_path(path)
        if "/" in path_norm:
            group = path_norm.split("/", 1)[0]
        else:
            group = "root"
        group_nodes.setdefault(group, []).append(n)

    group_edges = {}
    for group, nlist in group_nodes.items():
        idset = {n.get("id") for n in nlist}
        group_edges[group] = [
            e for e in edges
            if e.get("from") in idset and e.get("to") in idset
        ]
    return group_nodes, group_edges


def _graph_height(node_count: int) -> int:
    base = 360
    if node_count <= 20:
        return 420
    if node_count <= 60:
        return 520
    if node_count <= 120:
        return 620
    return 720


def _render_cytoscape(nodes: list, edges: list, height: int = 520):
    elements = []
    for n in nodes:
        kind = n.get("kind", "misc")
        path = n.get("path", "")
        name = n.get("name", "")
        if kind == "file":
            label = path
        elif name:
            label = f"{name} ({kind})"
        else:
            label = n.get("id", "")
        elements.append({
            "data": {
                "id": n.get("id"),
                "label": label,
                "kind": kind,
            }
        })

    for e in edges:
        elements.append({
            "data": {
                "id": f"{e.get('from')}->{e.get('to')}",
                "source": e.get("from"),
                "target": e.get("to"),
                "kind": e.get("kind", "link"),
            }
        })

    container_id = f"cy-{uuid.uuid4().hex}"
    elements_json = json.dumps(elements)
    html = f"""
    <div id="{container_id}" style="width: 100%; height: {height}px;"></div>
    <script src="https://unpkg.com/cytoscape/dist/cytoscape.min.js"></script>
    <script src="https://unpkg.com/dagre/dist/dagre.min.js"></script>
    <script src="https://unpkg.com/cytoscape-dagre/cytoscape-dagre.js"></script>
    <script>
      (function() {{
        var cy = cytoscape({{
          container: document.getElementById("{container_id}"),
          elements: {elements_json},
          wheelSensitivity: 0.2,
          layout: {{
            name: "dagre",
            rankDir: "LR",
            nodeSep: 50,
            rankSep: 80
          }},
          style: [
            {{
              selector: "node",
              style: {{
                "label": "data(label)",
                "font-size": "11px",
                "text-wrap": "wrap",
                "text-max-width": 160,
                "text-valign": "center",
                "text-halign": "center",
                "color": "#e5e7eb",
                "text-outline-width": 1,
                "text-outline-color": "#111827",
                "background-color": "#374151",
                "shape": "round-rectangle",
                "padding": "6px",
                "width": "label",
                "height": "label"
              }}
            }},
            {{
              selector: "node[kind = 'file']",
              style: {{ "background-color": "#1f2937" }}
            }},
            {{
              selector: "node[kind = 'function']",
              style: {{ "background-color": "#7c3aed" }}
            }},
            {{
              selector: "node[kind = 'class']",
              style: {{ "background-color": "#0f766e" }}
            }},
            {{
              selector: "edge",
              style: {{
                "curve-style": "bezier",
                "target-arrow-shape": "triangle",
                "line-color": "#64748b",
                "target-arrow-color": "#64748b",
                "width": 1
              }}
            }},
            {{
              selector: ".faded",
              style: {{
                "opacity": 0.15,
                "text-opacity": 0.15
              }}
            }}
          ]
        }});

        cy.on("tap", "node", function(evt) {{
          var node = evt.target;
          cy.elements().addClass("faded");
          node.closedNeighborhood().removeClass("faded");
        }});
        cy.on("tap", function(evt) {{
          if (evt.target === cy) {{
            cy.elements().removeClass("faded");
          }}
        }});
      }})();
    </script>
    """
    components.html(html, height=height + 10)


def _ensure_mkdocs_port(force_restart: bool = False) -> tuple[int | None, str | None]:
    def _is_local_port_open(port: int | None) -> bool:
        try:
            if port is None:
                return False
            port = int(port)
            if port <= 0:
                return False
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
                sock.settimeout(0.25)
                return sock.connect_ex(("127.0.0.1", port)) == 0
        except Exception:
            return False

    cached_port = st.session_state.get("mkdocs_port")
    if not force_restart and _is_local_port_open(cached_port):
        return int(cached_port), None
    port, err = serve_mkdocs(Path.cwd(), port=cached_port, force_restart=force_restart)
    if err:
        return None, err
    st.session_state["mkdocs_port"] = port
    return port, None


def _first_nav_doc_path(nav_node) -> str | None:
    if isinstance(nav_node, str):
        value = nav_node.strip()
        return value or None
    if isinstance(nav_node, list):
        for item in nav_node:
            found = _first_nav_doc_path(item)
            if found:
                return found
        return None
    if isinstance(nav_node, dict):
        for value in nav_node.values():
            found = _first_nav_doc_path(value)
            if found:
                return found
        return None
    return None


def _resolve_docs_entry_html(docs_dir: Path) -> str:
    mkdocs_config_path = Path.cwd() / "mkdocs.yml"
    try:
        if mkdocs_config_path.exists():
            mkdocs_config = yaml.safe_load(mkdocs_config_path.read_text(encoding="utf-8")) or {}
            nav_first_doc = _first_nav_doc_path(mkdocs_config.get("nav"))
            if nav_first_doc:
                nav_doc_path = Path(str(nav_first_doc).strip())
                if nav_doc_path.suffix.lower() == ".md" and nav_doc_path.name.lower() != "chat.md":
                    if (docs_dir / nav_doc_path.name).exists():
                        return f"{nav_doc_path.stem}.html"
    except Exception:
        pass

    if (docs_dir / "index.md").exists():
        return "index.html"

    try:
        md_candidates = sorted(
            [p for p in docs_dir.glob("*.md") if p.name.lower() != "chat.md"],
            key=lambda p: p.name.lower(),
        )
        if md_candidates:
            return f"{md_candidates[0].stem}.html"
    except Exception:
        pass
    return "index.html"


def _load_graph_payload(graph_path: str | None) -> tuple[str, dict | None]:
    if graph_path:
        graph_file = Path(graph_path)
    else:
        graph_file = Path("")
    if not graph_path or not graph_file.exists():
        return graph_path or "", None
    try:
        return str(graph_file), json.loads(graph_file.read_text(encoding="utf-8"))
    except Exception:
        return str(graph_file), None


def _open_repo_from_library(
    entry_or_repo_key: str,
    preferred_language: str | None = None,
    doc_variant: str = "technical",
    require_docs: bool = False,
    start_mkdocs: bool = True,
) -> tuple[dict | None, str | None]:
    library = _load_repo_library()
    direct_entry = library.get(entry_or_repo_key)
    if isinstance(direct_entry, dict):
        repo_name = (
            direct_entry.get("repo_name")
            or _parse_library_entry_key(entry_or_repo_key)[0]
            or ""
        ).strip()
    else:
        repo_name = (entry_or_repo_key or "").strip()
    if not repo_name:
        return None, "Invalid repository entry in library."

    doc_variant_norm = str(doc_variant or "technical").strip().lower()
    if doc_variant_norm not in {"technical", "functional"}:
        doc_variant_norm = "technical"

    entry_key, chosen_entry = _pick_library_repo_variant(
        repo_name=repo_name,
        preferred_language=preferred_language,
        require_docs=bool(require_docs),
        docs_kind=doc_variant_norm,
        library=library,
    )
    if not entry_key or not isinstance(chosen_entry, dict):
        return None, "Repository not found in library."

    resolved_entry = _resolve_library_entry_assets(repo_name, chosen_entry)
    if resolved_entry != chosen_entry:
        _upsert_repo_library_entry(
            repo_name,
            {**resolved_entry, "entry_key": entry_key},
            language=resolved_entry.get("language"),
        )
        chosen_entry = resolved_entry

    code_json = Path(chosen_entry.get("library_code_json") or "")
    if not code_json.exists():
        code_json = _resolve_repo_code_json(
            repo_name,
            preferred_language=chosen_entry.get("language"),
        )
    if code_json is None:
        return None, f"code.json not found for {repo_name}."

    docs_available = _activate_repo_assets(chosen_entry, doc_variant=doc_variant_norm)
    if not docs_available:
        _reset_docs_workspace()
    port = st.session_state.get("mkdocs_port")
    if start_mkdocs:
        generate_mkdocs_config(
            project_root=Path.cwd(),
            repo_name=repo_name,
            repo_url=chosen_entry.get("repo_url"),
            author=chosen_entry.get("owner"),
        )
        port, err = _ensure_mkdocs_port(force_restart=True)
        if err:
            return None, err

    graph_path = chosen_entry.get("library_graph_json") or chosen_entry.get("graph_path")
    graph_path, graph_data = _load_graph_payload(graph_path)

    state = {
        "repo_path": chosen_entry.get("repo_path") or str(Path.cwd() / "in" / repo_name),
        "repo_name": repo_name,
        "owner": chosen_entry.get("owner"),
        "output_dir": str(code_json.parent),
        "language": chosen_entry.get("language") or st.session_state.get("language") or "PT-BR",
        "result": {},
        "graph_path": graph_path,
        "graph_data": graph_data,
        "docs_generated": bool(docs_available),
        "docs_skipped": not bool(docs_available),
        "technical_docs_generated": bool(chosen_entry.get("docs_available")),
        "functional_docs_generated": bool(chosen_entry.get("functional_docs_available")),
        "doc_variant": doc_variant_norm,
        "library_entry_key": entry_key,
        "library_entry": chosen_entry,
        "technical_library_entry": chosen_entry if bool(chosen_entry.get("docs_available")) else None,
        "functional_library_entry": chosen_entry if bool(chosen_entry.get("functional_docs_available")) else None,
        "mkdocs_port": port,
        "docs_nonce": uuid.uuid4().hex,
        "generation_mode": str(st.session_state.get("generation_mode") or "Technical only"),
    }

    st.session_state["repo_name"] = repo_name
    st.session_state["repo_language"] = state.get("language")
    if state.get("language"):
        st.session_state["language"] = state["language"]
    _save_last_repo(repo_name)
    return state, None


def _open_repo_from_functional_library(
    entry_or_repo_key: str,
    preferred_language: str | None = None,
    require_docs: bool = False,
    start_mkdocs: bool = True,
) -> tuple[dict | None, str | None]:
    library = _load_functional_repo_library()
    direct_entry = library.get(entry_or_repo_key)
    if isinstance(direct_entry, dict):
        repo_name = (
            direct_entry.get("repo_name")
            or _parse_library_entry_key(entry_or_repo_key)[0]
            or ""
        ).strip()
    else:
        repo_name = (entry_or_repo_key or "").strip()
    if not repo_name:
        return None, "Invalid repository entry in functional library."

    entry_key, chosen_entry = _pick_functional_library_repo_variant(
        repo_name=repo_name,
        preferred_language=preferred_language,
        require_docs=bool(require_docs),
        library=library,
    )
    if not entry_key or not isinstance(chosen_entry, dict):
        return None, "Repository not found in functional library."

    resolved_entry = _resolve_functional_library_entry_assets(repo_name, chosen_entry)
    if resolved_entry != chosen_entry:
        _upsert_functional_repo_library_entry(
            repo_name,
            {**resolved_entry, "entry_key": entry_key},
            language=resolved_entry.get("language"),
        )
        chosen_entry = resolved_entry

    code_json = Path(chosen_entry.get("library_code_json") or "")
    code_json_exists = code_json.exists()
    if not code_json_exists:
        resolved_code = _resolve_repo_code_json(
            repo_name,
            preferred_language=chosen_entry.get("language"),
        )
        if resolved_code is not None:
            code_json = resolved_code
            code_json_exists = True

    docs_available = _activate_repo_assets(chosen_entry, doc_variant="functional")
    if not docs_available:
        _reset_docs_workspace()
    port = st.session_state.get("mkdocs_port")
    if start_mkdocs:
        generate_mkdocs_config(
            project_root=Path.cwd(),
            repo_name=repo_name,
            repo_url=chosen_entry.get("repo_url"),
            author=chosen_entry.get("owner"),
        )
        port, err = _ensure_mkdocs_port(force_restart=True)
        if err:
            return None, err

    graph_path = chosen_entry.get("library_graph_json") or chosen_entry.get("graph_path")
    graph_path, graph_data = _load_graph_payload(graph_path)

    if code_json_exists:
        output_dir = str(code_json.parent)
    else:
        output_dir = str(chosen_entry.get("output_dir") or (Path.cwd() / "out" / repo_name))

    state = {
        "repo_path": chosen_entry.get("repo_path") or str(Path.cwd() / "in" / repo_name),
        "repo_name": repo_name,
        "owner": chosen_entry.get("owner"),
        "output_dir": output_dir,
        "language": chosen_entry.get("language") or st.session_state.get("language") or "PT-BR",
        "result": {},
        "graph_path": graph_path,
        "graph_data": graph_data,
        "docs_generated": bool(docs_available),
        "docs_skipped": not bool(docs_available),
        "technical_docs_generated": False,
        "functional_docs_generated": bool(docs_available),
        "doc_variant": "functional",
        "library_entry_key": entry_key,
        "library_entry": chosen_entry,
        "technical_library_entry": None,
        "functional_library_entry": chosen_entry,
        "mkdocs_port": port,
        "docs_nonce": uuid.uuid4().hex,
        "generation_mode": str(st.session_state.get("generation_mode") or GEN_MODE_FUNCTIONAL_ONLY),
    }

    st.session_state["repo_name"] = repo_name
    st.session_state["repo_language"] = state.get("language")
    if state.get("language"):
        st.session_state["language"] = state["language"]
    _save_last_repo(repo_name)
    return state, None


def _activate_loaded_repo_doc_variant(state: dict, target_variant: str) -> tuple[bool, str]:
    target = str(target_variant or "technical").strip().lower()
    if target not in {"technical", "functional"}:
        target = "technical"

    repo_name = str(state.get("repo_name") or "").strip()
    language = _normalize_doc_language(
        state.get("language") or st.session_state.get("repo_language") or st.session_state.get("language")
    )

    if target == "functional":
        functional_entry = state.get("functional_library_entry")
        if not isinstance(functional_entry, dict):
            if repo_name:
                _, functional_entry = _pick_functional_library_repo_variant(
                    repo_name=repo_name,
                    preferred_language=language,
                    require_docs=True,
                )
        if isinstance(functional_entry, dict):
            if _activate_repo_assets(functional_entry, doc_variant="functional"):
                state["functional_library_entry"] = functional_entry
                return True, ""
        if _activate_local_functional_workspace():
            return True, ""
        return False, "Functional documentation is not available in workspace or library."

    technical_entry = state.get("technical_library_entry")
    if not isinstance(technical_entry, dict):
        if repo_name:
            _, technical_entry = _pick_library_repo_variant(
                repo_name=repo_name,
                preferred_language=language,
                require_docs=True,
                docs_kind="technical",
            )
    if isinstance(technical_entry, dict):
        if _activate_repo_assets(technical_entry, doc_variant="technical"):
            state["technical_library_entry"] = technical_entry
            if not isinstance(state.get("library_entry"), dict):
                state["library_entry"] = technical_entry
            return True, ""
    return False, "Technical documentation is not available in library for this repository/language."


def _render_loaded_repo(state: dict):
    repo_name = state.get("repo_name")
    if not repo_name:
        return
    repo_path = state.get("repo_path")
    owner = state.get("owner")
    result = state.get("result") or {}
    graph_data = state.get("graph_data")
    graph_path = state.get("graph_path")

    st.success("Repository loaded successfully!")
    if repo_path:
        st.code(repo_path)
    if owner is not None:
        st.code(f"Owner: {owner}")

    # Display interactive graphs
    st.subheader("📊 Repository Graphs")
    if graph_data is None and graph_path and os.path.exists(graph_path):
        try:
            with open(graph_path, "r", encoding="utf-8") as f:
                graph_data = json.load(f)
        except Exception:
            graph_data = None
    if graph_data:
        group_nodes, group_edges = _graph_groups(graph_data)

        with st.expander("General", expanded=False):
            nodes = graph_data.get("nodes", [])
            edges = graph_data.get("edges", [])
            _render_cytoscape(nodes, edges, height=_graph_height(len(nodes)))

        groups = sorted([g for g in group_nodes.keys() if g != "root"], key=lambda g: g.lower())
        if not groups:
            st.info("No subfolders to display.")
        for group in groups:
            with st.expander(group, expanded=False):
                nodes = group_nodes.get(group, [])
                edges = group_edges.get(group, [])
                if nodes:
                    _render_cytoscape(nodes, edges, height=_graph_height(len(nodes)))
                else:
                    st.info("No nodes to display.")
    else:
        st.warning("graph.json not found.")

    # Display additional information
    with st.expander("📋 Generated Files"):
        if result:
            st.json(result)
        else:
            st.caption("No generation data available.")

    active_doc_variant = str(state.get("doc_variant") or "technical").strip().lower()
    if active_doc_variant not in {"technical", "functional"}:
        active_doc_variant = "technical"
        state["doc_variant"] = "technical"

    functional_variant_available = bool(state.get("functional_docs_generated"))
    if not functional_variant_available:
        functional_variant_available = (
            (Path.cwd() / "functional_documentation.md").exists()
            and (Path.cwd() / "docs_functional").exists()
        )
    technical_variant_available = active_doc_variant == "technical"
    if not technical_variant_available:
        technical_entry = state.get("technical_library_entry")
        if not isinstance(technical_entry, dict) and repo_name:
            _, technical_entry = _pick_library_repo_variant(
                repo_name=repo_name,
                preferred_language=state.get("language"),
                require_docs=True,
                docs_kind="technical",
            )
        technical_variant_available = isinstance(technical_entry, dict)
        if technical_variant_available:
            state["technical_library_entry"] = technical_entry
    available_variants = []
    if technical_variant_available:
        available_variants.append("technical")
    if functional_variant_available:
        available_variants.append("functional")

    loaded_generation_mode = str(state.get("generation_mode") or "").strip()
    allow_view_switch = loaded_generation_mode == GEN_MODE_TECHNICAL_AND_FUNCTIONAL

    if (
        state.get("docs_generated")
        and active_doc_variant == "functional"
        and not allow_view_switch
    ):
        # Defensive resync after switching from another repo/view in the same session.
        _activate_loaded_repo_doc_variant(state, "functional")

    if allow_view_switch and len(available_variants) > 1:
        label_by_variant = {"technical": "Technical", "functional": "Functional"}
        default_variant = active_doc_variant if active_doc_variant in available_variants else available_variants[0]
        selected_label = st.radio(
            "Documentation view",
            [label_by_variant[v] for v in available_variants],
            index=available_variants.index(default_variant),
            horizontal=True,
            key=f"docs_view_variant_{repo_name}",
        )
        selected_variant = next(
            (variant for variant, label in label_by_variant.items() if label == selected_label),
            default_variant,
        )
        if selected_variant != active_doc_variant:
            switched, switch_err = _activate_loaded_repo_doc_variant(state, selected_variant)
            if not switched:
                st.warning(switch_err or f"Failed to activate {selected_variant} documentation view.")
            else:
                state["doc_variant"] = selected_variant
                state["docs_generated"] = True
                state["docs_skipped"] = False
                generate_mkdocs_config(
                    project_root=Path.cwd(),
                    repo_name=repo_name,
                    repo_url=(state.get("library_entry") or {}).get("repo_url"),
                    author=owner,
                )
                port, mkdocs_err = _ensure_mkdocs_port(force_restart=True)
                if mkdocs_err:
                    retry_port, retry_err = _ensure_mkdocs_port(force_restart=False)
                    if retry_err:
                        st.warning("Failed to refresh documentation server for selected view.")
                        st.code(mkdocs_err)
                    else:
                        port = retry_port
                else:
                    state["mkdocs_port"] = port
                    state["docs_nonce"] = uuid.uuid4().hex
                st.session_state["repo_state"] = state
                _safe_rerun()
                return

    if state.get("docs_generated"):
        if active_doc_variant == "functional":
            st.success("Functional markdown documentation loaded successfully!")
        else:
            st.success("Markdown documentation generated successfully!")
    elif state.get("docs_skipped"):
        st.info("⏭️ Skipped markdown documentation generation")

    port = state.get("mkdocs_port") or st.session_state.get("mkdocs_port")
    if not port and state.get("docs_generated"):
        port, err = _ensure_mkdocs_port(force_restart=False)
        if err:
            st.warning("Documentation server is not running.")
            st.code(err)
            port = None
        else:
            state["mkdocs_port"] = port
    if port:
        st.success("📖 Documentation server is running")
        docs_nonce = uuid.uuid4().hex
        state["docs_nonce"] = docs_nonce
        docs_dir = Path.cwd() / "docs"
        docs_entry_html = "index.html" if (docs_dir / "index.md").exists() else _resolve_docs_entry_html(docs_dir)
        docs_url = f"http://127.0.0.1:{port}/{docs_entry_html}?v={docs_nonce}&variant={active_doc_variant}"
        technical_md_path = Path.cwd() / "documentation.md"
        functional_md_path = Path.cwd() / "functional_documentation.md"
        if active_doc_variant == "functional":
            md_candidates = [functional_md_path, technical_md_path]
            export_suffix = "functional-documentation"
        else:
            md_candidates = [technical_md_path, functional_md_path]
            export_suffix = "documentation"
        docs_md_path = next((p for p in md_candidates if p.exists()), None)
        pdf_language = state.get("language") or st.session_state.get("language") or "PT-BR"
        col_docs, col_pdf, col_docx = st.columns([1, 1, 1])
        with col_docs:
            st.markdown(
                f"""
                <div class="open-docs-wrap odt-align-left">
                  <a class="open-docs-preview odt-link" href="{docs_url}" target="_blank" rel="noopener">
                    📖 Open docs
                  </a>
                </div>
                """,
                unsafe_allow_html=True,
            )
        with col_pdf:
            if docs_md_path is not None and docs_md_path.exists():
                try:
                    pdf_bytes = _get_pdf_bytes_from_file(
                        docs_md_path,
                        title=repo_name,
                        mtime=docs_md_path.stat().st_mtime,
                        language=pdf_language,
                        doc_kind=active_doc_variant,
                        render_cache_version=EXPORT_RENDER_CACHE_VERSION,
                    )
                    pdf_b64 = base64.b64encode(pdf_bytes).decode("ascii")
                    download_name = f"{repo_name}-{export_suffix}.pdf".replace('"', "")
                    st.markdown(
                        f"""
                        <div class="open-docs-wrap odt-align-center">
                          <a class="open-docs-preview odt-link"
                             href="data:application/pdf;base64,{pdf_b64}"
                             download="{download_name}">
                            ⬇️ Download PDF
                          </a>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                except Exception as e:
                    st.warning(f"PDF unavailable: {e}")
        with col_docx:
            if docs_md_path is not None and docs_md_path.exists():
                try:
                    docx_bytes = _get_docx_bytes_from_file(
                        docs_md_path,
                        title=repo_name,
                        mtime=docs_md_path.stat().st_mtime,
                        language=pdf_language,
                        doc_kind=active_doc_variant,
                        render_cache_version=EXPORT_RENDER_CACHE_VERSION,
                    )
                    docx_b64 = base64.b64encode(docx_bytes).decode("ascii")
                    docx_name = f"{repo_name}-{export_suffix}.docx".replace('"', "")
                    st.markdown(
                        f"""
                        <div class="open-docs-wrap odt-align-right">
                          <a class="open-docs-preview odt-link"
                             href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_b64}"
                             download="{docx_name}">
                            ⬇️ Download Word
                          </a>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                except Exception as e:
                    st.warning(f"Word unavailable: {e}")
            else:
                st.caption("Documentation file not found.")
        if allow_view_switch and len(available_variants) > 1:
            alt_variant_label = "Functional" if active_doc_variant == "technical" else "Technical"
            st.caption(f"Switch Documentation view to {alt_variant_label} to open/export that variant.")


def _render_chat(repo_name: str | None, language: str):
    _ensure_llm_state()
    st.markdown(
        """
        <style>
        :root {
            --chat-label-color: #9ca3af;
            --chat-select-bg: #111827;
            --chat-select-border: #374151;
            --chat-select-text: #e5e7eb;
            --chat-select-icon: #9ca3af;
            --chat-menu-bg: #0f172a;
            --chat-menu-border: #334155;
            --chat-menu-text: #e5e7eb;
            --chat-menu-hover: rgba(148, 163, 184, 0.18);
        }
        @media (prefers-color-scheme: light) {
            :root {
                --chat-label-color: #64748b;
                --chat-select-bg: #ffffff;
                --chat-select-border: #cbd5e1;
                --chat-select-text: #0f172a;
                --chat-select-icon: #475569;
                --chat-menu-bg: #ffffff;
                --chat-menu-border: #cbd5e1;
                --chat-menu-text: #0f172a;
                --chat-menu-hover: #e2e8f0;
            }
        }
        .block-container {max-width: 100%; padding-top: 1rem; padding-bottom: 1rem;}
        header[data-testid="stHeader"] {display: none;}
        footer {visibility: hidden;}
        .chat-title {
            font-size: 1.6rem;
            font-weight: 700;
            margin: 0.25rem 0 0.5rem 0;
            padding-left: 2.75rem;
        }
        .control-label {
            font-size: 0.8rem;
            color: var(--chat-label-color);
            font-weight: 600;
            text-align: left;
            margin-bottom: 0.25rem;
            max-width: 220px;
            margin-left: auto;
        }
        div[data-testid="stSelectbox"] > label {display: none;}
        div[data-testid="stSelectbox"] {
            max-width: 220px;
            margin-left: auto;
        }
        div[data-testid="stSelectbox"] div[data-baseweb="select"] > div {
            background: var(--chat-select-bg) !important;
            border: 1px solid var(--chat-select-border) !important;
            border-radius: 999px;
            min-height: 38px;
            padding-left: 8px;
            display: flex;
            align-items: center;
        }
        div[data-testid="stSelectbox"] div[data-baseweb="select"] > div > div:first-child,
        div[data-testid="stSelectbox"] div[data-baseweb="select"] > div > div:first-child * {
            color: var(--chat-select-text) !important;
            opacity: 1 !important;
            font-weight: 600;
            font-size: 0.85rem;
            line-height: 1.35;
        }
        div[data-testid="stSelectbox"] div[data-baseweb="select"] svg {
            color: var(--chat-select-icon) !important;
        }
        div[role="listbox"] {
            background: var(--chat-menu-bg) !important;
            border: 1px solid var(--chat-menu-border) !important;
        }
        div[role="option"] {
            color: var(--chat-menu-text) !important;
        }
        div[role="option"]:hover,
        div[role="option"][aria-selected="true"] {
            background: var(--chat-menu-hover) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    top_left, top_right = st.columns([3, 2])
    with top_left:
        st.markdown("<div class='chat-title'>Repository Chat</div>", unsafe_allow_html=True)
    with top_right:
        provider = st.session_state["llm_provider"]
        model_options = MODEL_OPTIONS.get(provider, MODEL_OPTIONS["gemini"])
        if st.session_state.get("llm_model") not in model_options:
            st.session_state["llm_model"] = model_options[0]
        st.markdown("<div class='control-label'>Model</div>", unsafe_allow_html=True)
        model_choice = st.selectbox(
            "Model",
            model_options,
            key="llm_model",
            label_visibility="collapsed",
        )
        temperature = 0.0
    use_system_key = st.session_state.get("llm_use_system_key", True)
    api_key = st.session_state.get("llm_api_key", "")
    bedrock_access_key = st.session_state.get("llm_bedrock_access_key", "")
    bedrock_secret_key = st.session_state.get("llm_bedrock_secret_key", "")
    if not use_system_key and not api_key:
        api_key = _get_runtime_key(provider) or ""
        st.session_state["llm_api_key"] = api_key
    if provider == "bedrock":
        _set_runtime_key(provider, None)
    else:
        _set_runtime_key(provider, None if use_system_key else api_key)
    _set_runtime_bedrock_credentials(
        provider,
        use_system_key,
        bedrock_access_key,
        bedrock_secret_key,
    )
    _save_llm_config(
        provider,
        model_choice,
        use_system_key,
        api_key=api_key,
        bedrock_access_key=bedrock_access_key,
        bedrock_secret_key=bedrock_secret_key,
    )
    if not repo_name:
        st.info("Load a repository first in the main app.")
        return

    code_json_path = _resolve_repo_code_json(repo_name)
    if code_json_path is None:
        st.error("code.json not found. Generate documentation first.")
        return

    index_key = f"code_index:{repo_name}"
    history_key = f"chat_history:{repo_name}"

    if index_key not in st.session_state:
        st.session_state[index_key] = _load_code_index(code_json_path)
    if history_key not in st.session_state:
        st.session_state[history_key] = _load_chat_history(repo_name)

    for msg in st.session_state[history_key]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    question = st.chat_input("Ask about the repository...")
    if question:
        st.session_state[history_key].append({"role": "user", "content": question})
        _save_chat_history(repo_name, st.session_state[history_key])
        with st.chat_message("user"):
            st.markdown(question)

        hits = _simple_search(st.session_state[index_key], question, k=5)
        doc_context = _extract_documentation_context(Path.cwd() / "documentation.md", question, k=5, max_chars=9000)
        try:
            answer = _answer_with_llm(
                question,
                hits,
                documentation_context=doc_context,
                language=language,
                provider=provider,
                model_name=model_choice,
                api_key=api_key,
                use_system_key=use_system_key,
                temperature=temperature,
            )
            st.session_state[history_key].append({"role": "assistant", "content": answer})
            _save_chat_history(repo_name, st.session_state[history_key])
            with st.chat_message("assistant"):
                st.markdown(answer)
        except Exception as e:
            st.error(f"Failed to query the model: {e}")


st.set_page_config(
    page_title="Repository Documentation",
    layout="centered"
)

_ensure_llm_state()
_sync_repo_library_from_disk()
_sync_functional_repo_library_from_disk()
_migrate_functional_assets_from_technical_library()

st.markdown(
    """
    <style>
      .block-container {
        padding-top: 3.0rem !important;
      }
      .open-docs-wrap { margin-top: 0.5rem; }
      .odt-align-left { text-align: left; }
      .odt-align-center { text-align: center; }
      .odt-align-right { text-align: right; }
      .open-docs-preview {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.6rem 0.95rem;
        border-radius: 999px;
        text-decoration: none !important;
        font-weight: 600;
        font-size: 0.9rem;
        transition: transform 0.15s ease, box-shadow 0.15s ease, filter 0.15s ease;
      }
      .odt-link {
        color: #7dd3fc;
        background: transparent;
        padding: 0.25rem 0;
        border-radius: 0;
        position: relative;
        font-size: 1rem;
        font-weight: 700;
      }
      .odt-link::after {
        content: "";
        position: absolute;
        left: 0;
        bottom: -3px;
        width: 100%;
        height: 2px;
        background: linear-gradient(90deg, #22d3ee, #60a5fa);
        transform: scaleX(0);
        transform-origin: left;
        transition: transform 0.15s ease;
      }
      .odt-link:hover::after { transform: scaleX(1); }
      @media (max-width: 640px) {
        .open-docs-preview { font-size: 0.95rem; }
      }
    </style>
    """,
    unsafe_allow_html=True,
)

view = _get_query_param("view")
if view == "chat":
    repo_param = _get_query_param("repo")
    repo_name = repo_param or st.session_state.get("repo_name") or _load_last_repo()
    language = _get_query_param("lang") or st.session_state.get("language") or "PT-BR"
    _render_chat(repo_name, language)
    st.stop()


_render_top_logo()
st.title("Repository Documentation")
st.write("Provide a repository source to start generating documentation.")

library_entries = _load_repo_library()
functional_library_entries = _load_functional_repo_library()
selected_library_repo_name = None
selected_library_entry: dict = {}
selected_library_variants: list[tuple[str, dict]] = []
selected_functional_library_repo_name = None
selected_functional_library_entry: dict = {}
selected_functional_library_variants: list[tuple[str, dict]] = []
selected_target_lang_entry: dict | None = None
selected_functional_target_lang_entry: dict | None = None
selected_functional_source_lang_entry: dict | None = None
translate_library_docs_on_load = False
translate_functional_library_docs_on_load = True
if _is_functional_only_mode(st.session_state.get("generation_mode")):
    st.session_state["generation_mode"] = GEN_MODE_FUNCTIONAL_ONLY
generation_mode = st.selectbox(
    "Documentation generation mode",
    [
        GEN_MODE_TECHNICAL_ONLY,
        GEN_MODE_TECHNICAL_AND_FUNCTIONAL,
        GEN_MODE_FUNCTIONAL_ONLY,
    ],
    index=0,
    help=(
        "Technical only: generate/load technical documentation. "
        "Technical + Functional: generate technical first, then functional. "
        "Functional only: generate/load functional documentation directly from repository artifacts."
    ),
    key="generation_mode",
    accept_new_options=False,
)
generation_mode = _normalize_generation_mode_value(generation_mode)

if "source_type" not in st.session_state:
    st.session_state["source_type"] = "Git URL"
if st.session_state.get("source_type") == SOURCE_TECH_LIBRARY_LEGACY:
    st.session_state["source_type"] = SOURCE_TECH_LIBRARY
if generation_mode == GEN_MODE_FUNCTIONAL_ONLY:
    source_options = ["Git URL", "ZIP file", "Local folder", SOURCE_FUNCTIONAL_LIBRARY]
    if st.session_state.get("source_type") not in source_options:
        st.session_state["source_type"] = "Git URL"
elif generation_mode == GEN_MODE_TECHNICAL_AND_FUNCTIONAL:
    source_options = ["Git URL", "ZIP file", "Local folder"]
    if st.session_state.get("source_type") not in source_options:
        st.session_state["source_type"] = "Git URL"
else:
    source_options = ["Git URL", "ZIP file", "Local folder", SOURCE_TECH_LIBRARY]
    if st.session_state.get("source_type") == SOURCE_FUNCTIONAL_LIBRARY:
        st.session_state["source_type"] = "Git URL"

source_type = st.radio(
    "Select repository source",
    source_options,
    key="source_type",
)

source = None

# -------- Input section --------
if source_type == "Git URL":
    source = st.text_input(
        "Git repository URL",
        placeholder="https://github.com/user/repo"
    )

elif source_type == "ZIP file":
    uploaded_zip = st.file_uploader(
        "Upload ZIP file",
        type=["zip"]
    )
    if uploaded_zip:
        temp_path = Path(uploaded_zip.name)
        temp_path.write_bytes(uploaded_zip.getvalue())
        source = str(temp_path)

elif source_type == "Local folder":
    source = st.text_input(
        "Local path",
        placeholder="C:/path/to/project"
    )

elif source_type == SOURCE_TECH_LIBRARY:
    if library_entries:
        grouped = _group_library_entries_by_repo(library_entries)
        repo_rows: list[tuple[str, str]] = []
        for repo_name, variants in grouped.items():
            latest = max((str(v.get("updated_at") or "") for _, v in variants), default="")
            repo_rows.append((repo_name, latest))
        repo_rows.sort(key=lambda row: row[1], reverse=True)
        labels = [row[0] for row in repo_rows]

        selected_library_repo_name = st.selectbox(
            "Saved repositories",
            labels,
            key="library_repo_select",
        )
        selected_library_variants = grouped.get(selected_library_repo_name, [])
        selected_library_variants.sort(
            key=lambda row: str(row[1].get("updated_at") or ""),
            reverse=True,
        )
        if selected_library_variants:
            selected_library_entry = selected_library_variants[0][1]

        available_langs = sorted(
            {
                _normalize_doc_language(entry.get("language"))
                for _, entry in selected_library_variants
            }
        )
        docs_langs = sorted(
            {
                _normalize_doc_language(entry.get("language"))
                for _, entry in selected_library_variants
                if entry.get("docs_available")
            }
        )
        languages_text = ", ".join(available_langs) if available_langs else "-"
        if docs_langs:
            docs_note = f"technical docs: {', '.join(docs_langs)}"
        else:
            docs_note = "no saved docs"
        formatted_updated_at = _format_library_updated_at(selected_library_entry.get("updated_at"))
        repo_label = selected_library_repo_name
        st.caption(
            f"{repo_label} • languages: {languages_text} • {docs_note} • "
            f"last update on {formatted_updated_at} (UTC-3)"
        )
    else:
        st.caption("No repositories saved in library.")

elif source_type == SOURCE_FUNCTIONAL_LIBRARY:
    if functional_library_entries:
        grouped = _group_functional_library_entries_by_repo(functional_library_entries)
        repo_rows: list[tuple[str, str]] = []
        for repo_name, variants in grouped.items():
            latest = max((str(v.get("updated_at") or "") for _, v in variants), default="")
            repo_rows.append((repo_name, latest))
        repo_rows.sort(key=lambda row: row[1], reverse=True)
        labels = [row[0] for row in repo_rows]

        selected_functional_library_repo_name = st.selectbox(
            "Saved functional repositories",
            labels,
            key="functional_library_repo_select",
        )
        selected_functional_library_variants = grouped.get(selected_functional_library_repo_name, [])
        selected_functional_library_variants.sort(
            key=lambda row: str(row[1].get("updated_at") or ""),
            reverse=True,
        )
        if selected_functional_library_variants:
            selected_functional_library_entry = selected_functional_library_variants[0][1]

        available_langs = sorted(
            {
                _normalize_doc_language(entry.get("language"))
                for _, entry in selected_functional_library_variants
            }
        )
        docs_langs = sorted(
            {
                _normalize_doc_language(entry.get("language"))
                for _, entry in selected_functional_library_variants
                if entry.get("functional_docs_available")
            }
        )
        languages_text = ", ".join(available_langs) if available_langs else "-"
        docs_note = f"functional docs: {', '.join(docs_langs)}" if docs_langs else "no saved docs"
        formatted_updated_at = _format_library_updated_at(
            selected_functional_library_entry.get("updated_at")
        )
        repo_label = selected_functional_library_repo_name
        st.caption(
            f"{repo_label} • languages: {languages_text} • {docs_note} • "
            f"last update on {formatted_updated_at} (UTC-3)"
        )
    else:
        st.caption("No repositories saved in functional library.")

# Language selection (format: PT-BR, EN-US, etc.)
language = st.selectbox(
    "Select documentation language",
    ["PT-BR", "EN-US", "ES-ES", "FR-FR", "DE-DE"],
    index=0,
    help="Select output language using format PT-BR, EN-US, etc."
)
st.session_state["language"] = language

if source_type == SOURCE_TECH_LIBRARY and selected_library_repo_name:
    target_lang = _normalize_doc_language(language)
    selected_target_lang_entry = next(
        (
            entry
            for _, entry in selected_library_variants
            if _normalize_doc_language(entry.get("language")) == target_lang
            and bool(entry.get("docs_available"))
        ),
        None,
    )
    _, source_entry = _pick_library_repo_variant(
        repo_name=selected_library_repo_name,
        preferred_language=selected_library_entry.get("language"),
        require_docs=True,
        docs_kind="technical",
        library=library_entries,
    )
    source_lang = _normalize_doc_language(source_entry.get("language")) if isinstance(source_entry, dict) else "-"

    if generation_mode == GEN_MODE_FUNCTIONAL_ONLY:
        if isinstance(source_entry, dict):
            st.caption(
                "Functional documentation will be generated from saved technical docs "
                f"({source_lang}) and stored in Functional Library as {target_lang}. "
                "Existing functional docs in that language will be replaced."
            )
        else:
            st.caption("Selected repository has no saved technical documentation to derive functional docs.")
    else:
        if isinstance(selected_target_lang_entry, dict):
            st.caption(f"Saved documentation already exists in {target_lang}.")
        elif isinstance(source_entry, dict):
            translate_library_docs_on_load = st.checkbox(
                "Translate saved documentation on load",
                value=True,
                key="translate_library_docs_on_load",
                help=(
                    f"Translate existing documentation from {source_lang} to {target_lang} in one LLM call, "
                    "without remapping the repository."
                ),
            )
        else:
            st.caption("Selected repository has no saved documentation to translate.")
elif source_type == SOURCE_FUNCTIONAL_LIBRARY and selected_functional_library_repo_name:
    target_lang = _normalize_doc_language(language)
    selected_functional_target_lang_entry = next(
        (
            entry
            for _, entry in selected_functional_library_variants
            if _normalize_doc_language(entry.get("language")) == target_lang
            and bool(entry.get("functional_docs_available"))
        ),
        None,
    )
    _, selected_functional_source_lang_entry = _pick_functional_library_repo_variant(
        repo_name=selected_functional_library_repo_name,
        preferred_language=selected_functional_library_entry.get("language"),
        require_docs=True,
        library=functional_library_entries,
    )
    if isinstance(selected_functional_target_lang_entry, dict):
        st.caption(f"Saved functional documentation is available in {target_lang}.")
    elif isinstance(selected_functional_source_lang_entry, dict):
        source_lang = _normalize_doc_language(selected_functional_source_lang_entry.get("language"))
        translate_functional_library_docs_on_load = st.checkbox(
            "Auto-translate functional docs on load",
            value=True,
            key="translate_functional_library_docs_on_load",
            help=(
                "If the selected language is missing, translate from an available saved functional "
                "documentation and save it in Functional Library."
            ),
        )
        if translate_functional_library_docs_on_load:
            st.caption(
                "No saved functional documentation in "
                f"{target_lang}. It will be translated from {source_lang} and saved in Functional Library."
            )
        else:
            st.caption(
                f"No saved functional documentation in {target_lang}. "
                "Disable this only if you want to load existing language variants only."
            )
    else:
        st.caption("Selected repository has no saved functional documentation to load or translate.")

st.subheader("LLM settings")
provider = st.radio(
    "Provider",
    options=list(PROVIDER_LABELS.keys()),
    format_func=lambda p: PROVIDER_LABELS[p],
    horizontal=True,
    key="llm_provider",
)
model_options = MODEL_OPTIONS.get(provider, MODEL_OPTIONS["gemini"])
if st.session_state.get("llm_model") not in model_options:
    st.session_state["llm_model"] = model_options[0]
model_name = st.selectbox(
    "Model",
    model_options,
    key="llm_model",
    help="Choose the model for the selected provider.",
)
use_system_key = st.checkbox(
    "Use system key",
    key="llm_use_system_key",
    help=(
        "Gemini and GPT can use an API key. Bedrock uses AWS credentials "
        "from the system (.aws/env/profile/role) or manual AWS keys. "
        "In Bedrock manual mode, no extra env export is required."
    ),
)
api_key = st.session_state.get("llm_api_key", "")
bedrock_access_key = st.session_state.get("llm_bedrock_access_key", "")
bedrock_secret_key = st.session_state.get("llm_bedrock_secret_key", "")
if not use_system_key and provider in {"gemini", "openai"}:
    api_key = st.text_input(
        "API key",
        type="password",
        key="llm_api_key",
        placeholder="Paste your key here",
    )
if provider == "bedrock":
    st.caption(
        "Bedrock credentials: use system credentials from `.aws`/profile/role, "
        "or disable `Use system key` and provide AWS keys manually. "
        "In manual mode, this app injects the keys at runtime (no shell export needed)."
    )
    if not use_system_key:
        bedrock_access_key = st.text_input(
            "AWS access key ID",
            key="llm_bedrock_access_key",
            placeholder="AKIA...",
            help="Required in manual mode. No additional env setup is needed.",
        )
        bedrock_secret_key = st.text_input(
            "AWS secret access key",
            type="password",
            key="llm_bedrock_secret_key",
            placeholder="Paste your AWS secret key",
            help="Required in manual mode. No additional env setup is needed.",
        )
if provider == "openai" and use_system_key and not resolve_api_key(provider):
    st.caption("No OPENAI_API_KEY found. Provide a key to use GPT.")
if use_system_key:
    _set_runtime_key(provider, None)
else:
    _set_runtime_key(provider, None if provider == "bedrock" else api_key)
_set_runtime_bedrock_credentials(
    provider,
    use_system_key,
    bedrock_access_key,
    bedrock_secret_key,
)
effective_api_key = ""
if provider != "bedrock":
    effective_api_key = "" if use_system_key else (api_key or _get_runtime_key(provider) or "")
if provider != "bedrock" and not use_system_key and effective_api_key and not api_key:
    st.session_state["llm_api_key"] = effective_api_key
_save_llm_config(
    provider,
    model_name,
    use_system_key,
    api_key=effective_api_key,
    bedrock_access_key=bedrock_access_key,
    bedrock_secret_key=bedrock_secret_key,
)

llm_signature = json.dumps(
    {
        "provider": provider,
        "model": model_name,
        "use_system_key": bool(use_system_key),
        "api_key": "" if use_system_key else effective_api_key,
        "bedrock_access_key": (
            ""
            if use_system_key or provider != "bedrock"
            else str(bedrock_access_key or "")
        ),
        "bedrock_secret_key": (
            ""
            if use_system_key or provider != "bedrock"
            else str(bedrock_secret_key or "")
        ),
    },
    sort_keys=True,
)
if "llm_saved_signature" not in st.session_state:
    st.session_state["llm_saved_signature"] = ""
if "llm_settings_saved_confirmed" not in st.session_state:
    st.session_state["llm_settings_saved_confirmed"] = False
if st.session_state.get("llm_saved_signature") != llm_signature:
    st.session_state["llm_settings_saved_confirmed"] = False

if st.button("Save LLM settings", key="save_llm_settings"):
    if (
        provider == "bedrock"
        and not use_system_key
        and (
            not str(bedrock_access_key or "").strip()
            or not str(bedrock_secret_key or "").strip()
        )
    ):
        st.error(
            "Bedrock manual mode requires both AWS access key ID and AWS secret access key."
        )
    else:
        _save_llm_config(
            provider,
            model_name,
            use_system_key,
            api_key=effective_api_key,
            bedrock_access_key=bedrock_access_key,
            bedrock_secret_key=bedrock_secret_key,
        )
        st.session_state["llm_saved_signature"] = llm_signature
        st.session_state["llm_settings_saved_confirmed"] = True
        st.success("Global LLM settings saved.")

runtime_doc_sections = _localized_default_doc_sections(language)
runtime_functional_sections = _localized_default_functional_sections(language)
doc_sections_errors: list[str] = []
functional_sections_errors: list[str] = []
doc_sections_pending_confirmation = False
functional_sections_pending_confirmation = False
needs_technical_generation = generation_mode in {GEN_MODE_TECHNICAL_ONLY, GEN_MODE_TECHNICAL_AND_FUNCTIONAL}
needs_functional_generation = generation_mode in {
    GEN_MODE_TECHNICAL_AND_FUNCTIONAL,
    GEN_MODE_FUNCTIONAL_ONLY,
}
functional_generation_requested = (
    needs_functional_generation
    and not (
        generation_mode == GEN_MODE_FUNCTIONAL_ONLY
        and source_type == SOURCE_FUNCTIONAL_LIBRARY
    )
)

if source_type != SOURCE_TECH_LIBRARY and needs_technical_generation:
    _render_doc_sections_editor(language=language)
    runtime_doc_sections, doc_sections_errors = _resolve_runtime_doc_sections()
    doc_sections_pending_confirmation = _sections_pending_confirmation(functional=False)
if functional_generation_requested and not (
    source_type == SOURCE_TECH_LIBRARY and generation_mode == GEN_MODE_TECHNICAL_AND_FUNCTIONAL
):
    _render_functional_sections_editor(language=language)
    runtime_functional_sections, functional_sections_errors = _resolve_runtime_functional_sections()
    functional_sections_pending_confirmation = _sections_pending_confirmation(functional=True)


# -------- Action section --------
if not st.session_state.get("llm_settings_saved_confirmed", False):
    st.info("Save LLM settings before loading a repository.")


def _create_doc_progress_callback():
    doc_progress_ui = {"bar": None, "status": None, "cost": None}

    def _on_doc_generation_progress(event: dict):
        if not isinstance(event, dict):
            return

        is_multi_pass = bool(event.get("is_multi_pass"))
        total_calls = max(1, int(event.get("total_calls") or 1))
        current_call = max(1, int(event.get("current_call") or 1))
        phase = str(event.get("phase") or "generation")
        event_name = str(event.get("event") or "")
        event_message = str(event.get("message") or "").strip()
        if event_message:
            msg_base = event_message
        elif is_multi_pass:
            msg_base = "Large repository detected. Documentation generation may take a few minutes."
        else:
            msg_base = "Generating documentation in multiple steps."

        def _phase_label(value: str) -> str:
            if value == "final_cleanup":
                return "final polish"
            if value == "functional_final_cleanup":
                return "functional final polish"
            if value in {"section_writing", "generation"}:
                return "section writing"
            if value == "functional_section_writing":
                return "functional section writing"
            if value == "final_writing":
                return "final writing"
            if value == "evidence_extraction":
                return "evidence extraction"
            if value == "functional_chunk_extraction":
                return "functional chunk extraction"
            return "documentation step"

        if doc_progress_ui["status"] is None:
            doc_progress_ui["status"] = st.empty()
        if doc_progress_ui["bar"] is None:
            doc_progress_ui["bar"] = st.progress(0)
        if doc_progress_ui["cost"] is None:
            doc_progress_ui["cost"] = st.empty()

        total_cost_usd = event.get("total_cost_usd")
        call_cost_usd = event.get("call_cost_usd")
        cost_available = bool(event.get("cost_available"))

        def _render_cost_line():
            if not cost_available:
                return
            try:
                total_val = float(total_cost_usd or 0.0)
            except Exception:
                return
            call_suffix = ""
            try:
                if call_cost_usd is not None:
                    call_val = float(call_cost_usd)
                    call_suffix = f" (+${call_val:.6f} last call)"
            except Exception:
                call_suffix = ""
            doc_progress_ui["cost"].caption(
                f"Estimated generation cost: ${total_val:.6f} USD{call_suffix}"
            )

        if event_name == "plan":
            doc_progress_ui["status"].info(msg_base)
            doc_progress_ui["bar"].progress(0)
            _render_cost_line()
            return

        if event_name == "call_start":
            pct = min(99, max(0, int(((current_call - 1) / total_calls) * 100)))
            phase_label = _phase_label(phase)
            doc_progress_ui["bar"].progress(pct)
            doc_progress_ui["status"].info(
                f"{msg_base} Processing {phase_label} {current_call}/{total_calls}..."
            )
            _render_cost_line()
            return

        if event_name == "call_end":
            pct = min(99, max(1, int((current_call / total_calls) * 100)))
            phase_label = _phase_label(phase)
            doc_progress_ui["bar"].progress(pct)
            doc_progress_ui["status"].info(
                f"{msg_base} Completed {phase_label} {current_call}/{total_calls}."
            )
            _render_cost_line()
            return

        if event_name == "done":
            doc_progress_ui["bar"].progress(100)
            doc_progress_ui["status"].success("Documentation generation completed.")
            _render_cost_line()

    return _on_doc_generation_progress


retry_state = st.session_state.get("doc_generation_retry")
if source_type != SOURCE_TECH_LIBRARY and isinstance(retry_state, dict):
    retry_repo_name = str(retry_state.get("repo_name") or "repository")
    st.warning(
        f"Previous documentation generation failed for {retry_repo_name}. "
        "You can retry from the last successful LLM call."
    )
    if st.button(
        "Retry documentation from last successful call",
        key="retry_doc_generation",
    ):
        with st.spinner("Retrying documentation from last successful call..."):
            try:
                retry_type = str(retry_state.get("retry_type") or "technical").strip().lower()
                code_json_path = Path(str(retry_state.get("code_json_path") or ""))
                checkpoint_path = Path(str(retry_state.get("checkpoint_path") or ""))
                if not code_json_path.exists():
                    raise FileNotFoundError("code.json not found for retry.")
                if not checkpoint_path.exists():
                    raise FileNotFoundError("Resume checkpoint not found.")

                retry_language = _normalize_doc_language(retry_state.get("language") or language)
                retry_provider = str(retry_state.get("provider") or provider)
                retry_model_name = str(retry_state.get("model_name") or model_name)
                retry_use_system_key = bool(retry_state.get("use_system_key", use_system_key))
                retry_api_key = "" if retry_use_system_key else str(retry_state.get("api_key") or "")
                retry_sections = retry_state.get("documentation_sections")
                retry_functional_sections = retry_state.get("functional_sections")
                if not isinstance(retry_sections, dict) or not retry_sections:
                    retry_sections = _localized_default_doc_sections(retry_language)
                if not isinstance(retry_functional_sections, dict) or not retry_functional_sections:
                    retry_functional_sections = _localized_default_functional_sections(retry_language)

                if retry_type == "functional_code":
                    generate_functional_doc(
                        str(code_json_path),
                        "functional_documentation.md",
                        retry_language,
                        provider=retry_provider,
                        model_name=retry_model_name,
                        api_key=retry_api_key,
                        use_system_key=retry_use_system_key,
                        progress_callback=_create_doc_progress_callback(),
                        functional_sections=retry_functional_sections,
                        resume_from_checkpoint=True,
                        checkpoint_path=str(checkpoint_path),
                    )
                    separate_output(
                        "functional_documentation.md",
                        documentation_sections=retry_functional_sections,
                        output_docs_dir="docs_functional",
                        keep_chat_file=False,
                    )
                elif retry_type == "functional":
                    technical_md_path = Path(
                        str(retry_state.get("technical_md_path") or (Path.cwd() / "documentation.md"))
                    )
                    if not technical_md_path.exists():
                        raise FileNotFoundError("Technical documentation source not found for functional retry.")
                    generate_functional_doc_from_technical(
                        str(technical_md_path),
                        "functional_documentation.md",
                        retry_language,
                        provider=retry_provider,
                        model_name=retry_model_name,
                        api_key=retry_api_key,
                        use_system_key=retry_use_system_key,
                        progress_callback=_create_doc_progress_callback(),
                        functional_sections=retry_functional_sections,
                        resume_from_checkpoint=True,
                        checkpoint_path=str(checkpoint_path),
                    )
                    separate_output(
                        "functional_documentation.md",
                        documentation_sections=retry_functional_sections,
                        output_docs_dir="docs_functional",
                        keep_chat_file=False,
                    )
                else:
                    generate_doc(
                        str(code_json_path),
                        "documentation.md",
                        retry_language,
                        provider=retry_provider,
                        model_name=retry_model_name,
                        api_key=retry_api_key,
                        use_system_key=retry_use_system_key,
                        progress_callback=_create_doc_progress_callback(),
                        documentation_sections=retry_sections,
                        resume_from_checkpoint=True,
                        checkpoint_path=str(checkpoint_path),
                    )
                    separate_output(
                        "documentation.md",
                        documentation_sections=retry_sections,
                    )
                    if _normalize_generation_mode_value(retry_state.get("generation_mode")) == GEN_MODE_TECHNICAL_AND_FUNCTIONAL:
                        functional_checkpoint = Path(
                            str(
                                retry_state.get("functional_checkpoint_path")
                                or (Path(str(retry_state.get("output_dir") or code_json_path.parent))
                                    / "functional_doc_gen_resume.json")
                            )
                        )
                        generate_functional_doc_from_technical(
                            "documentation.md",
                            "functional_documentation.md",
                            retry_language,
                            provider=retry_provider,
                            model_name=retry_model_name,
                            api_key=retry_api_key,
                            use_system_key=retry_use_system_key,
                            progress_callback=_create_doc_progress_callback(),
                            functional_sections=retry_functional_sections,
                            checkpoint_path=str(functional_checkpoint),
                        )
                        separate_output(
                            "functional_documentation.md",
                            documentation_sections=retry_functional_sections,
                            output_docs_dir="docs_functional",
                            keep_chat_file=False,
                        )

                retry_generation_mode = _normalize_generation_mode_value(
                    retry_state.get("generation_mode")
                )
                if retry_type in {"functional", "functional_code"} and retry_generation_mode != GEN_MODE_TECHNICAL_AND_FUNCTIONAL:
                    if not _activate_local_functional_workspace():
                        st.error("Failed to activate functional documentation in workspace after retry.")
                        st.stop()

                retry_repo_name = str(retry_state.get("repo_name") or code_json_path.parent.name)
                retry_repo_path = str(retry_state.get("repo_path") or "")
                retry_owner = str(retry_state.get("owner") or "")
                retry_output_dir = str(retry_state.get("output_dir") or code_json_path.parent)
                retry_source = retry_state.get("source")
                retry_source_type = retry_state.get("source_type")
                retry_graph_hint = retry_state.get("graph_path") or str(
                    Path(retry_output_dir) / "graphs" / "graph.json"
                )

                generate_mkdocs_config(
                    project_root=Path.cwd(),
                    repo_name=retry_repo_name,
                    repo_url=retry_source if retry_source_type == "Git URL" else None,
                    author=retry_owner,
                )

                port, err = _ensure_mkdocs_port(force_restart=True)
                if err:
                    st.error("MkDocs failed to start.")
                    st.code(err)
                    port = None

                graph_path, graph_data = _load_graph_payload(retry_graph_hint)

                retry_mode_norm = _normalize_generation_mode_value(retry_state.get("generation_mode"))
                if not (
                    retry_type in {"functional", "functional_code"}
                    and retry_mode_norm != GEN_MODE_TECHNICAL_AND_FUNCTIONAL
                ):
                    snapshot_payload = _snapshot_repo_assets(
                        repo_name=retry_repo_name,
                        code_json_path=code_json_path,
                        language=retry_language,
                        graph_path=graph_path,
                        include_functional=False,
                    )
                    _upsert_repo_library_entry(
                        retry_repo_name,
                        {
                            "repo_path": retry_repo_path,
                            "owner": retry_owner,
                            "output_dir": retry_output_dir,
                            "graph_path": graph_path,
                            "repo_url": retry_source if retry_source_type == "Git URL" else None,
                            "source": retry_source,
                            "source_type": retry_source_type,
                            "language": retry_language,
                            "docs_available": bool(snapshot_payload.get("docs_available")),
                            "functional_docs_available": False,
                            **snapshot_payload,
                        },
                        language=retry_language,
                    )
                should_snapshot_functional = bool(
                    str(retry_type or "") == "functional"
                    or str(retry_type or "") == "functional_code"
                    or _normalize_generation_mode_value(retry_state.get("generation_mode")) == GEN_MODE_TECHNICAL_AND_FUNCTIONAL
                )
                functional_snapshot_payload: dict = {}
                if should_snapshot_functional:
                    functional_snapshot_payload = _snapshot_functional_assets(
                        repo_name=retry_repo_name,
                        language=retry_language,
                        graph_path=graph_path,
                        code_json_path=code_json_path,
                    )
                    _upsert_functional_repo_library_entry(
                        retry_repo_name,
                        {
                            "repo_path": retry_repo_path,
                            "owner": retry_owner,
                            "output_dir": retry_output_dir,
                            "graph_path": graph_path,
                            "repo_url": retry_source if retry_source_type == "Git URL" else None,
                            "source": retry_source,
                            "source_type": retry_source_type,
                            "language": retry_language,
                            "docs_available": bool(functional_snapshot_payload.get("docs_available")),
                            "functional_docs_available": bool(
                                functional_snapshot_payload.get("functional_docs_available")
                            ),
                            **functional_snapshot_payload,
                        },
                        language=retry_language,
                    )

                st.session_state["repo_name"] = retry_repo_name
                st.session_state["repo_language"] = retry_language
                _save_last_repo(retry_repo_name)
                st.session_state["repo_state"] = {
                    "repo_path": retry_repo_path,
                    "repo_name": retry_repo_name,
                    "owner": retry_owner,
                    "output_dir": retry_output_dir,
                    "language": retry_language,
                    "result": retry_state.get("result") if isinstance(retry_state.get("result"), dict) else {},
                    "graph_path": graph_path,
                    "graph_data": graph_data,
                    "docs_generated": True,
                    "docs_skipped": False,
                    "functional_docs_generated": bool(
                        should_snapshot_functional
                        and functional_snapshot_payload.get("functional_docs_available")
                    ),
                    "doc_variant": (
                        "functional"
                        if retry_type in {"functional", "functional_code"}
                        and _normalize_generation_mode_value(retry_state.get("generation_mode")) != GEN_MODE_TECHNICAL_AND_FUNCTIONAL
                        else "technical"
                    ),
                    "generation_mode": str(retry_state.get("generation_mode") or "Technical only"),
                    "mkdocs_port": port,
                    "docs_nonce": uuid.uuid4().hex,
                }
                st.session_state.pop("doc_generation_retry", None)
                _safe_rerun()
            except Exception as retry_err:
                st.error(f"Failed to resume documentation generation: {retry_err}")

load_button_label = "Load repository"
if generation_mode == GEN_MODE_TECHNICAL_AND_FUNCTIONAL:
    load_button_label = "Generate technical + functional docs"
elif generation_mode == GEN_MODE_FUNCTIONAL_ONLY:
    if source_type == SOURCE_FUNCTIONAL_LIBRARY:
        load_button_label = "Load functional docs"
    else:
        load_button_label = "Generate functional docs"

if st.button(load_button_label, disabled=not st.session_state.get("llm_settings_saved_confirmed", False)):
    if (
        provider == "bedrock"
        and not use_system_key
        and (
            not str(bedrock_access_key or "").strip()
            or not str(bedrock_secret_key or "").strip()
        )
    ):
        st.error(
            "Bedrock manual mode requires both AWS access key ID and AWS secret access key."
        )
    elif source_type == SOURCE_TECH_LIBRARY and generation_mode == GEN_MODE_FUNCTIONAL_ONLY:
        st.error(
            "Functional-only mode supports Git URL, ZIP file, Local folder, or Functional Library."
        )
    elif source_type == SOURCE_TECH_LIBRARY and generation_mode == GEN_MODE_TECHNICAL_AND_FUNCTIONAL:
        st.error(
            "Technical + Functional mode is available for Git URL, ZIP, or Local folder sources. "
            "For Technical Library, use Technical only or Functional only."
        )
    elif source_type != SOURCE_TECH_LIBRARY and needs_technical_generation and doc_sections_pending_confirmation:
        st.error("Click 'Confirm custom sections' before generating documentation.")
    elif source_type != SOURCE_TECH_LIBRARY and needs_technical_generation and doc_sections_errors:
        for err in doc_sections_errors:
            st.error(err)
    elif functional_generation_requested and functional_sections_pending_confirmation:
        st.error("Click 'Confirm custom sections' before generating functional documentation.")
    elif functional_generation_requested and functional_sections_errors:
        for err in functional_sections_errors:
            st.error(err)
    elif source_type == SOURCE_FUNCTIONAL_LIBRARY and generation_mode == GEN_MODE_FUNCTIONAL_ONLY:
        if not selected_functional_library_repo_name:
            st.error("Select a saved repository from functional library.")
        else:
            selected_target_lang = _normalize_doc_language(language)
            source_lang_for_load = selected_target_lang
            should_translate_functional = False
            if not isinstance(selected_functional_target_lang_entry, dict):
                source_entry = selected_functional_source_lang_entry
                if not isinstance(source_entry, dict):
                    st.error("Selected repository has no saved functional documentation to load or translate.")
                    st.stop()
                source_lang_for_load = _normalize_doc_language(source_entry.get("language"))
                if (
                    source_lang_for_load != selected_target_lang
                    and not bool(translate_functional_library_docs_on_load)
                ):
                    st.error(
                        "Selected repository has no saved functional documentation in the selected language. "
                        "Enable 'Auto-translate functional docs on load' or choose an available language."
                    )
                    st.stop()
                should_translate_functional = (
                    source_lang_for_load != selected_target_lang
                    and bool(translate_functional_library_docs_on_load)
                )
            with st.spinner("Loading functional documentation from functional library..."):
                repo_state, err = _open_repo_from_functional_library(
                    selected_functional_library_repo_name,
                    preferred_language=source_lang_for_load,
                    require_docs=True,
                    start_mkdocs=False,
                )
            if err:
                st.error(f"Failed to open functional documentation from functional library: {err}")
                st.session_state.pop("repo_state", None)
            else:
                translated_functional = False
                translated_from_lang = source_lang_for_load
                merged_functional_entry = None
                if should_translate_functional:
                    source_md_candidates = [
                        Path.cwd() / "functional_documentation.md",
                        Path.cwd() / "documentation.md",
                    ]
                    source_md_path = next((p for p in source_md_candidates if p.exists()), None)
                    if source_md_path is None:
                        st.error("Functional documentation file not found for translation.")
                        st.session_state.pop("repo_state", None)
                        st.stop()
                    try:
                        with st.spinner(
                            f"Translating functional documentation from {source_lang_for_load} "
                            f"to {selected_target_lang}..."
                        ):
                            translated_md = _translate_documentation_with_llm(
                                documentation_md=source_md_path.read_text(encoding="utf-8"),
                                target_language=selected_target_lang,
                                provider=provider,
                                model_name=model_name,
                                api_key=effective_api_key,
                                use_system_key=use_system_key,
                                temperature=0.0,
                            )
                            Path.cwd().joinpath("functional_documentation.md").write_text(
                                translated_md,
                                encoding="utf-8",
                            )
                            translated_sections = _derive_sections_from_marked_documentation(
                                translated_md,
                                runtime_functional_sections,
                            )
                            separate_output(
                                "functional_documentation.md",
                                documentation_sections=translated_sections,
                                output_docs_dir="docs_functional",
                                keep_chat_file=False,
                            )
                        if not _activate_local_functional_workspace():
                            st.error("Failed to activate translated functional documentation in workspace.")
                            st.session_state.pop("repo_state", None)
                            st.stop()

                        repo_name_for_snapshot = (
                            repo_state.get("repo_name") or selected_functional_library_repo_name
                        )
                        code_json_path = Path(repo_state.get("output_dir") or "") / "code.json"
                        if not code_json_path.exists() and repo_name_for_snapshot:
                            resolved_code = _resolve_repo_code_json(
                                repo_name=repo_name_for_snapshot,
                                preferred_language=source_lang_for_load,
                            )
                            if resolved_code:
                                code_json_path = resolved_code
                        if not code_json_path.exists():
                            entry_code_json = Path(
                                str((repo_state.get("library_entry") or {}).get("library_code_json") or "")
                            )
                            if entry_code_json.exists():
                                code_json_path = entry_code_json

                        source_meta = repo_state.get("library_entry") or selected_functional_library_entry
                        if repo_name_for_snapshot:
                            snapshot_payload = _snapshot_functional_assets(
                                repo_name=repo_name_for_snapshot,
                                language=selected_target_lang,
                                graph_path=repo_state.get("graph_path"),
                                code_json_path=code_json_path if code_json_path.exists() else None,
                            )
                            output_dir_for_entry = (
                                str(code_json_path.parent)
                                if code_json_path.exists()
                                else str(repo_state.get("output_dir") or "")
                            )
                            _, merged_functional_entry = _upsert_functional_repo_library_entry(
                                repo_name_for_snapshot,
                                {
                                    "repo_path": repo_state.get("repo_path"),
                                    "owner": repo_state.get("owner"),
                                    "output_dir": output_dir_for_entry,
                                    "graph_path": repo_state.get("graph_path"),
                                    "repo_url": source_meta.get("repo_url"),
                                    "source": source_meta.get("source"),
                                    "source_type": SOURCE_FUNCTIONAL_LIBRARY,
                                    "language": selected_target_lang,
                                    "docs_available": bool(snapshot_payload.get("docs_available")),
                                    "functional_docs_available": bool(
                                        snapshot_payload.get("functional_docs_available")
                                    ),
                                    **snapshot_payload,
                                },
                                language=selected_target_lang,
                            )

                        repo_state["language"] = selected_target_lang
                        translated_functional = True
                    except Exception as translate_err:
                        st.error(f"Failed to translate functional documentation: {translate_err}")
                        st.session_state.pop("repo_state", None)
                        st.stop()

                st.session_state["repo_state"] = repo_state
                loaded_name = repo_state.get("repo_name") or selected_functional_library_repo_name
                if isinstance(merged_functional_entry, dict):
                    repo_state["library_entry"] = merged_functional_entry
                active_library_entry = repo_state.get("library_entry") or selected_functional_library_entry
                active_lang = selected_target_lang if translated_functional else source_lang_for_load
                generate_mkdocs_config(
                    project_root=Path.cwd(),
                    repo_name=loaded_name,
                    repo_url=(active_library_entry or {}).get("repo_url"),
                    author=repo_state.get("owner"),
                )
                port, mkdocs_err = _ensure_mkdocs_port(force_restart=True)
                if mkdocs_err:
                    st.error("MkDocs failed to start.")
                    st.code(mkdocs_err)
                    st.session_state.pop("repo_state", None)
                    st.stop()
                repo_state["mkdocs_port"] = port
                repo_state["docs_generated"] = True
                repo_state["docs_skipped"] = False
                repo_state["functional_docs_generated"] = True
                repo_state["doc_variant"] = "functional"
                repo_state["generation_mode"] = generation_mode
                repo_state["docs_nonce"] = uuid.uuid4().hex
                repo_state["language"] = active_lang
                if translated_functional:
                    st.session_state["repo_language"] = active_lang
                    st.session_state["language"] = active_lang
                    st.success(
                        "Repository loaded from functional library and documentation translated: "
                        f"{loaded_name} [{translated_from_lang} -> {selected_target_lang}]"
                    )
                else:
                    st.session_state["repo_language"] = active_lang
                    st.session_state["language"] = active_lang
                    st.success(
                        "Repository loaded from functional library: "
                        f"{loaded_name} [{active_lang}]"
                    )
                st.session_state.pop("doc_generation_retry", None)
                _safe_rerun()
    elif source_type == SOURCE_TECH_LIBRARY and generation_mode == "Functional only (from existing technical doc)":
        if not selected_library_repo_name:
            st.error("Select a saved repository from library.")
        else:
            selected_target_lang = _normalize_doc_language(language)
            with st.spinner("Loading technical documentation from library..."):
                repo_state, err = _open_repo_from_library(
                    selected_library_repo_name,
                    preferred_language=selected_target_lang,
                    doc_variant="technical",
                    require_docs=True,
                    start_mkdocs=False,
                )
            if err:
                st.error(f"Failed to open technical documentation from library: {err}")
                st.session_state.pop("repo_state", None)
            else:
                docs_md_path = Path.cwd() / "documentation.md"
                if not docs_md_path.exists():
                    st.error("Technical documentation file not found for functional generation.")
                    st.session_state.pop("repo_state", None)
                elif "<!-- SECTION:" not in docs_md_path.read_text(encoding="utf-8"):
                    st.error(
                        "Saved technical documentation is incomplete for functional derivation. "
                        "Load a repository with a valid generated technical documentation."
                    )
                    st.session_state.pop("repo_state", None)
                else:
                    try:
                        progress_cb = _create_doc_progress_callback()
                        generate_functional_doc_from_technical(
                            str(docs_md_path),
                            "functional_documentation.md",
                            selected_target_lang,
                            provider=provider,
                            model_name=model_name,
                            api_key=effective_api_key,
                            use_system_key=use_system_key,
                            progress_callback=progress_cb,
                            functional_sections=runtime_functional_sections,
                            checkpoint_path=str(
                                Path(repo_state.get("output_dir") or Path.cwd())
                                / "functional_doc_gen_resume.json"
                            ),
                        )
                        separate_output(
                            "functional_documentation.md",
                            documentation_sections=runtime_functional_sections,
                            output_docs_dir="docs_functional",
                            keep_chat_file=False,
                        )
                    except Exception as functional_err:
                        st.error(f"Failed to generate functional documentation: {functional_err}")
                        st.session_state.pop("repo_state", None)
                    else:
                        repo_name_for_snapshot = repo_state.get("repo_name") or selected_library_repo_name
                        code_json_path = Path(repo_state.get("output_dir") or "") / "code.json"
                        if not code_json_path.exists() and repo_name_for_snapshot:
                            resolved_code = _resolve_repo_code_json(
                                repo_name=repo_name_for_snapshot,
                                preferred_language=selected_target_lang,
                            )
                            if resolved_code:
                                code_json_path = resolved_code
                        if not code_json_path.exists():
                            entry_code_json = Path(
                                str((repo_state.get("library_entry") or {}).get("library_code_json") or "")
                            )
                            if entry_code_json.exists():
                                code_json_path = entry_code_json

                        merged_functional_entry = None
                        source_meta = repo_state.get("library_entry") or selected_library_entry
                        if repo_name_for_snapshot:
                            snapshot_payload = _snapshot_functional_assets(
                                repo_name=repo_name_for_snapshot,
                                language=selected_target_lang,
                                graph_path=repo_state.get("graph_path"),
                                code_json_path=code_json_path if code_json_path.exists() else None,
                            )
                            output_dir_for_entry = (
                                str(code_json_path.parent)
                                if code_json_path.exists()
                                else str(repo_state.get("output_dir") or "")
                            )
                            _, merged_functional_entry = _upsert_functional_repo_library_entry(
                                repo_name_for_snapshot,
                                {
                                    "repo_path": repo_state.get("repo_path"),
                                    "owner": repo_state.get("owner"),
                                    "output_dir": output_dir_for_entry,
                                    "graph_path": repo_state.get("graph_path"),
                                    "repo_url": source_meta.get("repo_url"),
                                    "source": source_meta.get("source"),
                                    "source_type": SOURCE_TECH_LIBRARY,
                                    "language": selected_target_lang,
                                    "docs_available": bool(snapshot_payload.get("docs_available")),
                                    "functional_docs_available": bool(
                                        snapshot_payload.get("functional_docs_available")
                                    ),
                                    **snapshot_payload,
                                },
                                language=selected_target_lang,
                            )
                            if not _activate_repo_assets(merged_functional_entry, doc_variant="functional"):
                                if not _activate_local_functional_workspace():
                                    st.error("Failed to activate functional docs after generation.")
                                    st.session_state.pop("repo_state", None)
                                    st.stop()
                        else:
                            if not _activate_local_functional_workspace():
                                st.error("Failed to activate generated functional docs in workspace.")
                                st.session_state.pop("repo_state", None)

                        effective_repo_name = repo_name_for_snapshot or selected_library_repo_name
                        generate_mkdocs_config(
                            project_root=Path.cwd(),
                            repo_name=effective_repo_name,
                            repo_url=source_meta.get("repo_url"),
                            author=repo_state.get("owner"),
                        )
                        port, err = _ensure_mkdocs_port(force_restart=True)
                        if err:
                            st.error("MkDocs failed to start.")
                            st.code(err)
                            port = None
                        repo_state["mkdocs_port"] = port
                        repo_state["language"] = selected_target_lang
                        repo_state["docs_generated"] = True
                        repo_state["docs_skipped"] = False
                        repo_state["functional_docs_generated"] = True
                        repo_state["doc_variant"] = "functional"
                        repo_state["generation_mode"] = generation_mode
                        repo_state["docs_nonce"] = uuid.uuid4().hex
                        if isinstance(merged_functional_entry, dict):
                            repo_state["library_entry"] = merged_functional_entry
                        st.session_state["repo_language"] = selected_target_lang
                        st.session_state["language"] = selected_target_lang
                        st.session_state["repo_state"] = repo_state
                        st.success(
                            f"Functional documentation generated from technical docs and saved to functional "
                            f"library: {effective_repo_name} [{selected_target_lang}]"
                        )
                        st.session_state.pop("doc_generation_retry", None)
                        _safe_rerun()
    elif source_type == SOURCE_TECH_LIBRARY:
        if not selected_library_repo_name:
            st.error("Select a saved repository from library.")
        else:
            with st.spinner("Loading repository from library..."):
                selected_target_lang = _normalize_doc_language(language)
                repo_state, err = _open_repo_from_library(
                    selected_library_repo_name,
                    preferred_language=selected_target_lang,
                )
                if err:
                    st.error(f"Failed to open from library: {err}")
                    st.session_state.pop("repo_state", None)
                else:
                    loaded_lang = _normalize_doc_language(repo_state.get("language"))
                    did_translate_docs = False
                    repo_state["doc_variant"] = "technical"

                    if translate_library_docs_on_load and loaded_lang != selected_target_lang:
                        docs_md_path = Path.cwd() / "documentation.md"
                        if not docs_md_path.exists():
                            st.warning(
                                "Saved documentation file not found for translation; loading original language content."
                            )
                        else:
                            try:
                                with st.spinner(
                                    f"Translating documentation from {loaded_lang} to {selected_target_lang}..."
                                ):
                                    translated_md = _translate_documentation_with_llm(
                                        documentation_md=docs_md_path.read_text(encoding="utf-8"),
                                        target_language=selected_target_lang,
                                        provider=provider,
                                        model_name=model_name,
                                        api_key=effective_api_key,
                                        use_system_key=use_system_key,
                                        temperature=0.0,
                                    )
                                    docs_md_path.write_text(translated_md, encoding="utf-8")
                                    translated_sections = _derive_sections_from_marked_documentation(
                                        translated_md,
                                        runtime_doc_sections,
                                    )
                                    separate_output(
                                        "documentation.md",
                                        documentation_sections=translated_sections,
                                    )

                                repo_name_for_snapshot = repo_state.get("repo_name") or (
                                    selected_library_repo_name or ""
                                )
                                code_json_path = Path(repo_state.get("output_dir") or "") / "code.json"
                                if not code_json_path.exists() and repo_name_for_snapshot:
                                    resolved_code = _resolve_repo_code_json(
                                        repo_name_for_snapshot,
                                        preferred_language=loaded_lang,
                                    )
                                    if resolved_code:
                                        code_json_path = resolved_code

                                if repo_name_for_snapshot and code_json_path.exists():
                                    source_meta = repo_state.get("library_entry") or selected_library_entry
                                    snapshot_payload = _snapshot_repo_assets(
                                        repo_name=repo_name_for_snapshot,
                                        code_json_path=code_json_path,
                                        language=selected_target_lang,
                                        graph_path=repo_state.get("graph_path"),
                                        include_functional=bool(
                                            (repo_state.get("library_entry") or {}).get("functional_docs_available")
                                        ),
                                    )
                                    _upsert_repo_library_entry(
                                        repo_name_for_snapshot,
                                        {
                                            "repo_path": repo_state.get("repo_path"),
                                            "owner": repo_state.get("owner"),
                                            "output_dir": str(code_json_path.parent),
                                            "graph_path": repo_state.get("graph_path"),
                                            "repo_url": source_meta.get("repo_url"),
                                            "source": source_meta.get("source"),
                                            "source_type": source_meta.get("source_type")
                                            or SOURCE_TECH_LIBRARY,
                                            "language": selected_target_lang,
                                            "docs_available": bool(snapshot_payload.get("docs_available")),
                                            "functional_docs_available": bool(
                                                snapshot_payload.get("functional_docs_available")
                                            ),
                                            **snapshot_payload,
                                        },
                                        language=selected_target_lang,
                                    )
                                did_translate_docs = True
                                repo_state["language"] = selected_target_lang
                                repo_state["docs_generated"] = True
                                repo_state["docs_skipped"] = False
                                st.session_state["repo_language"] = selected_target_lang
                                st.session_state["language"] = selected_target_lang
                            except Exception as translate_err:
                                st.error(f"Failed to translate documentation: {translate_err}")

                    active_technical_entry = repo_state.get("library_entry") or selected_library_entry
                    if (
                        bool(repo_state.get("docs_generated"))
                        and isinstance(active_technical_entry, dict)
                    ):
                        if not _activate_repo_assets(active_technical_entry, doc_variant="technical"):
                            st.error("Failed to activate technical documentation in workspace.")
                            st.session_state.pop("repo_state", None)
                            st.stop()

                    generate_mkdocs_config(
                        project_root=Path.cwd(),
                        repo_name=repo_state.get("repo_name") or selected_library_repo_name or "repository",
                        repo_url=(active_technical_entry or {}).get("repo_url"),
                        author=repo_state.get("owner"),
                    )
                    port, mkdocs_err = _ensure_mkdocs_port(force_restart=True)
                    if mkdocs_err:
                        st.error("MkDocs failed to start.")
                        st.code(mkdocs_err)
                        st.session_state.pop("repo_state", None)
                        st.stop()
                    repo_state["mkdocs_port"] = port
                    repo_state["doc_variant"] = "technical"
                    repo_state["generation_mode"] = generation_mode
                    repo_state["docs_nonce"] = uuid.uuid4().hex

                    st.session_state["repo_state"] = repo_state
                    loaded_name = repo_state.get("repo_name") or selected_library_repo_name
                    final_lang = _normalize_doc_language(repo_state.get("language"))
                    if did_translate_docs:
                        st.success(
                            f"Repository loaded from technical library and documentation translated: "
                            f"{loaded_name} [{loaded_lang} -> {final_lang}]"
                        )
                    else:
                        st.success(f"Repository loaded from technical library: {loaded_name} [{final_lang}]")
                    st.session_state.pop("doc_generation_retry", None)
                    _safe_rerun()
    else:
        if not source:
            st.error("Please provide a valid source.")
        else:
            with st.spinner("Loading repository..."):
                repo_path = ""
                repo_name = ""
                owner = ""
                output_dir = ""
                result = {}
                code_json_path: Path | None = None
                graph_path = ""
                resolved_api_key = effective_api_key
                try:
                    repo_path, repo_name, owner = load_repository(source)
                    output_dir = str(Path.cwd() / f"out/{repo_name}")
                    st.session_state["repo_name"] = repo_name
                    st.session_state["repo_language"] = _normalize_doc_language(language)
                    _save_last_repo(repo_name)

                    docs_generated = False
                    functional_docs_generated = False
                    resolved_api_key = effective_api_key

                    # Generate documentation + repo map
                    with st.spinner("Generating documentation..."):
                        result = repo_map(repo_path, output_dir)

                        code_json_path = Path(output_dir) / "code.json"

                        if needs_technical_generation:
                            generate_doc(
                                str(code_json_path),
                                "documentation.md",
                                language,
                                provider=provider,
                                model_name=model_name,
                                api_key=resolved_api_key,
                                use_system_key=use_system_key,
                                progress_callback=_create_doc_progress_callback(),
                                documentation_sections=runtime_doc_sections,
                            )
                            separate_output(
                                "documentation.md",
                                documentation_sections=runtime_doc_sections,
                            )
                            docs_generated = True

                        if generation_mode == GEN_MODE_TECHNICAL_AND_FUNCTIONAL:
                            generate_functional_doc_from_technical(
                                "documentation.md",
                                "functional_documentation.md",
                                language,
                                provider=provider,
                                model_name=model_name,
                                api_key=resolved_api_key,
                                use_system_key=use_system_key,
                                progress_callback=_create_doc_progress_callback(),
                                functional_sections=runtime_functional_sections,
                                checkpoint_path=str(Path(output_dir) / "functional_doc_gen_resume.json"),
                            )
                            separate_output(
                                "functional_documentation.md",
                                documentation_sections=runtime_functional_sections,
                                output_docs_dir="docs_functional",
                                keep_chat_file=False,
                            )
                            functional_docs_generated = True
                        elif generation_mode == GEN_MODE_FUNCTIONAL_ONLY:
                            generate_functional_doc(
                                str(code_json_path),
                                "functional_documentation.md",
                                language,
                                provider=provider,
                                model_name=model_name,
                                api_key=resolved_api_key,
                                use_system_key=use_system_key,
                                progress_callback=_create_doc_progress_callback(),
                                functional_sections=runtime_functional_sections,
                                checkpoint_path=str(
                                    Path(output_dir) / "functional_code_doc_gen_resume.json"
                                ),
                            )
                            separate_output(
                                "functional_documentation.md",
                                documentation_sections=runtime_functional_sections,
                                output_docs_dir="docs_functional",
                                keep_chat_file=False,
                            )
                            if not _activate_local_functional_workspace():
                                st.error("Failed to activate generated functional docs in workspace.")
                                st.stop()
                            functional_docs_generated = True

                        project_root = Path.cwd()

                        generate_mkdocs_config(
                            project_root=project_root,
                            repo_name=repo_name,
                            repo_url=source if source_type == "Git URL" else None,
                            author=owner,
                        )

                        port, err = _ensure_mkdocs_port(force_restart=True)
                        if err:
                            st.error("MkDocs failed to start.")
                            st.code(err)
                            port = None

                        graph_path = result.get("graph_file") or str(Path(output_dir) / "graphs" / "graph.json")
                        graph_path, graph_data = _load_graph_payload(graph_path)

                        technical_entry_key = None
                        technical_entry = None
                        if generation_mode != GEN_MODE_FUNCTIONAL_ONLY:
                            snapshot_payload = _snapshot_repo_assets(
                                repo_name=repo_name,
                                code_json_path=code_json_path,
                                language=language,
                                graph_path=graph_path,
                                include_functional=False,
                            )
                            technical_entry_key, technical_entry = _upsert_repo_library_entry(
                                repo_name,
                                {
                                    "repo_path": repo_path,
                                    "owner": owner,
                                    "output_dir": output_dir,
                                    "graph_path": graph_path,
                                    "repo_url": source if source_type == "Git URL" else None,
                                    "source": source,
                                    "source_type": source_type,
                                    "language": language,
                                    "docs_available": bool(snapshot_payload.get("docs_available")),
                                    "functional_docs_available": False,
                                    **snapshot_payload,
                                },
                                language=language,
                            )
                        functional_entry_key = None
                        functional_entry = None
                        if functional_docs_generated:
                            functional_snapshot_payload = _snapshot_functional_assets(
                                repo_name=repo_name,
                                language=language,
                                graph_path=graph_path,
                                code_json_path=code_json_path,
                            )
                            functional_entry_key, functional_entry = _upsert_functional_repo_library_entry(
                                repo_name,
                                {
                                    "repo_path": repo_path,
                                    "owner": owner,
                                    "output_dir": output_dir,
                                    "graph_path": graph_path,
                                    "repo_url": source if source_type == "Git URL" else None,
                                    "source": source,
                                    "source_type": source_type,
                                    "language": language,
                                    "docs_available": bool(functional_snapshot_payload.get("docs_available")),
                                    "functional_docs_available": bool(
                                        functional_snapshot_payload.get("functional_docs_available")
                                    ),
                                    **functional_snapshot_payload,
                                },
                                language=language,
                            )

                        st.session_state["repo_state"] = {
                            "repo_path": repo_path,
                            "repo_name": repo_name,
                            "owner": owner,
                            "output_dir": output_dir,
                            "language": language,
                            "result": result,
                            "graph_path": graph_path,
                            "graph_data": graph_data,
                            "docs_generated": bool(docs_generated or functional_docs_generated),
                            "docs_skipped": False,
                            "technical_docs_generated": bool(docs_generated),
                            "functional_docs_generated": functional_docs_generated,
                            "doc_variant": "functional" if (functional_docs_generated and not docs_generated) else "technical",
                            "library_entry_key": technical_entry_key,
                            "library_entry": technical_entry,
                            "technical_library_entry": technical_entry,
                            "functional_library_entry_key": functional_entry_key,
                            "functional_library_entry": functional_entry,
                            "mkdocs_port": port,
                            "docs_nonce": uuid.uuid4().hex,
                            "generation_mode": generation_mode,
                        }
                        st.session_state.pop("doc_generation_retry", None)
                        _safe_rerun()

                except Exception as e:
                    st.error(f"Failed to load repository: {e}")
                    st.session_state.pop("repo_state", None)
                    retry_checkpoint = Path(output_dir) / "doc_gen_resume.json" if output_dir else Path()
                    retry_functional_checkpoint = (
                        Path(output_dir) / "functional_doc_gen_resume.json" if output_dir else Path()
                    )
                    retry_functional_code_checkpoint = (
                        Path(output_dir) / "functional_code_doc_gen_resume.json" if output_dir else Path()
                    )
                    technical_md_path = Path.cwd() / "documentation.md"
                    has_retry_technical = (
                        code_json_path is not None
                        and code_json_path.exists()
                        and retry_checkpoint.exists()
                    )
                    has_retry_functional = (
                        code_json_path is not None
                        and code_json_path.exists()
                        and retry_functional_checkpoint.exists()
                        and technical_md_path.exists()
                    )
                    has_retry_functional_code = (
                        code_json_path is not None
                        and code_json_path.exists()
                        and retry_functional_code_checkpoint.exists()
                    )
                    if has_retry_technical or has_retry_functional or has_retry_functional_code:
                        if has_retry_functional_code:
                            retry_type = "functional_code"
                        elif has_retry_functional:
                            retry_type = "functional"
                        else:
                            retry_type = "technical"
                        checkpoint_value = (
                            retry_functional_code_checkpoint
                            if retry_type == "functional_code"
                            else (retry_functional_checkpoint if retry_type == "functional" else retry_checkpoint)
                        )
                        st.session_state["doc_generation_retry"] = {
                            "repo_path": repo_path,
                            "repo_name": repo_name,
                            "owner": owner,
                            "output_dir": output_dir,
                            "code_json_path": str(code_json_path),
                            "graph_path": graph_path,
                            "language": _normalize_doc_language(language),
                            "provider": provider,
                            "model_name": model_name,
                            "api_key": resolved_api_key,
                            "use_system_key": use_system_key,
                            "retry_type": retry_type,
                            "generation_mode": generation_mode,
                            "documentation_sections": runtime_doc_sections,
                            "functional_sections": runtime_functional_sections,
                            "technical_md_path": str(technical_md_path),
                            "functional_checkpoint_path": str(retry_functional_checkpoint),
                            "functional_code_checkpoint_path": str(retry_functional_code_checkpoint),
                            "include_functional_snapshot": bool(
                                generation_mode == GEN_MODE_TECHNICAL_AND_FUNCTIONAL
                                or retry_type in {"functional", "functional_code"}
                            ),
                            "source": source,
                            "source_type": source_type,
                            "result": result if isinstance(result, dict) else {},
                            "checkpoint_path": str(checkpoint_value),
                        }
                        st.info("Retry is available: use the button to continue from the last successful LLM call.")
                    else:
                        st.session_state.pop("doc_generation_retry", None)

if st.session_state.get("repo_state"):
    _render_loaded_repo(st.session_state["repo_state"])

